from flask import Flask, render_template, request, redirect, session, url_for, jsonify, flash, send_file, send_from_directory
import pymysql
from functools import wraps
from datetime import datetime, timedelta
import json
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side, NamedStyle
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os
from babel.dates import format_datetime
import locale
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from io import BytesIO
from datetime import datetime

#import logging
#from datetime import datetime

# Configuración del logging
#   logging.basicConfig(
#    filename='almacen.log',
#    level=logging.DEBUG,
#    format='%(asctime)s - %(levelname)s - %(message)s'
#)

# Configurar el locale en español
try:
    locale.setlocale(locale.LC_TIME, 'es_ES.UTF-8')
except:
    try:
        locale.setlocale(locale.LC_TIME, 'es_ES')
    except:
        locale.setlocale(locale.LC_TIME, 'Spanish')

def fecha_en_español(fecha):
    """Formatea una fecha en español con el formato: 'dd de MMMM del yyyy'"""
    meses = {
        'January': 'Enero',
        'February': 'Febrero',
        'March': 'Marzo',
        'April': 'Abril',
        'May': 'Mayo',
        'June': 'Junio',
        'July': 'Julio',
        'August': 'Agosto',
        'September': 'Septiembre',
        'October': 'Octubre',
        'November': 'Noviembre',
        'December': 'Diciembre'
    }
    
    fecha_str = fecha.strftime('%d de %B del %Y')
    for mes_en, mes_es in meses.items():
        fecha_str = fecha_str.replace(mes_en, mes_es)
    return fecha_str

app = Flask(__name__)
app.secret_key = 'tu_clave_secreta_aqui'

# Conexión a la base de datos
def conexion():
    return pymysql.connect(
        host="192.168.8.217",
        user="xd",
        passwd="xddd",
        db="AlmacenMunicipal"
    )


# Decorador para verificar login
def requiere_login(rol_requerido):
    def decorador(func):
        @wraps(func)
        def envoltura(*args, **kwargs):
            if 'usuario' not in session:
                return redirect(url_for('login'))
            
            rol_usuario = session.get('rol')
            
            # Si se requiere rol admin y el usuario no es admin, redirigir
            if rol_requerido == 'admin' and rol_usuario != 'admin':
                return redirect(url_for('inventario'))
                
            # Si se requiere rol normal, permitir acceso a normales y admins
            if rol_requerido == 'normal':
                return func(*args, **kwargs)
                
            return func(*args, **kwargs)
        return envoltura
    return decorador

@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        correo = request.form["correo"]
        password = request.form["password"]
        
        conn = conexion()
        try:
            with conn.cursor() as cursor:
                cursor.execute("""
                    SELECT id_usuario, nombre, correo, contraseña, rol 
                    FROM Usuario 
                    WHERE correo = %s AND contraseña = %s
                """, (correo, password))
                usuario = cursor.fetchone()
                
                if usuario:
                    session['usuario'] = usuario[1]
                    session['rol'] = usuario[4]
                    if usuario[4] == 'admin':
                        return redirect(url_for('index'))
                    else:
                        # Usuario normal va directo a la vista simplificada del inventario
                        return redirect(url_for('inventario_normal'))
                else:
                    return render_template("login.html", error="Credenciales incorrectas")
        except Exception as e:
            print(e)
            return render_template("login.html", error="Error al iniciar sesión")
        finally:
            conn.close()
    
    return render_template("login.html")

@app.route("/inventario_normal")
@requiere_login('normal')
def inventario_normal():
    conn = conexion()
    try:
        with conn.cursor() as cursor:
            # Obtener categorías
            cursor.execute("""
                SELECT DISTINCT 
                    c.id_categoria, 
                    c.nombre 
                FROM Categoria c
                INNER JOIN Producto p ON c.id_categoria = p.id_categoria
                ORDER BY c.nombre
            """)
            categorias = cursor.fetchall()
            
            # Obtener procedencias únicas
            cursor.execute("""
                SELECT DISTINCT procedencia 
                FROM Producto 
                WHERE procedencia IS NOT NULL 
                ORDER BY procedencia
            """)
            procedencias = [row[0] for row in cursor.fetchall()]
            
            # Obtener productos con toda la información
            cursor.execute("""
                SELECT 
                    p.id_producto,
                    p.codigo,
                    p.nombre,
                    p.descripcion,
                    p.cantidad,
                    p.unidad_medida,
                    c.nombre as categoria,
                    DATE_FORMAT(p.fecha_ingreso, '%d/%m/%Y %H:%i') as fecha_ingreso,
                    COALESCE(p.procedencia, 'No especificada') as procedencia,
                    CASE
                        WHEN p.cantidad <= 0 THEN 'agotado'
                        ELSE 'disponible'
                    END as estado
                FROM Producto p
                LEFT JOIN Categoria c ON p.id_categoria = c.id_categoria
                ORDER BY p.fecha_ingreso DESC, p.id_producto DESC
            """)
            productos = cursor.fetchall()
            
            return render_template("inventario_normal.html", 
                                 productos=productos,
                                 categorias=categorias,
                                 procedencias=procedencias)
    except Exception as e:
        print(f"Error al cargar inventario: {str(e)}")
        return render_template("inventario_normal.html", error="Error al cargar el inventario")
    finally:
        conn.close()

@app.route("/registro", methods=["GET", "POST"])
def registro():
    if request.method == "POST":
        nombre = request.form["nombre"]
        apellido = request.form["apellido"]
        correo = request.form["correo"]
        password = request.form["password"]
        
        conn = conexion()
        try:
            with conn.cursor() as cursor:
                cursor.execute("""
                    INSERT INTO Usuario (nombre, apellido, correo, contraseña, rol)
                    VALUES (%s, %s, %s, %s, 'normal')
                """, (nombre, apellido, correo, password))
                conn.commit()
                return redirect(url_for('login'))
        except Exception as e:
            print(e)
            return render_template("registro.html", error="Error al registrar usuario")
        finally:
            conn.close()
    
    return render_template("registro.html")

@app.route("/")
@app.route("/index")
@requiere_login('admin')
def index():
    return render_template("index.html")

@app.route("/inventario")
@requiere_login('admin')
def inventario():
    conn = conexion()
    try:
        with conn.cursor() as cursor:
            # Obtener categorías
            cursor.execute("""
                SELECT DISTINCT 
                    c.id_categoria, 
                    c.nombre 
                FROM Categoria c
                INNER JOIN Producto p ON c.id_categoria = p.id_categoria
                ORDER BY c.nombre
            """)
            categorias = cursor.fetchall()
            
            # Obtener procedencias únicas
            cursor.execute("""
                SELECT DISTINCT procedencia 
                FROM Producto 
                WHERE procedencia IS NOT NULL 
                ORDER BY procedencia
            """)
            procedencias = [row[0] for row in cursor.fetchall()]
            
            # Obtener productos con toda la información
            cursor.execute("""
                SELECT 
                    p.id_producto,
                    p.codigo,
                    p.nombre,
                    p.descripcion,
                    p.cantidad,
                    p.unidad_medida,
                    c.nombre as categoria,
                    p.fecha_ingreso,
                    COALESCE(p.procedencia, 'No especificada') as procedencia,
                    CASE
                        WHEN p.cantidad <= 0 AND EXISTS (
                            SELECT 1
                            FROM DetallePrestamo dp
                            JOIN Prestamo pr ON dp.id_prestamo = pr.id_prestamo
                            LEFT JOIN Devolucion d ON dp.id_detalle_prestamo = d.id_detalle_prestamo
                            WHERE dp.id_producto = p.id_producto
                            AND d.id_devolucion IS NULL
                        ) THEN 'prestado'
                        WHEN p.cantidad <= 0 THEN 'agotado'
                        ELSE 'disponible'
                    END as estado,
                    COALESCE(p.cantidad_inicial, p.cantidad) as cantidad_inicial,
                    (
                        SELECT COALESCE(SUM(dp.cantidad), 0)
                        FROM DetallePrestamo dp
                        JOIN Prestamo pr ON dp.id_prestamo = pr.id_prestamo
                        LEFT JOIN Devolucion d ON dp.id_detalle_prestamo = d.id_detalle_prestamo
                        WHERE dp.id_producto = p.id_producto
                        AND d.id_devolucion IS NULL
                    ) as prestamos_activos
                FROM Producto p
                LEFT JOIN Categoria c ON p.id_categoria = c.id_categoria
                ORDER BY p.fecha_ingreso DESC, p.id_producto DESC
            """)
            productos = cursor.fetchall()
            
            return render_template("inventario.html", 
                                 productos=productos,
                                 categorias=categorias,
                                 procedencias=procedencias)
    except Exception as e:
        print(f"Error al cargar inventario: {str(e)}")
        flash('Error al cargar el inventario', 'danger')
        return redirect(url_for('index'))
    finally:
        conn.close()
@app.route("/ingreso_producto", methods=["GET", "POST"])
@requiere_login('admin')
def ingreso_producto():
    unidades_medida = ['kg', 'l', 'm', 'g', 'ud', 'gal', 'lb', 'cm', 'mm', 'km', 'ml', 'oz', 'm²', 'm³', 'mg', 'dl', 'cl', 't', 'pulg', 'mts', 'u', 'min', 'h', 'seg', 'ha', 'hl', 'qq', 'dam', 'dm', 'dcm']
    if request.method == "POST":
        conn = conexion()
        try:
            with conn.cursor() as cursor:
                # Obtener la fecha del formulario o usar la fecha actual
                fecha_ingreso = datetime.strptime(
                    request.form.get('fecha_ingreso', datetime.now().strftime('%Y-%m-%dT%H:%M')),
                    '%Y-%m-%dT%H:%M'
                )
                
                cursor.execute("""
                    INSERT INTO Producto 
                    (nombre, descripcion, cantidad, unidad_medida, 
                     id_categoria, fecha_ingreso, procedencia, estado)
                    VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
                """, (
                    request.form['nombre'],
                    request.form['descripcion'],
                    float(request.form['cantidad']),
                    request.form['unidad_medida'],
                    request.form['categoria'],
                    fecha_ingreso,
                    request.form['procedencia'],
                    'disponible'
                ))
                
                id_producto = cursor.lastrowid
                
                # 2. Generar y actualizar el código
                # Modificamos la consulta para obtener el último número usado
                cursor.execute("""
                    SELECT MAX(CAST(
                        SUBSTRING_INDEX(codigo, '-', -1) AS UNSIGNED
                    )) as ultimo_numero
                    FROM Producto p
                    JOIN Categoria c ON p.id_categoria = c.id_categoria
                    WHERE c.id_categoria = %s
                    AND codigo REGEXP 'ALM-[A-Z]{3}-[0-9]{2}-[0-9]{4}'
                """, (request.form['categoria'],))
                
                ultimo_numero = cursor.fetchone()[0] or 0
                nuevo_numero = ultimo_numero + 1
                
                # Obtener el prefijo de la categoría
                cursor.execute("""
                    SELECT nombre FROM Categoria WHERE id_categoria = %s
                """, (request.form['categoria'],))
                categoria_nombre = cursor.fetchone()[0]
                categoria_prefijo = ''.join(e for e in categoria_nombre.upper() if e.isalnum())[:3]
                
                # Generar el nuevo código
                año_actual = str(datetime.now().year)[-2:]
                nuevo_codigo = f"ALM-{categoria_prefijo}-{año_actual}-{str(nuevo_numero).zfill(4)}"
                
                # Actualizar el código del producto
                cursor.execute("""
                    UPDATE Producto 
                    SET codigo = %s 
                    WHERE id_producto = %s
                """, (nuevo_codigo, id_producto))
                
                # 3. Registrar en historial
                detalle_historial = (
                    f"Ingreso inicial de {request.form['nombre']} - "
                    f"Cantidad: {float(request.form['cantidad'])} - "
                    f"Código asignado: {nuevo_codigo}"
                )
                
                cursor.execute("""
                    INSERT INTO HistorialModificaciones 
                    (id_producto, tipo_modificacion, cantidad, 
                     fecha_modificacion, detalle)
                    VALUES (%s, 'ingreso', %s, %s, %s)
                """, (id_producto, float(request.form['cantidad']), 
                      datetime.now(), detalle_historial))
                
                conn.commit()
                
                return render_template('ingreso_producto.html', 
                                    categorias=obtener_categorias(),
                                    mensaje=f"Producto registrado exitosamente con código: {nuevo_codigo}",
                                    tipo_mensaje="success",
                                    unidades_medida=unidades_medida,
                                    now=datetime.now())
                
        except Exception as e:
            conn.rollback()
            print("Error al ingresar producto:", str(e))
            return render_template('ingreso_producto.html', 
                                categorias=obtener_categorias(),
                                mensaje=f"Error al registrar el producto: {str(e)}",
                                tipo_mensaje="danger",
                                unidades_medida=unidades_medida,
                                now=datetime.now())
        finally:
            conn.close()
    else:
        return render_template('ingreso_producto.html', 
                             categorias=obtener_categorias(),
                             unidades_medida=unidades_medida,
                             now=datetime.now())

@app.route("/ingresar_producto", methods=["POST"])
@requiere_login('admin')
def ingresar_producto():
    conn = conexion()
    try:
        nombre = request.form['nombre']
        descripcion = request.form['descripcion']
        cantidad = float(request.form['cantidad'])
        unidad_medida = request.form['unidad_medida']
        id_categoria = request.form['categoria']
        fecha_ingreso = datetime.now()
        procedencia = request.form['procedencia']
        
        with conn.cursor() as cursor:
            cursor.execute("""
                INSERT INTO Producto (
                    nombre, descripcion, cantidad, cantidad_inicial, 
                    unidad_medida, fecha_ingreso, procedencia, 
                    estado, id_categoria, codigo
                ) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
            """, (
                nombre, descripcion, cantidad, cantidad,  # cantidad_inicial = cantidad
                unidad_medida, fecha_ingreso, procedencia,
                'disponible', id_categoria, None  # código se generará después
            ))
            
            id_producto = cursor.lastrowid
            
            # Generar código único
            cursor.execute("""
                SELECT nombre FROM Categoria WHERE id_categoria = %s
            """, (id_categoria,))
            categoria_nombre = cursor.fetchone()[0]
            categoria_prefijo = ''.join(e for e in categoria_nombre.upper() if e.isalnum())[:3]
            
            año_actual = str(datetime.now().year)[-2:]
            cursor.execute("""
                SELECT MAX(CAST(SUBSTRING_INDEX(codigo, '-', -1) AS UNSIGNED))
                FROM Producto
                WHERE codigo LIKE %s
            """, (f'ALM-{categoria_prefijo}-{año_actual}-%',))
            
            ultimo_numero = cursor.fetchone()[0] or 0
            nuevo_numero = str(ultimo_numero + 1).zfill(4)
            nuevo_codigo = f"ALM-{categoria_prefijo}-{año_actual}-{nuevo_numero}"
            
            # Actualizar el código del producto
            cursor.execute("""
                UPDATE Producto 
                SET codigo = %s 
                WHERE id_producto = %s
            """, (nuevo_codigo, id_producto))
            
            # Registrar en historial
            detalle_historial = (
                f"Ingreso inicial de {nombre} - "
                f"Cantidad: {cantidad} - "
                f"Código asignado: {nuevo_codigo}"
            )
            
            cursor.execute("""
                INSERT INTO HistorialModificaciones 
                (id_producto, tipo_modificacion, cantidad, 
                 fecha_modificacion, detalle)
                VALUES (%s, 'ingreso', %s, %s, %s)
            """, (id_producto, cantidad, datetime.now(), detalle_historial))
            
            conn.commit()
            return redirect(url_for('productos'))
            
    except Exception as e:
        conn.rollback()
        print("Error al ingresar producto:", str(e))
        return f"Error al ingresar el producto: {str(e)}"
    finally:
        conn.close()
def asignar_numero_secuencial(conn, tabla, id_campo, id_valor, fecha_campo, fecha_valor):
    """
    Asigna un número secuencial a un registro en la tabla especificada.
    Hace commit inmediato del número secuencial para que sea visible para otras transacciones
    y evitar condiciones de carrera.
    """
    import time
    import random
    
    try:
        # Usar un cursor separado para la asignación del número secuencial
        # Esto permite hacer commit inmediato sin afectar la transacción externa
        cursor = conn.cursor()
        
        try:
            # Verificar si ya existe un número asignado
            cursor.execute(f"""
                SELECT numero_secuencial 
                FROM {tabla} 
                WHERE {id_campo} = %s
            """, (id_valor,))
            
            resultado = cursor.fetchone()
            if resultado:
                cursor.close()
                return resultado[0]
            
            # Estrategia de reintentos con recálculo del máximo en cada intento
            max_intentos = 100
            for intento in range(max_intentos):
                try:
                    # Obtener el máximo actual del año específico
                    cursor.execute(f"""
                        SELECT COALESCE(MAX(numero_secuencial), 0)
                        FROM {tabla}
                        WHERE YEAR({fecha_campo}) = YEAR(%s)
                    """, (fecha_valor,))
                    
                    max_actual_año = cursor.fetchone()[0] or 0
                    siguiente_numero = max_actual_año + 1
                    
                    # Verificar si este número ya existe (puede ser de otro año si la restricción única es global)
                    cursor.execute(f"""
                        SELECT COUNT(*) 
                        FROM {tabla}
                        WHERE numero_secuencial = %s
                    """, (siguiente_numero,))
                    
                    existe = cursor.fetchone()[0] > 0
                    
                    # Si el número ya existe, usar el máximo global + 1
                    if existe:
                        cursor.execute(f"""
                            SELECT COALESCE(MAX(numero_secuencial), 0)
                            FROM {tabla}
                        """)
                        max_actual_global = cursor.fetchone()[0] or 0
                        siguiente_numero = max_actual_global + 1
                    
                    # Intentar insertar con el número calculado
                    cursor.execute(f"""
                        INSERT INTO {tabla} ({id_campo}, numero_secuencial, {fecha_campo})
                        VALUES (%s, %s, %s)
                    """, (id_valor, siguiente_numero, fecha_valor))
                    
                    # Hacer commit inmediato para que el número sea visible para otras transacciones
                    conn.commit()
                    
                    # Obtener el número asignado
                    cursor.execute(f"""
                        SELECT numero_secuencial 
                        FROM {tabla} 
                        WHERE {id_campo} = %s
                    """, (id_valor,))
                    
                    resultado = cursor.fetchone()
                    cursor.close()
                    if resultado:
                        return resultado[0]
                    else:
                        raise Exception("No se pudo obtener el número asignado")
                        
                except Exception as e:
                    error_str = str(e)
                    # Si es un error de duplicado, hacer rollback y reintentar
                    if "Duplicate entry" in error_str or "1062" in error_str:
                        conn.rollback()
                        
                        # Verificar nuevamente si ya existe (puede haber sido creado por otra transacción)
                        cursor.execute(f"""
                            SELECT numero_secuencial 
                            FROM {tabla} 
                            WHERE {id_campo} = %s
                        """, (id_valor,))
                        resultado = cursor.fetchone()
                        if resultado:
                            cursor.close()
                            return resultado[0]
                        
                        # Recalcular el máximo con la misma lógica que arriba
                        cursor.execute(f"""
                            SELECT COALESCE(MAX(numero_secuencial), 0)
                            FROM {tabla}
                            WHERE YEAR({fecha_campo}) = YEAR(%s)
                        """, (fecha_valor,))
                        max_actual_año = cursor.fetchone()[0] or 0
                        siguiente_numero = max_actual_año + 1
                        
                        # Verificar si este número ya existe
                        cursor.execute(f"""
                            SELECT COUNT(*) 
                            FROM {tabla}
                            WHERE numero_secuencial = %s
                        """, (siguiente_numero,))
                        existe = cursor.fetchone()[0] > 0
                        
                        # Si el número ya existe, usar el máximo global + 1
                        if existe:
                            cursor.execute(f"""
                                SELECT COALESCE(MAX(numero_secuencial), 0)
                                FROM {tabla}
                            """)
                            max_actual_global = cursor.fetchone()[0] or 0
                            siguiente_numero = max_actual_global + 1
                        
                        # Esperar antes de reintentar con jitter aleatorio
                        sleep_time = 0.01 * (intento + 1) + random.uniform(0, 0.02)
                        time.sleep(sleep_time)
                        continue
                    else:
                        # Si es otro tipo de error, hacer rollback y lanzarlo
                        conn.rollback()
                        cursor.close()
                        raise
            
            # Si llegamos aquí, todos los intentos fallaron
            cursor.close()
            raise Exception(f"No se pudo asignar número secuencial después de {max_intentos} intentos debido a duplicados")
            
        except Exception as e:
            cursor.close()
            raise
            
    except Exception as e:
        print(f"Error al asignar número secuencial: {str(e)}")
        raise
def obtener_categorias():
    conn = conexion()
    try:
        with conn.cursor() as cursor:
            cursor.execute("SELECT id_categoria, nombre FROM Categoria")
            return cursor.fetchall()
    except Exception as e:
        print("Error al obtener categorías:", str(e))
        return []
    finally:
        conn.close()

@app.route("/salida_producto")
@requiere_login('admin')
def salida_producto():
    conn = conexion()
    try:
        with conn.cursor() as cursor:
            cursor.execute("""
                SELECT 
                    p.id_producto,
                    p.codigo,
                    p.nombre,
                    p.descripcion,
                    p.cantidad,
                    p.unidad_medida,
                    c.nombre as categoria,
                    p.estado,
                    p.procedencia
                FROM Producto p
                LEFT JOIN Categoria c ON p.id_categoria = c.id_categoria
                WHERE p.cantidad > 0  -- Solo productos con stock disponible
                ORDER BY 
                    CASE 
                        WHEN p.estado = 'disponible' THEN 1
                        ELSE 2
                    END,
                    p.nombre
            """)
            productos = cursor.fetchall()
            return render_template("salida_producto.html", productos=productos)
    finally:
        conn.close()

from datetime import datetime  # Asegúrate de que esta importación esté al inicio del archivo

@app.route("/confirmar_salida", methods=["POST"])
@requiere_login('admin')
def confirmar_salida():
    productos_ids = request.form.getlist('productos[]')
    conn = conexion()
    try:
        with conn.cursor() as cursor:
            cursor.execute("""
                SELECT p.*, c.nombre as categoria_nombre
                FROM Producto p
                LEFT JOIN Categoria c ON p.id_categoria = c.id_categoria
                WHERE p.id_producto IN %s
            """, (tuple(productos_ids),))
            productos_seleccionados = cursor.fetchall()
            
            # Obtener las áreas disponibles
            cursor.execute("SELECT DISTINCT area FROM AreaUsuario ORDER BY area")
            areas = cursor.fetchall()
            
            # Obtener fecha y hora actual
            now = datetime.now()
            
            return render_template("confirmar_salida.html", 
                                 productos_seleccionados=productos_seleccionados,
                                 areas=areas,
                                 now=now)  # Pasar la variable now al template
    except Exception as e:
        print("Error al preparar salida:", str(e))
        return "Error al preparar la salida de productos"
    finally:
        conn.close()

def actualizar_estado_producto(cursor, id_producto):
    """
    Actualiza el estado del producto basado en su stock y préstamos activos
    """
    # Obtener cantidad total y cantidad prestada
    cursor.execute("""
        SELECT 
            p.cantidad as stock_actual,
            COALESCE(SUM(dp.cantidad), 0) as cantidad_prestada
        FROM Producto p
        LEFT JOIN DetallePrestamo dp ON p.id_producto = dp.id_producto
        LEFT JOIN Prestamo pr ON dp.id_prestamo = pr.id_prestamo
        LEFT JOIN Devolucion d ON pr.id_prestamo = d.id_prestamo
        WHERE p.id_producto = %s
        AND d.id_devolucion IS NULL  -- Solo préstamos no devueltos
        GROUP BY p.id_producto, p.cantidad
    """, (id_producto,))
    
    resultado = cursor.fetchone()
    if resultado:
        stock_actual = float(resultado[0])
        cantidad_prestada = float(resultado[1])
        
        # Determinar el estado
        if stock_actual <= 0:
            if cantidad_prestada > 0:
                nuevo_estado = 'prestado'
            else:
                nuevo_estado = 'agotado'
        else:
            nuevo_estado = 'disponible'
        
        # Actualizar el estado
        cursor.execute("""
            UPDATE Producto 
            SET estado = %s 
            WHERE id_producto = %s
        """, (nuevo_estado, id_producto))
def asignar_numero_secuencial_salida(conn, id_salida, es_combustible=False):
    """
    Asigna un número secuencial a una salida, diferenciando entre productos y combustibles.
    """
    try:
        with conn.cursor() as cursor:
            año_actual = datetime.now().year
            tipo_salida = 'combustible' if es_combustible else 'producto'
            
            # Obtener o crear secuencia para el año y tipo
            cursor.execute("""
                INSERT INTO secuencias_salida (tipo_salida, año, ultimo_numero)
                VALUES (%s, %s, 0)
                ON DUPLICATE KEY UPDATE id_secuencia=LAST_INSERT_ID(id_secuencia)
            """, (tipo_salida, año_actual))
            
            # Incrementar y obtener siguiente número
            cursor.execute("""
                UPDATE secuencias_salida 
                SET ultimo_numero = ultimo_numero + 1
                WHERE tipo_salida = %s AND año = %s
                RETURNING ultimo_numero
            """, (tipo_salida, año_actual))
            
            nuevo_numero = cursor.fetchone()[0]
            
            # Actualizar la salida con el nuevo número
            prefijo = 'COMB' if es_combustible else 'PROD'
            numero_formato = f"{prefijo}-{año_actual}-{nuevo_numero:04d}"
            
            cursor.execute("""
                UPDATE salidaproducto 
                SET numero_salida = %s
                WHERE id_salida = %s
            """, (numero_formato, id_salida))
            
            conn.commit()
            return numero_formato
            
    except Exception as e:
        conn.rollback()
        print(f"Error al asignar número secuencial: {str(e)}")
        raise

def registrar_salida(conn, datos_salida):
    try:
        with conn.cursor() as cursor:
            # Primero verificamos si el producto es combustible
            cursor.execute("""
                SELECT c.nombre 
                FROM producto p 
                JOIN categoria c ON p.id_categoria = c.id_categoria
                WHERE p.id_producto = %s
            """, (datos_salida['producto'],))
            
            categoria = cursor.fetchone()[0].lower()
            es_combustible = 'gasolina' in categoria or 'petróleo' in categoria
            
            if es_combustible:
                # Si es combustible, registramos en la tabla de vales sin afectar la numeración de salidas
                cursor.execute("""
                    INSERT INTO salidaproducto 
                    (id_prestatario, id_producto, cantidad, fecha_salida, motivo, autorizado_por)
                    VALUES (%s, %s, %s, %s, %s, %s)
                """, (datos_salida['prestatario'], datos_salida['producto'], 
                     datos_salida['cantidad'], datos_salida['fecha'],
                     datos_salida['motivo'], datos_salida['autorizado_por']))
                
                id_salida = cursor.lastrowid
                # No asignamos número de salida para combustibles
                numero_salida = None
            else:
                # Si es producto normal, procedemos con la numeración regular
                cursor.execute("""
                    INSERT INTO salidaproducto 
                    (id_prestatario, id_producto, cantidad, fecha_salida, motivo, autorizado_por)
                    VALUES (%s, %s, %s, %s, %s, %s)
                """, (datos_salida['prestatario'], datos_salida['producto'], 
                     datos_salida['cantidad'], datos_salida['fecha'],
                     datos_salida['motivo'], datos_salida['autorizado_por']))
                
                id_salida = cursor.lastrowid
                
                # Obtener el último número de salida (solo para productos no combustibles)
                cursor.execute("""
                    SELECT COALESCE(MAX(CAST(numero_salida AS UNSIGNED)), 0) + 1
                    FROM salidaproducto sp
                    JOIN producto p ON sp.id_producto = p.id_producto
                    JOIN categoria c ON p.id_categoria = c.id_categoria
                    WHERE YEAR(sp.fecha_salida) = YEAR(CURRENT_DATE)
                    AND c.nombre NOT IN ('Gasolina', 'Petróleo')
                """)
                
                numero_salida = cursor.fetchone()[0]
                
                # Actualizar el número de salida
                cursor.execute("""
                    UPDATE salidaproducto 
                    SET numero_salida = %s
                    WHERE id_salida = %s
                """, (numero_salida, id_salida))
            
            conn.commit()
            return id_salida, numero_salida
            
    except Exception as e:
        conn.rollback()
        print(f"Error al registrar salida: {str(e)}")
        raise
@app.route("/procesar_salida", methods=["POST"])
@requiere_login('admin')
def procesar_salida():
    conn = conexion()
    try:
        # Datos del prestatario
        nombre = request.form['nombre']
        apellido = request.form['apellido']
        dni = request.form['dni']
        telefono = request.form['telefono']
        
        # Datos de la salida
        autorizado_por = request.form['autorizado_por']
        motivo = request.form['motivo']
        observacion_autorizacion = request.form['observacion_autorizacion']
        area_usuario = request.form['area']
        fecha_salida = datetime.strptime(request.form['fecha_salida'], '%Y-%m-%dT%H:%M')
        
        # Arrays de productos
        productos_ids = request.form.getlist('producto_id[]')
        cantidades = request.form.getlist('cantidad[]')
        observaciones = request.form.getlist('observacion[]')
        
        # Inicializar la lista de productos_salida
        productos_salida = []
        
        # Formatear el motivo con la observación
        motivo_completo = f"{motivo} - Observación: {observacion_autorizacion}" if observacion_autorizacion else motivo
        
        with conn.cursor() as cursor:
            # Registrar o actualizar prestatario
            cursor.execute("""
                INSERT INTO Prestatario (nombre, apellido, telefono, dni)
                VALUES (%s, %s, %s, %s)
                ON DUPLICATE KEY UPDATE id_prestatario=LAST_INSERT_ID(id_prestatario)
            """, (nombre, apellido, telefono, dni))
            id_prestatario = cursor.lastrowid
            
            cantidad_total = sum(float(c) for c in cantidades)
            
            # Crear la salida principal con el motivo formateado
            cursor.execute("""
                INSERT INTO SalidaProducto 
                (id_prestatario, fecha_salida, motivo, autorizado_por, cantidad)
                VALUES (%s, %s, %s, %s, %s)
            """, (id_prestatario, fecha_salida, motivo_completo, autorizado_por, cantidad_total))
            
            id_salida = cursor.lastrowid
            
            # Asignar número secuencial
            numero_salida = asignar_numero_secuencial(
                conn,
                'NumeroSalida',
                'id_salida',
                id_salida,
                'fecha_salida',
                fecha_salida
            )
            
            for i in range(len(productos_ids)):
                # Verificar stock disponible
                cursor.execute("""
                    SELECT codigo, nombre, unidad_medida, cantidad 
                    FROM Producto 
                    WHERE id_producto = %s
                """, (productos_ids[i],))
                producto_info = cursor.fetchone()
                
                if not producto_info:
                    raise ValueError(f"Producto con ID {productos_ids[i]} no encontrado")
                
                cantidad_salida = float(cantidades[i])
                if float(producto_info[3]) < cantidad_salida:
                    raise ValueError(f"Stock insuficiente para {producto_info[1]}")
                
                # Registrar detalle de salida
                cursor.execute("""
                    INSERT INTO DetalleSalida 
                    (id_salida, id_producto, cantidad_salida, 
                     cantidad_disponible_retorno, area_usuario, fecha_salida, 
                     observacion)
                    VALUES (%s, %s, %s, %s, %s, %s, %s)
                """, (
                    id_salida,
                    productos_ids[i],
                    cantidad_salida,
                    cantidad_salida,
                    area_usuario,
                    fecha_salida,
                    observaciones[i]
                ))
                
                # Actualizar stock
                cursor.execute("""
                    UPDATE Producto 
                    SET cantidad = cantidad - %s 
                    WHERE id_producto = %s
                """, (cantidad_salida, productos_ids[i]))
                
                # Registrar en historial
                cursor.execute("""
                    INSERT INTO HistorialModificaciones 
                    (id_producto, tipo_modificacion, cantidad, fecha_modificacion, detalle)
                    VALUES (%s, 'salida', %s, %s, %s)
                """, (
                    productos_ids[i],
                    cantidad_salida,
                    fecha_salida,
                    f"Salida #{numero_salida} - {nombre} {apellido} - DNI: {dni} - Área: {area_usuario} - {motivo}"
                ))
                
                productos_salida.append({
                    'codigo': producto_info[0],
                    'nombre': producto_info[1],
                    'cantidad': cantidad_salida,
                    'unidad': producto_info[2],
                    'area_usuario': area_usuario,
                    'observacion': observaciones[i],
                    'motivo': motivo
                })
            
            conn.commit()
            
            # Generar documento Word
            doc = generar_documento_salida(
                nombre=nombre,
                apellido=apellido,
                dni=dni,
                productos=productos_salida,
                fecha_salida=fecha_salida,
                id_salida=id_salida,
                autorizado_por=autorizado_por,
                observacion_autorizacion=observacion_autorizacion,
                telefono=telefono,
                numero_salida=numero_salida
            )
            
            return send_file(
                doc,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                as_attachment=True,
                download_name=f'salida_{str(numero_salida).zfill(4)}.docx'
            )
            
    except Exception as e:
        conn.rollback()
        print(f"Error al procesar salida: {str(e)}")
        return jsonify({"success": False, "message": str(e)})
    finally:
        conn.close()

def generar_documento_salida(nombre, apellido, dni, productos, fecha_salida, id_salida, autorizado_por, observacion_autorizacion, telefono, numero_salida):
    doc = Document()
    
    # Configuración de estilos
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    
    # Número de salida en esquina superior derecha
    num_salida = doc.add_paragraph()
    num_salida.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    num_run = num_salida.add_run(f'N° {str(numero_salida).zfill(4)}')
    num_run.font.size = Pt(11)
    num_run.bold = True
    
    # Título
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run('\nCONSTANCIA DE SALIDA DE MATERIALES\n')
    title_run.bold = True
    title_run.font.size = Pt(14)
    
    # Fecha
    fecha = doc.add_paragraph()
    fecha_run = fecha.add_run(f'Piás, {fecha_en_español(fecha_salida)}')
    fecha.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Datos del solicitante
    doc.add_paragraph()
    datos_solicitante = doc.add_paragraph()
    datos_solicitante.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    datos_solicitante.add_run('Por medio del presente documento se deja constancia que el Sr(a). ')
    datos_solicitante.add_run(f'{nombre} {apellido}').bold = True
    datos_solicitante.add_run(', identificado con DNI N° ')
    datos_solicitante.add_run(dni).bold = True
    datos_solicitante.add_run(', con número de teléfono ')
    datos_solicitante.add_run(telefono).bold = True
    datos_solicitante.add_run(', ha retirado los siguientes materiales del Almacén Central:')
    
    doc.add_paragraph()
    
    # Tabla de materiales
    tabla = doc.add_table(rows=1, cols=6)  # Cambiado de 5 a 6 columnas
    tabla.style = 'Table Grid'
    tabla.autofit = True
    
    # Encabezados de tabla
    encabezados = tabla.rows[0].cells
    for i, texto in enumerate(['CÓDIGO', 'DESCRIPCIÓN', 'CANTIDAD', 'U.M.', 'ÁREA USUARIO', 'OBSERVACIÓN']):  # Agregado 'OBSERVACIÓN'
        encabezados[i].text = texto
        encabezados[i].paragraphs[0].runs[0].bold = True
        encabezados[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Productos
    for producto in productos:
        row = tabla.add_row().cells
        row[0].text = producto['codigo']
        row[1].text = producto['nombre']
        row[2].text = str(producto['cantidad'])
        row[3].text = producto['unidad']
        row[4].text = producto['area_usuario']
        row[5].text = producto['observacion']  # Agregada la observación del producto
        for cell in row:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Motivo de la salida
    doc.add_paragraph()
    motivo_parrafo = doc.add_paragraph()
    motivo_parrafo.add_run('MOTIVO DE LA SALIDA: ').bold = True
    motivo_parrafo.add_run(productos[0]['motivo'])
    
    # Autorización
    doc.add_paragraph()
    auth_parrafo = doc.add_paragraph()
    auth_parrafo.add_run('AUTORIZADO POR: ').bold = True
    auth_parrafo.add_run(autorizado_por)
    
    # Observación de la autorización
    if observacion_autorizacion:
        doc.add_paragraph()
        obs_auth_parrafo = doc.add_paragraph()
        obs_auth_parrafo.add_run('OBSERVACIÓN DE LA AUTORIZACIÓN: ').bold = True
        obs_auth_parrafo.add_run(observacion_autorizacion)
    
    # Espacio para firmas
    doc.add_paragraph('\n\n')
    
    # Tabla de firmas (3 columnas)
    firma_table = doc.add_table(rows=1, cols=3)
    firma_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Firma del solicitante
    firma_solicitante = firma_table.rows[0].cells[0]
    p_solicitante = firma_solicitante.add_paragraph('_____________________')
    p_solicitante.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_nombre_solicitante = firma_solicitante.add_paragraph(
        f'{nombre} {apellido}\n'
        f'DNI: {dni}\n'
        'SOLICITANTE'
    )
    p_nombre_solicitante.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Firma del autorizante
    firma_autorizante = firma_table.rows[0].cells[1]
    p_autorizante = firma_autorizante.add_paragraph('_____________________')
    p_autorizante.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_nombre_autorizante = firma_autorizante.add_paragraph(
        f'{autorizado_por}\n'
        'AUTORIZANTE'
    )
    p_nombre_autorizante.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Firma del encargado de almacén
    firma_encargado = firma_table.rows[0].cells[2]
    p_encargado = firma_encargado.add_paragraph('_____________________')
    p_encargado.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_nombre_encargado = firma_encargado.add_paragraph(
        'FABIAN LOZANO CRUZ\n'
        'ASISTENTE DE ALMACÉN'
    )
    p_nombre_encargado.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Pie de página
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run('ÁREA DE LOGÍSTICA - ALMACÉN CENTRAL')
    footer_run.font.size = Pt(8)
    footer_run.bold = True
    
    # Guardar documento
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output

@app.route("/prestamo_nuevo")
@requiere_login('admin')
def prestamo_nuevo():
    conn = conexion()
    try:
        with conn.cursor() as cursor:
            cursor.execute("""
                SELECT 
                    p.id_producto,
                    p.codigo,
                    p.nombre,
                    p.descripcion,
                    p.cantidad,
                    p.unidad_medida,
                    p.procedencia,
                    p.estado,
                    CASE
                        WHEN p.cantidad <= 0 AND (
                            SELECT COALESCE(SUM(dp.cantidad), 0)
                            FROM DetallePrestamo dp
                            JOIN Prestamo pr ON dp.id_prestamo = pr.id_prestamo
                            LEFT JOIN Devolucion d ON pr.id_prestamo = d.id_prestamo
                            WHERE dp.id_producto = p.id_producto
                            AND d.id_devolucion IS NULL
                        ) > 0 THEN 'prestado'
                        WHEN p.cantidad <= 0 THEN 'agotado'
                        ELSE 'disponible'
                    END as estado_actual,
                    c.nombre as categoria_nombre
                FROM Producto p
                LEFT JOIN Categoria c ON p.id_categoria = c.id_categoria
                GROUP BY 
                    p.id_producto, p.codigo, p.nombre, p.descripcion, p.cantidad,
                    p.unidad_medida, p.procedencia, p.estado, c.nombre
                ORDER BY p.nombre
            """)
            productos = cursor.fetchall()
            return render_template("prestamo_nuevo.html", productos=productos)
    except Exception as e:
        print("Error al cargar productos:", str(e))
        return "Error al cargar los productos: " + str(e)
    finally:
        conn.close()

@app.route("/confirmar_prestamo", methods=["POST"])
@requiere_login('admin')
def confirmar_prestamo():
    productos_ids = request.form.getlist('productos[]')
    conn = conexion()
    try:
        with conn.cursor() as cursor:
            cursor.execute("""
                SELECT p.*, c.nombre as categoria_nombre
                FROM Producto p
                LEFT JOIN Categoria c ON p.id_categoria = c.id_categoria
                WHERE p.id_producto IN %s
            """, (tuple(productos_ids),))
            productos_seleccionados = cursor.fetchall()
            
            # Generar documento Word en lugar de Excel
            productos_prestamo = [{
                'codigo': prod[1],  # código del producto
                'nombre': prod[2],  # nombre del producto
                'cantidad': 0,      # la cantidad se llenará en el formulario
                'unidad': prod[5]   # unidad de medida
            } for prod in productos_seleccionados]
            
            return render_template(
                "confirmar_prestamo.html", 
                productos_seleccionados=productos_seleccionados,
                now=datetime.now()
            )
            
    except Exception as e:
        print("Error al preparar préstamo:", str(e))
        flash('Error al preparar el préstamo', 'danger')
        return redirect(url_for('prestamo_nuevo'))
    finally:
        conn.close()

@app.route("/procesar_prestamo", methods=["POST"])
@requiere_login('admin')
def procesar_prestamo():
    conn = conexion()
    try:
        with conn.cursor() as cursor:
            # Datos del prestatario
            dni = request.form["dni"]
            nombre = request.form["nombre"]
            apellido = request.form["apellido"]
            telefono = request.form["telefono"]
            fecha_prestamo = datetime.strptime(request.form["fecha_prestamo"], '%Y-%m-%dT%H:%M')
            fecha_devolucion = datetime.strptime(request.form["fecha_devolucion"], '%Y-%m-%dT%H:%M')
            motivo_prestamo = request.form.get('motivo_prestamo')
            
            # Nuevos campos de autorización
            autorizado_por = request.form.get('autorizado_por')
            observacion_autorizacion = request.form.get('observacion_autorizacion')
            
            # Registrar o actualizar prestatario
            cursor.execute("""
                INSERT INTO Prestatario (nombre, apellido, telefono, dni)
                VALUES (%s, %s, %s, %s)
                ON DUPLICATE KEY UPDATE id_prestatario=LAST_INSERT_ID(id_prestatario)
            """, (nombre, apellido, telefono, dni))
            
            id_prestatario = cursor.lastrowid
            
            # Crear el préstamo
            cursor.execute("""
                INSERT INTO Prestamo 
                (id_prestatario, nombre_solicitante, apellido_solicitante, 
                dni_solicitante, telefono_solicitante, fecha_prestamo, 
                fecha_devolucion_planeada, motivo_prestamo, autorizado_por, 
                observacion_autorizacion)
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
            """, (id_prestatario, nombre, apellido, dni, telefono, 
                  fecha_prestamo, fecha_devolucion, motivo_prestamo,
                  autorizado_por, observacion_autorizacion))
            
            id_prestamo = cursor.lastrowid
            
            # Obtener el siguiente número secuencial del año actual
            cursor.execute("""
                SELECT COALESCE(MAX(numero_secuencial), 0) + 1
                FROM NumeroPrestamo 
                WHERE YEAR(fecha_prestamo) = YEAR(NOW())
            """)
            siguiente_numero = cursor.fetchone()[0]
            
            # Registrar el número secuencial
            cursor.execute("""
                INSERT INTO NumeroPrestamo 
                (id_prestamo, numero_secuencial, fecha_prestamo)
                VALUES (%s, %s, %s)
            """, (id_prestamo, siguiente_numero, fecha_prestamo))
            
            # Procesar productos
            productos_ids = request.form.getlist("producto_id[]")
            cantidades = request.form.getlist("cantidad[]")
            observaciones = request.form.getlist("observacion[]")
            productos_prestamo = []
            
            for i in range(len(productos_ids)):
                id_producto = productos_ids[i]
                cantidad = float(cantidades[i])
                observacion = observaciones[i]
                
                cursor.execute("""
                    SELECT codigo, nombre, unidad_medida 
                    FROM Producto 
                    WHERE id_producto = %s
                """, (id_producto,))
                producto_info = cursor.fetchone()
                
                productos_prestamo.append({
                    'codigo': producto_info[0],
                    'nombre': producto_info[1],
                    'cantidad': cantidad,
                    'unidad': producto_info[2],
                    'observacion': observacion
                })
                
                # Registrar en DetallePrestamo
                cursor.execute("""
                    INSERT INTO DetallePrestamo 
                    (id_prestamo, id_producto, cantidad, observacion)
                    VALUES (%s, %s, %s, %s)
                """, (id_prestamo, id_producto, cantidad, observacion))
                
                # Actualizar stock del producto
                cursor.execute("""
                    UPDATE Producto 
                    SET cantidad = cantidad - %s 
                    WHERE id_producto = %s
                """, (cantidad, id_producto))
                
                # Registrar en historial
                detalle_historial = (
                    f"Préstamo #{siguiente_numero} - "
                    f"Solicitante: {nombre} {apellido} - "
                    f"DNI: {dni} - "
                    f"Cantidad: {cantidad} {producto_info[2]}"
                )
                
                cursor.execute("""
                    INSERT INTO HistorialModificaciones 
                    (id_producto, tipo_modificacion, cantidad, 
                     fecha_modificacion, detalle)
                    VALUES (%s, 'prestamo', %s, %s, %s)
                """, (id_producto, cantidad, datetime.now(), detalle_historial))
            
            conn.commit()
            
            # Generar documento
            doc = generar_documento_prestamo(
                siguiente_numero,
                dni,
                nombre,
                apellido,
                telefono,
                productos_prestamo,
                es_devolucion=False,
                fecha_prestamo=fecha_prestamo,
                fecha_devolucion=fecha_devolucion,
                motivo_prestamo=motivo_prestamo,
                autorizado_por=autorizado_por,
                observacion_autorizacion=observacion_autorizacion
            )
            
            # Actualizar el nombre del archivo con el número secuencial
            return send_file(
                doc,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                as_attachment=True,
                download_name=f'prestamo_{str(siguiente_numero).zfill(4)}.docx'
            )
            
    except Exception as e:
        conn.rollback()
        print(f"Error al procesar préstamo: {str(e)}")
        flash('Error al procesar el préstamo', 'danger')
        return redirect(url_for('prestamo_nuevo'))
    finally:
        conn.close()

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from datetime import datetime

def generar_documento_prestamo(id_prestamo, dni, nombre, apellido, telefono, productos, 
                             es_devolucion=False, fecha_prestamo=None, fecha_devolucion=None, 
                             motivo_prestamo=None, autorizado_por=None, observacion_autorizacion=None):
    doc = Document()
    
    # Configurar márgenes
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Número de préstamo
    num_prestamo = doc.add_paragraph()
    num_prestamo_run = num_prestamo.add_run(f'N° de Préstamo: {str(id_prestamo).zfill(4)}')
    num_prestamo.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    num_prestamo_run.font.size = Pt(12)
    num_prestamo_run.font.bold = True
    
    # Título
    titulo = doc.add_paragraph()
    titulo_run = titulo.add_run('PRÉSTAMO DE MATERIALES DE ALMACÉN')
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titulo_run.font.size = Pt(14)
    titulo_run.font.bold = True
    
    # Fecha
    fecha = doc.add_paragraph()
    fecha.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if es_devolucion:
        fecha_texto = fecha_devolucion
        tipo_fecha = "Devolución"
    else:
        fecha_texto = fecha_prestamo
        tipo_fecha = "Préstamo"
    fecha.add_run(f'Fecha de {tipo_fecha}: {fecha_texto.strftime("%d/%m/%Y")}')
    
    # Contenido principal
    contenido = doc.add_paragraph()
    contenido.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Justificar el texto principal
    contenido.add_run('Por medio del presente documento se deja constancia que el Sr(a). ')
    contenido.add_run(f'{nombre} {apellido}').bold = True
    contenido.add_run(', identificado con DNI N° ')
    contenido.add_run(dni).bold = True
    contenido.add_run(', con número de teléfono ')
    contenido.add_run(telefono).bold = True
    contenido.add_run(', solicita en calidad de préstamo los siguientes materiales:')
    
    # Tabla de productos
    tabla = doc.add_table(rows=1, cols=5)  # Cambiado de 4 a 5 columnas
    tabla.style = 'Table Grid'
    encabezados = tabla.rows[0].cells
    for i, texto in enumerate(['CÓDIGO', 'DESCRIPCIÓN', 'CANTIDAD', 'U.M.', 'OBSERVACIÓN']):  # Agregado 'OBSERVACIÓN'
        encabezados[i].text = texto
        encabezados[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        encabezados[i].paragraphs[0].runs[0].bold = True
    
    for producto in productos:
        row = tabla.add_row().cells
        row[0].text = producto['codigo']
        row[1].text = producto['nombre']
        row[2].text = str(producto['cantidad'])
        row[3].text = producto['unidad']
        row[4].text = producto['observacion'] if producto.get('observacion') else ''  # Nueva columna
        for cell in row:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Agregar compromiso de devolución después de la tabla de productos
    doc.add_paragraph()
    compromiso = doc.add_paragraph()
    compromiso.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    compromiso.add_run('COMPROMISO DE DEVOLUCIÓN: ').bold = True
    compromiso.add_run(
        f'El solicitante se compromete a devolver los materiales prestados en buen estado y sin alteraciones '
        f'antes del día {fecha_devolucion.strftime("%d/%m/%Y")} a las {fecha_devolucion.strftime("%H:%M")} horas. '
        f'En caso de incumplimiento, pérdida, daño o deterioro, el solicitante asumirá la responsabilidad '
        f'de la reposición o reparación correspondiente, además de las sanciones administrativas que correspondan.'
    )

    # Motivo del préstamo
    if motivo_prestamo:
        doc.add_paragraph()
        motivo = doc.add_paragraph()
        motivo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Justificar el motivo
        motivo.add_run('MOTIVO DEL PRÉSTAMO: ').bold = True
        motivo.add_run(motivo_prestamo)
    
    # Autorización
    if autorizado_por:
        doc.add_paragraph()
        auth = doc.add_paragraph()
        auth.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Justificar la autorización
        auth.add_run('AUTORIZADO POR: ').bold = True
        auth.add_run(autorizado_por)
    
    # Observación de autorización
    if observacion_autorizacion:
        doc.add_paragraph()
        obs = doc.add_paragraph()
        obs.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Justificar la observación
        obs.add_run('OBSERVACIÓN DE LA AUTORIZACIÓN: ').bold = True
        obs.add_run(observacion_autorizacion)
    
    # Espacio para firmas
    doc.add_paragraph('\n\n')
    
    # Tabla de firmas
    firma_table = doc.add_table(rows=1, cols=3)
    firma_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Firma del solicitante
    firma_solicitante = firma_table.rows[0].cells[0]
    p_solicitante = firma_solicitante.add_paragraph('_____________________')
    p_solicitante.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_nombre_solicitante = firma_solicitante.add_paragraph(
        f'{nombre} {apellido}\n'
        f'DNI: {dni}\n'
        'SOLICITANTE'
    )
    p_nombre_solicitante.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Firma del autorizante
    firma_autorizante = firma_table.rows[0].cells[1]
    p_autorizante = firma_autorizante.add_paragraph('_____________________')
    p_autorizante.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_nombre_autorizante = firma_autorizante.add_paragraph(
        f'{autorizado_por}\n'
        'AUTORIZANTE'
    )
    p_nombre_autorizante.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Firma del encargado de almacén
    firma_encargado = firma_table.rows[0].cells[2]
    p_encargado = firma_encargado.add_paragraph('_____________________')
    p_encargado.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_nombre_encargado = firma_encargado.add_paragraph(
        'FABIAN LOZANO CRUZ\n'
        'ASISTENTE DE ALMACÉN'
    )
    p_nombre_encargado.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Pie de página
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run('ÁREA DE LOGÍSTICA - ALMACÉN MUNICIPAL')
    footer_run.font.size = Pt(8)
    footer_run.bold = True
    
    # Guardar en memoria
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output

@app.route('/descargar_documento/<filename>')
@requiere_login('admin')
def descargar_documento(filename):
    try:
        return send_from_directory(
            os.path.join(app.root_path, 'static', 'documentos'),
            filename,
            as_attachment=True
        )
    except Exception as e:
        flash('Error al descargar el documento', 'danger')
        return redirect(url_for('prestamos'))

@app.route("/devolucion")
@requiere_login('admin')
def devolucion():
    conn = conexion()
    try:
        with conn.cursor() as cursor:
            cursor.execute("""
                SELECT 
                    d.id_devolucion,
                    COALESCE(nd.numero_secuencial, 0) as numero_devolucion,
                    p.id_prestamo,
                    COALESCE(np.numero_secuencial, 0) as numero_prestamo,
                    pr.codigo,
                    pr.nombre,
                    dp.cantidad,
                    pr.unidad_medida,
                    pre.nombre,
                    pre.apellido,
                    DATE_FORMAT(p.fecha_prestamo, '%d/%m/%Y') as fecha_prestamo,
                    DATE_FORMAT(d.fecha_devolucion, '%d/%m/%Y') as fecha_devolucion
                FROM Devolucion d
                LEFT JOIN NumeroDevolucion nd ON d.id_devolucion = nd.id_devolucion
                JOIN DetallePrestamo dp ON d.id_detalle_prestamo = dp.id_detalle_prestamo
                JOIN Prestamo p ON dp.id_prestamo = p.id_prestamo
                LEFT JOIN NumeroPrestamo np ON p.id_prestamo = np.id_prestamo
                JOIN Producto pr ON dp.id_producto = pr.id_producto
                JOIN Prestatario pre ON p.id_prestatario = pre.id_prestatario
                ORDER BY d.fecha_devolucion DESC
            """)
            devoluciones = cursor.fetchall()
            return render_template('devolucion.html', devoluciones=devoluciones)
    except Exception as e:
        print(f"Error en la ruta /devolucion: {str(e)}")
        flash("Error al cargar las devoluciones", "danger")
        return redirect(url_for('index'))
    finally:
        conn.close()
def reporte_inventario():
    return render_template("reportes/inventario.html")

@app.route("/reporte_inventario_fecha")
@requiere_login('admin')
def reporte_inventario_fecha():
    return render_template("reportes/inventario_fecha.html")

@app.route("/reporte_prestamos")
@requiere_login('admin')
def reporte_prestamos():
    return render_template("reportes/prestamos.html")

@app.route("/reporte_devoluciones")
@requiere_login('admin')
def reporte_devoluciones():
    return render_template("reportes/devoluciones.html")

@app.route("/reportes/solicitudes_prestamo")
@requiere_login('admin')
def solicitudes_prestamo():
    conn = conexion()
    try:
        with conn.cursor() as cursor:
            cursor.execute("""
                WITH PrestamosNumerados AS (
                    SELECT DISTINCT
                        p.id_prestamo,
                        np.numero_secuencial as numero_prestamo,
                        p.fecha_prestamo,
                        p.nombre_solicitante,
                        p.apellido_solicitante,
                        p.dni_solicitante,
                        p.telefono_solicitante,
                        GROUP_CONCAT(DISTINCT CONCAT(pr.nombre, ' (', dp.cantidad, ' ', 
                            pr.unidad_medida, ')') SEPARATOR ', ') as productos,
                        CASE 
                            WHEN d.fecha_devolucion IS NOT NULL THEN 'Devuelto'
                            WHEN p.fecha_devolucion_planeada < NOW() THEN 'Vencido'
                            ELSE 'Activo'
                        END as estado,
                        COALESCE(d.fecha_devolucion, p.fecha_devolucion_planeada) as fecha_devolucion,
                        p.autorizado_por,
                        p.observacion_autorizacion
                    FROM Prestamo p
                    LEFT JOIN NumeroPrestamo np ON p.id_prestamo = np.id_prestamo
                    JOIN DetallePrestamo dp ON p.id_prestamo = dp.id_prestamo
                    JOIN Producto pr ON dp.id_producto = pr.id_producto
                    LEFT JOIN Devolucion d ON p.id_prestamo = d.id_prestamo
                    GROUP BY 
                        p.id_prestamo, 
                        np.numero_secuencial,
                        p.fecha_prestamo,
                        p.nombre_solicitante,
                        p.apellido_solicitante,
                        p.dni_solicitante,
                        p.telefono_solicitante,
                        d.fecha_devolucion,
                        p.fecha_devolucion_planeada,
                        p.autorizado_por,
                        p.observacion_autorizacion
                )
                SELECT * FROM PrestamosNumerados
                ORDER BY YEAR(fecha_prestamo) DESC, 
                         numero_prestamo DESC
            """)
            solicitudes = cursor.fetchall()
            
            # Asignar números secuenciales a préstamos que no los tengan
            for solicitud in solicitudes:
                if solicitud[1] is None or solicitud[1] == 0:  # Si no tiene número secuencial
                    numero_prestamo = asignar_numero_secuencial(
                        conn,
                        'NumeroPrestamo',
                        'id_prestamo',
                        solicitud[0],  # id_prestamo
                        'fecha_prestamo',
                        solicitud[2]   # fecha_prestamo
                    )
            
            # Volver a obtener los datos actualizados
            cursor.execute("""
                WITH PrestamosNumerados AS (
                    SELECT DISTINCT
                        p.id_prestamo,
                        np.numero_secuencial as numero_prestamo,
                        p.fecha_prestamo,
                        p.nombre_solicitante,
                        p.apellido_solicitante,
                        p.dni_solicitante,
                        p.telefono_solicitante,
                        GROUP_CONCAT(DISTINCT CONCAT(pr.nombre, ' (', dp.cantidad, ' ', 
                            pr.unidad_medida, ')') SEPARATOR ', ') as productos,
                        CASE 
                            WHEN d.fecha_devolucion IS NOT NULL THEN 'Devuelto'
                            WHEN p.fecha_devolucion_planeada < NOW() THEN 'Vencido'
                            ELSE 'Activo'
                        END as estado,
                        COALESCE(d.fecha_devolucion, p.fecha_devolucion_planeada) as fecha_devolucion,
                        p.autorizado_por,
                        p.observacion_autorizacion
                    FROM Prestamo p
                    JOIN NumeroPrestamo np ON p.id_prestamo = np.id_prestamo
                    JOIN DetallePrestamo dp ON p.id_prestamo = dp.id_prestamo
                    JOIN Producto pr ON dp.id_producto = pr.id_producto
                    LEFT JOIN Devolucion d ON p.id_prestamo = d.id_prestamo
                    GROUP BY 
                        p.id_prestamo, 
                        np.numero_secuencial,
                        p.fecha_prestamo,
                        p.nombre_solicitante,
                        p.apellido_solicitante,
                        p.dni_solicitante,
                        p.telefono_solicitante,
                        d.fecha_devolucion,
                        p.fecha_devolucion_planeada,
                        p.autorizado_por,
                        p.observacion_autorizacion
                )
                SELECT * FROM PrestamosNumerados
                ORDER BY YEAR(fecha_prestamo) DESC, 
                         numero_prestamo DESC
            """)
            solicitudes = cursor.fetchall()
            
            # Convertir las fechas a objetos datetime si son strings
            solicitudes = [list(s) for s in solicitudes]
            for s in solicitudes:
                if isinstance(s[2], str):
                    s[2] = datetime.strptime(s[2], '%Y-%m-%d %H:%M:%S')
                if isinstance(s[9], str):
                    s[9] = datetime.strptime(s[9], '%Y-%m-%d %H:%M:%S')
            
            return render_template("solicitudes_prestamo.html", solicitudes=solicitudes)
    finally:
        conn.close()
@app.route("/descargar_solicitud_prestamo/<int:id_prestamo>")
@requiere_login('admin')
def descargar_solicitud_prestamo(id_prestamo):
    conn = conexion()
    try:
        with conn.cursor() as cursor:
            # Obtener el número secuencial del préstamo
            cursor.execute("""
                SELECT numero_secuencial
                FROM NumeroPrestamo
                WHERE id_prestamo = %s
            """, (id_prestamo,))
            
            result = cursor.fetchone()
            if not result:
                flash('Número de préstamo no encontrado', 'danger')
                return redirect(url_for('solicitudes_prestamo'))
                
            numero_prestamo = result[0]
            
            # Obtener información del préstamo y prestatario con motivo y autorización
            cursor.execute("""
                SELECT 
                    p.nombre_solicitante,
                    p.apellido_solicitante,
                    p.dni_solicitante,
                    p.telefono_solicitante,
                    p.fecha_prestamo,
                    p.fecha_devolucion_planeada,
                    p.motivo_prestamo,
                    p.autorizado_por,
                    p.observacion_autorizacion
                FROM Prestamo p
                WHERE p.id_prestamo = %s
            """, (id_prestamo,))
            
            prestamo = cursor.fetchone()
            if not prestamo:
                flash('Préstamo no encontrado', 'danger')
                return redirect(url_for('solicitudes_prestamo'))
            
            # Obtener productos del préstamo con sus observaciones
            cursor.execute("""
                SELECT 
                    pr.codigo,
                    pr.nombre,
                    dp.cantidad,
                    pr.unidad_medida,
                    dp.observacion
                FROM DetallePrestamo dp
                JOIN Producto pr ON dp.id_producto = pr.id_producto
                WHERE dp.id_prestamo = %s
            """, (id_prestamo,))
            
            productos = [{
                'codigo': p[0],
                'nombre': p[1],
                'cantidad': float(p[2]),
                'unidad': p[3],
                'observacion': p[4] if p[4] else ''
            } for p in cursor.fetchall()]
            
            # Generar documento con los nuevos campos
            doc = generar_documento_prestamo(
                numero_prestamo,  # Usar el número secuencial de NumeroPrestamo
                prestamo[2],     # dni
                prestamo[0],     # nombre
                prestamo[1],     # apellido
                prestamo[3],     # telefono
                productos,
                es_devolucion=False,
                fecha_prestamo=prestamo[4],
                fecha_devolucion=prestamo[5],
                motivo_prestamo=prestamo[6],
                autorizado_por=prestamo[7],
                observacion_autorizacion=prestamo[8]
            )
            
            return send_file(
                doc,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                as_attachment=True,
                download_name=f'prestamo_{str(numero_prestamo).zfill(4)}.docx'  # Cambiado a 4 dígitos
            )
            
    except Exception as e:
        print("Error al descargar solicitud:", str(e))
        flash('Error al descargar la solicitud', 'danger')
        return redirect(url_for('solicitudes_prestamo'))
    finally:
        conn.close()

@app.route("/logout")
def logout():
    session.clear()
    return redirect(url_for('login'))

@app.route("/historial")
@requiere_login('admin')
def historial():
    conn = conexion()
    try:
        with conn.cursor() as cursor:
            cursor.execute("""
                WITH UltimoMovimiento AS (
                    SELECT 
                        id_producto,
                        MAX(fecha_modificacion) as fecha_modificacion
                    FROM (
                        -- Movimientos de HistorialModificaciones
                        SELECT id_producto, fecha_modificacion
                        FROM HistorialModificaciones
                        UNION ALL
                        -- Movimientos de SalidaProducto
                        SELECT id_producto, fecha_salida as fecha_modificacion
                        FROM SalidaProducto
                        UNION ALL
                        -- Movimientos de DetallePrestamo
                        SELECT dp.id_producto, p.fecha_prestamo as fecha_modificacion
                        FROM DetallePrestamo dp
                        JOIN Prestamo p ON dp.id_prestamo = p.id_prestamo
                    ) TodosLosMovimientos
                    GROUP BY id_producto
                )
                SELECT 
                    p.id_producto,
                    p.codigo,
                    p.nombre,
                    c.nombre as categoria,
                    p.cantidad,
                    p.unidad_medida,
                    p.estado,
                    COALESCE(um.fecha_modificacion, NULL) as ultimo_movimiento,
                    CASE 
                        WHEN um.fecha_modificacion IS NULL THEN 1 
                        ELSE 0 
                    END as orden_nulos
                FROM Producto p
                LEFT JOIN Categoria c ON p.id_categoria = c.id_categoria
                LEFT JOIN UltimoMovimiento um ON p.id_producto = um.id_producto
                ORDER BY orden_nulos, ultimo_movimiento DESC
            """)
            productos = cursor.fetchall()
            
            # Convertir las fechas a objetos datetime si son strings
            productos = [list(p) for p in productos]
            for p in productos:
                if isinstance(p[7], str):
                    try:
                        p[7] = datetime.strptime(p[7], '%Y-%m-%d %H:%M:%S')
                    except (ValueError, TypeError):
                        p[7] = None
            
            return render_template("historial.html", productos=productos)
    except Exception as e:
        print("Error al cargar historial:", str(e))
        flash('Error al cargar el historial', 'danger')
        return redirect(url_for('index'))
    finally:
        conn.close()
@app.route("/api/historial/<int:id_producto>")
@requiere_login('admin')
def api_historial(id_producto):
    conn = conexion()
    try:
        with conn.cursor() as cursor:
            # Primero obtenemos el nombre del producto
            cursor.execute("""
                SELECT nombre FROM Producto WHERE id_producto = %s
            """, (id_producto,))
            producto = cursor.fetchone()
            
            if not producto:
                return jsonify({'error': 'Producto no encontrado'}), 404
            
            # Obtenemos todos los movimientos
            cursor.execute("""
                SELECT 
                    tipo_modificacion,
                    cantidad,
                    fecha_modificacion,
                    detalle
                FROM HistorialModificaciones
                WHERE id_producto = %s
                ORDER BY fecha_modificacion DESC
            """, (id_producto,))
            
            movimientos = cursor.fetchall()
            
            # Convertimos a formato JSON
            return jsonify([{
                'tipo_modificacion': mov[0],
                'cantidad': float(mov[1]),
                'fecha': mov[2].strftime('%d/%m/%Y %H:%M'),
                'detalle': mov[3]
            } for mov in movimientos])
            
    except Exception as e:
        print("Error al obtener historial:", str(e))
        return jsonify({'error': str(e)}), 500
    finally:
        conn.close()

@app.route("/devoluciones")
@requiere_login('admin')
def devoluciones():
    conn = conexion()
    try:
        with conn.cursor(pymysql.cursors.DictCursor):
            cursor.execute("""
                SELECT 
                    p.id_prestamo,
                    dp.id_producto,
                    pr.nombre as nombre_producto,
                    pr.unidad_medida,
                    dp.cantidad,
                    pt.nombre as nombre_solicitante,
                    pt.apellido as apellido_solicitante,
                    p.fecha_prestamo,
                    p.fecha_devolucion_planeada
                FROM Prestamo p
                JOIN DetallePrestamo dp ON p.id_prestamo = dp.id_prestamo
                JOIN Producto pr ON dp.id_producto = pr.id_producto
                JOIN Prestatario pt ON p.id_prestatario = pt.id_prestatario
                LEFT JOIN Devolucion d ON p.id_prestamo = d.id_prestamo
                WHERE d.id_devolucion IS NULL
                ORDER BY p.fecha_devolucion_planeada
            """)
            prestamos = cursor.fetchall()
            return render_template("devoluciones.html", 
                                 prestamos=prestamos,
                                 now=datetime.now())
    except Exception as e:
        print("Error al cargar devoluciones:", str(e))
        return "Error al cargar las devoluciones"
    finally:
        conn.close()

@app.route("/corregir_detalles_salida")
@requiere_login('admin')
def corregir_detalles_salida():
    conn = conexion()
    try:
        with conn.cursor() as cursor:
            # Actualizar todos los registros donde cantidad_disponible_retorno es 0 o NULL
            cursor.execute("""
                UPDATE DetalleSalida 
                SET cantidad_disponible_retorno = cantidad_salida 
                WHERE (cantidad_disponible_retorno = 0 OR cantidad_disponible_retorno IS NULL)
                AND cantidad_salida > 0
            """)
            
            conn.commit()
            return "Registros corregidos correctamente"
    except Exception as e:
        conn.rollback()
        return f"Error: {str(e)}"
    finally:
        conn.close()

@app.route("/procesar_devolucion", methods=["POST"])
@requiere_login('admin')
def procesar_devolucion():
    conn = conexion()
    try:
        detalles_ids = request.form.getlist('detalles[]')
        
        if not detalles_ids:
            flash('No se seleccionaron productos para devolver', 'warning')
            return redirect(url_for('devoluciones_pendientes'))
            
        with conn.cursor() as cursor:
            for detalle_id in detalles_ids:
                # Primero verificamos si ya existe una devolución para este detalle
                cursor.execute("""
                    SELECT COUNT(*) 
                    FROM Devolucion 
                    WHERE id_detalle_prestamo = %s
                """, (detalle_id,))
                
                if cursor.fetchone()[0] > 0:
                    continue  # Si ya existe una devolución, saltamos este detalle
                
                # Registrar la devolución
                cursor.execute("""
                    INSERT INTO Devolucion 
                    (id_detalle_prestamo, fecha_devolucion)
                    VALUES (%s, NOW())
                """, (detalle_id,))
                
                # Obtener información del préstamo y producto
                cursor.execute("""
                    SELECT 
                        dp.cantidad,
                        dp.id_producto,
                        p.id_prestamo
                    FROM DetallePrestamo dp
                    JOIN Prestamo p ON dp.id_prestamo = p.id_prestamo
                    WHERE dp.id_detalle_prestamo = %s
                """, (detalle_id,))
                
                detalle_prestamo = cursor.fetchone()
                if detalle_prestamo:
                    cantidad_devuelta = float(detalle_prestamo[0])
                    id_producto = detalle_prestamo[1]
                    id_prestamo = detalle_prestamo[2]
                    
                    # Actualizar stock del producto (una sola vez)
                    cursor.execute("""
                        UPDATE Producto 
                        SET cantidad = cantidad + %s,
                            estado = CASE 
                                WHEN cantidad + %s > 0 THEN 'disponible'
                                ELSE estado
                            END
                        WHERE id_producto = %s
                    """, (cantidad_devuelta, cantidad_devuelta, id_producto))
                    
                    # Registrar en historial
                    cursor.execute("""
                        INSERT INTO HistorialModificaciones 
                        (id_producto, tipo_modificacion, cantidad, fecha_modificacion, detalle)
                        VALUES (%s, 'devolucion', %s, NOW(), %s)
                    """, (
                        id_producto,
                        cantidad_devuelta,
                        f"Devolución del préstamo PRES-{str(id_prestamo).zfill(4)}"
                    ))
            
            conn.commit()
            flash('Devolución procesada exitosamente', 'success')
            return redirect(url_for('devolucion'))
            
    except Exception as e:
        conn.rollback()
        print(f"Error al procesar devolución: {str(e)}")
        flash(f'Error al procesar la devolución: {str(e)}', 'danger')
        return redirect(url_for('devoluciones_pendientes'))
    finally:
        conn.close()

@app.route("/producto_nuevo", methods=["GET", "POST"])
@requiere_login('admin')
def producto_nuevo():
    if request.method == "POST":
        conn = conexion()
        try:
            with conn.cursor() as cursor:
                # Insertar el producto
                cursor.execute("""
                    INSERT INTO Producto 
                    (nombre, descripcion, cantidad, unidad_medida, 
                     fecha_ingreso, procedencia, id_categoria)
                    VALUES (%s, %s, %s, %s, %s, %s, %s)
                """, (request.form['nombre'], request.form['descripcion'],
                      float(request.form['cantidad']), request.form['unidad_medida'],
                      datetime.now(), request.form['procedencia'],
                      request.form['categoria']))
                
                # Obtener el ID del producto recién insertado
                producto_id = cursor.lastrowid
                
                # Generar y actualizar el código
                cursor.execute("""
                    UPDATE Producto p
                    JOIN Categoria c ON p.id_categoria = c.id_categoria
                    SET p.codigo = CONCAT(
                        'ALM-',
                        LEFT(UPPER(REGEXP_REPLACE(c.nombre, '[^A-Za-z0-9]', '')), 3),
                        '-',
                        RIGHT(YEAR(CURRENT_DATE), 2),
                        '-',
                        LPAD((
                            SELECT COUNT(*) 
                            FROM (SELECT * FROM Producto) AS p2 
                            WHERE p2.id_categoria = p.id_categoria 
                            AND p2.id_producto <= p.id_producto
                        ), 4, '0')
                    )
                    WHERE p.id_producto = %s
                """, (producto_id,))
                
                # Obtener el código generado
                cursor.execute("SELECT codigo FROM Producto WHERE id_producto = %s", (producto_id,))
                codigo_generado = cursor.fetchone()[0]
                
                conn.commit()
                return redirect(url_for('productos'))
                
        except Exception as e:
            conn.rollback()
            print("Error al crear producto:", str(e))
            return "Error al crear el producto: " + str(e)
        finally:
            conn.close()
    else:
        return render_template("producto_nuevo.html", categorias=obtener_categorias())

# Reporte de Préstamos Activos
@app.route("/reportes")
@requiere_login('admin')
def reportes():
    try:
        conn = conexion()
        with conn.cursor() as cursor:
            # Contar productos con stock bajo
            cursor.execute("""
                SELECT COUNT(*) 
                FROM Producto 
                WHERE cantidad <= 10
            """)
            productos_bajos = cursor.fetchone()[0]
            
            # Contar préstamos activos
            cursor.execute("""
                SELECT COUNT(DISTINCT p.id_prestamo)
                FROM Prestamo p
                LEFT JOIN Devolucion d ON p.id_prestamo = d.id_prestamo
                WHERE d.id_devolucion IS NULL
            """)
            prestamos_activos = cursor.fetchone()[0]
            
            return render_template('reportes.html',
                                productos_bajos=productos_bajos,
                                prestamos_activos=prestamos_activos)
    except Exception as e:
        print("Error en reportes:", str(e))
        flash('Error al cargar los reportes', 'danger')
        return redirect(url_for('index'))
    finally:
        if conn:
            conn.close()

@app.route("/reportes/excel/prestamos")
@requiere_login('admin')
def excel_prestamos():
    try:
        excel_file = generar_excel_prestamos()
        return send_file(
            excel_file,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=f'Prestamos_Activos_{datetime.now().strftime("%Y%m%d_%H%M")}.xlsx'
        )
    except Exception as e:
        print("Error al generar excel de préstamos:", str(e))
        flash('Error al generar el reporte de préstamos', 'danger')
        return redirect(url_for('reportes'))

@app.route("/reportes/descargar_inventario")
@requiere_login('admin')
def descargar_inventario():
    try:
        excel_file = generar_excel_inventario()
        return send_file(
            excel_file,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=f'Inventario_{datetime.now().strftime("%Y%m%d_%H%M")}.xlsx'
        )
    except Exception as e:
        print("Error al descargar inventario:", str(e))
        flash('Error al generar el inventario', 'danger')
        return redirect(url_for('reportes'))

@app.route("/reportes/descargar_prestamos")
@requiere_login('admin')
def descargar_prestamos():
    try:
        excel_file = generar_excel_prestamos()
        return send_file(
            excel_file,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=f'Prestamos_Activos_{datetime.now().strftime("%Y%m%d_%H%M")}.xlsx'
        )
    except Exception as e:
        print("Error al descargar préstamos:", str(e))
        flash('Error al generar el reporte de préstamos', 'danger')
        return redirect(url_for('reportes'))

@app.route("/reportes/descargar_historial")
@requiere_login('admin')
def descargar_historial():
    try:
        excel_file = generar_excel_historial()
        return send_file(
            excel_file,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=f'Historial_{datetime.now().strftime("%Y%m%d_%H%M")}.xlsx'
        )
    except Exception as e:
        print("Error al descargar historial:", str(e))
        flash('Error al generar el historial', 'danger')
        return redirect(url_for('reportes'))

@app.route("/reportes/descargar_agotados")
@requiere_login('admin')
def descargar_agotados():
    try:
        excel_file = generar_excel_agotados()
        return send_file(
            excel_file,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=f'Productos_Agotados_{datetime.now().strftime("%Y%m%d_%H%M")}.xlsx'
        )
    except Exception as e:
        print("Error al descargar agotados:", str(e))
        flash('Error al generar el reporte de agotados', 'danger')
        return redirect(url_for('reportes'))

@app.route("/reportes/prestamos_activos")
@requiere_login('admin')
def prestamos_activos():
    try:
        excel_file = generar_excel_prestamos()
        return send_file(
            excel_file,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=f'Prestamos_Activos_{datetime.now().strftime("%Y%m%d_%H%M")}.xlsx'
        )
    except Exception as e:
        print("Error al generar reporte de préstamos:", str(e))
        flash('Error al generar el reporte de préstamos', 'danger')
        return redirect(url_for('reportes'))

def generar_excel_agotados():
    conn = conexion()
    wb = Workbook()
    ws = wb.active
    ws.title = "PRODUCTOS AGOTADOS"
    
    try:
        # Configuración de estilos
        title_font = Font(name='Arial', size=16, bold=True, color="FFFFFF")
        subtitle_font = Font(name='Arial', size=11, color="666666", bold=True)
        header_font = Font(name='Arial', size=11, bold=True, color="FFFFFF")
        data_font = Font(name='Arial', size=10)
        
        # Colores
        title_fill = PatternFill(start_color="C00000", end_color="C00000", fill_type="solid")  # Rojo oscuro
        subtitle_fill = PatternFill(start_color="F2F2F2", end_color="F2F2F2", fill_type="solid")
        header_fill = PatternFill(start_color="A50000", end_color="A50000", fill_type="solid")  # Rojo más oscuro
        
        # Bordes
        thin_border = Border(
            left=Side(style='thin', color="CCCCCC"),
            right=Side(style='thin', color="CCCCCC"),
            top=Side(style='thin', color="CCCCCC"),
            bottom=Side(style='thin', color="CCCCCC")
        )
        
        # Consulta SQL mejorada
        with conn.cursor() as cursor:
            cursor.execute("""
                SELECT 
                    p.codigo,
                    p.nombre,
                    c.nombre as categoria,
                    p.cantidad,
                    p.unidad_medida,
                    p.procedencia,
                    COALESCE(
                        (SELECT SUM(dp.cantidad)
                         FROM DetallePrestamo dp
                         JOIN Prestamo pr ON dp.id_prestamo = pr.id_prestamo
                         LEFT JOIN Devolucion d ON pr.id_prestamo = d.id_prestamo
                         WHERE dp.id_producto = p.id_producto
                         AND d.id_devolucion IS NULL), 0
                    ) as cantidad_prestada,
                    p.fecha_ingreso,
                    CASE 
                        WHEN p.cantidad = 0 AND p.estado = 'prestado' THEN 'Todo prestado'
                        ELSE 'Sin stock'
                    END as estado,
                    (SELECT MAX(fecha_modificacion)
                     FROM HistorialModificaciones
                     WHERE id_producto = p.id_producto) as ultima_modificacion
                FROM Producto p
                LEFT JOIN Categoria c ON p.id_categoria = c.id_categoria
                WHERE p.estado IN ('agotado', 'prestado') OR p.cantidad <= 0
                ORDER BY p.fecha_ingreso DESC, p.codigo
            """)
            productos = cursor.fetchall()
            
            # Título principal
            ws.merge_cells('A1:J1')
            title_cell = ws['A1']
            title_cell.value = "MUNICIPALIDAD DISTRITAL DE PIÁS"
            title_cell.font = title_font
            title_cell.fill = title_fill
            title_cell.alignment = Alignment(horizontal='center', vertical='center')
            
            # Subtítulo con fecha
            ws.merge_cells('A2:J2')
            subtitle_cell = ws['A2']
            subtitle_cell.value = f"REPORTE DE PRODUCTOS AGOTADOS - Generado el {datetime.now().strftime('%d/%m/%Y %H:%M')}"
            subtitle_cell.font = subtitle_font
            subtitle_cell.fill = subtitle_fill
            subtitle_cell.alignment = Alignment(horizontal='center', vertical='center')
            
            # Información adicional
            ws.merge_cells('A3:J3')
            info_cell = ws['A3']
            info_cell.value = "ALMACÉN CENTRAL - CONTROL DE INVENTARIO"
            info_cell.font = subtitle_font
            info_cell.alignment = Alignment(horizontal='center', vertical='center')
            
            # Encabezados
            headers = ['CÓDIGO', 'PRODUCTO', 'CATEGORÍA', 'STOCK', 'UM', 
                      'PROCEDENCIA', 'PRESTADO', 'ÚLTIMO INGRESO', 'ESTADO',
                      'ÚLTIMA MODIFICACIÓN']
            
            for col, header in enumerate(headers, 1):
                cell = ws.cell(row=5, column=col, value=header)
                cell.font = header_font
                cell.fill = header_fill
                cell.border = thin_border
                cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
            
            # Datos con colores alternados
            if not productos:
                ws.merge_cells('A6:J6')
                no_data_cell = ws['A6']
                no_data_cell.value = "No hay productos agotados en el inventario"
                no_data_cell.alignment = Alignment(horizontal='center')
                no_data_cell.font = data_font
            else:
                for row_idx, prod in enumerate(productos, 6):
                    row_color = "FFF2F2" if row_idx % 2 == 0 else "FFFFFF"  # Rojo muy claro alternado con blanco
                    for col_idx, value in enumerate(prod, 1):
                        cell = ws.cell(row=row_idx, column=col_idx, value=value)
                        cell.font = data_font
                        cell.fill = PatternFill(start_color=row_color, end_color=row_color, fill_type="solid")
                        cell.border = thin_border
                        
                        if isinstance(value, datetime):
                            cell.value = value.strftime('%d/%m/%Y %H:%M')
                        elif isinstance(value, float):
                            cell.value = f"{value:,.2f}"
                        
                        cell.alignment = Alignment(horizontal='center', vertical='center')
            
            # Ajustar anchos
            ws.column_dimensions['A'].width = 15  # CÓDIGO
            ws.column_dimensions['B'].width = 30  # PRODUCTO
            ws.column_dimensions['C'].width = 20  # CATEGORÍA
            ws.column_dimensions['D'].width = 10  # STOCK
            ws.column_dimensions['E'].width = 8   # UM
            ws.column_dimensions['F'].width = 20  # PROCEDENCIA
            ws.column_dimensions['G'].width = 12  # PRESTADO
            ws.column_dimensions['H'].width = 18  # ÚLTIMO INGRESO
            ws.column_dimensions['I'].width = 15  # ESTADO
            ws.column_dimensions['J'].width = 20  # ÚLTIMA MODIFICACIÓN
            
            # Altura de filas
            ws.row_dimensions[1].height = 35  # Título
            ws.row_dimensions[2].height = 25  # Subtítulo
            ws.row_dimensions[3].height = 25  # Info adicional
            ws.row_dimensions[5].height = 30  # Encabezados
            
            # Pie de página
            footer_row = len(productos) + 7
            ws.merge_cells(f'A{footer_row}:J{footer_row}')
            footer_cell = ws[f'A{footer_row}']
            footer_cell.value = "ÁREA DE LOGÍSTICA - ALMACÉN CENTRAL"
            footer_cell.font = Font(size=8, color="666666", bold=True)
            footer_cell.alignment = Alignment(horizontal='center')
                
    except Exception as e:
        print(f"Error en generar_excel_agotados: {str(e)}")
        raise
    finally:
        conn.close()
    
    excel_file = BytesIO()
    wb.save(excel_file)
    excel_file.seek(0)
    
    return excel_file

def generar_excel_inventario():
    conn = conexion()
    wb = Workbook()
    ws = wb.active
    ws.title = "INVENTARIO ACTUAL"
    
    try:
        # Estilos
        title_font = Font(name='Arial', size=16, bold=True, color="FFFFFF")
        subtitle_font = Font(name='Arial', size=11, color="666666", bold=True)
        header_font = Font(name='Arial', size=11, bold=True, color="FFFFFF")
        data_font = Font(name='Arial', size=10)
        
        # Colores
        title_fill = PatternFill(start_color="2F75B5", end_color="2F75B5", fill_type="solid")
        subtitle_fill = PatternFill(start_color="F2F2F2", end_color="F2F2F2", fill_type="solid")
        header_fill = PatternFill(start_color="305496", end_color="305496", fill_type="solid")
        
        # Bordes
        thin_border = Border(
            left=Side(style='thin', color="CCCCCC"),
            right=Side(style='thin', color="CCCCCC"),
            top=Side(style='thin', color="CCCCCC"),
            bottom=Side(style='thin', color="CCCCCC")
        )
        
        # Consulta SQL actualizada para manejar estados
        with conn.cursor() as cursor:
            cursor.execute("""
                SELECT 
                    p.codigo,
                    p.nombre,
                    c.nombre as categoria,
                    p.cantidad,
                    p.unidad_medida,
                    p.procedencia,
                    p.descripcion,
                    p.fecha_ingreso,
                    CASE 
                        WHEN p.cantidad = 0 AND EXISTS (
                            SELECT 1 
                            FROM DetallePrestamo dp
                            JOIN Prestamo pr ON dp.id_prestamo = pr.id_prestamo
                            LEFT JOIN Devolucion d ON dp.id_detalle_prestamo = d.id_detalle_prestamo
                            WHERE dp.id_producto = p.id_producto
                            AND d.id_devolucion IS NULL
                        ) THEN 'PRESTADO'
                        WHEN p.cantidad = 0 THEN 'AGOTADO'
                        ELSE 'DISPONIBLE'
                    END as estado
                FROM Producto p
                LEFT JOIN Categoria c ON p.id_categoria = c.id_categoria
                ORDER BY p.fecha_ingreso DESC, p.codigo ASC
            """)
            productos = cursor.fetchall()
            
            # Título principal
            ws.merge_cells('A1:I1')
            title_cell = ws['A1']
            title_cell.value = "MUNICIPALIDAD DISTRITAL DE PIÁS"
            title_cell.font = title_font
            title_cell.fill = title_fill
            title_cell.alignment = Alignment(horizontal='center', vertical='center')
            
            # Subtítulo con fecha
            ws.merge_cells('A2:I2')
            subtitle_cell = ws['A2']
            subtitle_cell.value = f"REPORTE DE INVENTARIO ACTUAL - Generado el {datetime.now().strftime('%d/%m/%Y %H:%M')}"
            subtitle_cell.font = subtitle_font
            subtitle_cell.fill = subtitle_fill
            subtitle_cell.alignment = Alignment(horizontal='center', vertical='center')
            
            # Información adicional
            ws.merge_cells('A3:I3')
            info_cell = ws['A3']
            info_cell.value = "ALMACÉN CENTRAL"
            info_cell.font = subtitle_font
            info_cell.alignment = Alignment(horizontal='center', vertical='center')
            
            # Encabezados
            headers = ['CÓDIGO', 'PRODUCTO', 'CATEGORÍA', 'CANTIDAD', 'UM', 
                      'PROCEDENCIA', 'DESCRIPCIÓN', 'FECHA INGRESO', 'ESTADO']
            
            for col, header in enumerate(headers, 1):
                cell = ws.cell(row=5, column=col, value=header)
                cell.font = header_font
                cell.fill = header_fill
                cell.border = thin_border
                cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
            
            # Datos con colores alternados
            if not productos:
                ws.merge_cells('A6:I6')
                no_data_cell = ws['A6']
                no_data_cell.value = "No hay productos registrados en el inventario"
                no_data_cell.alignment = Alignment(horizontal='center')
                no_data_cell.font = data_font
            else:
                for row_idx, prod in enumerate(productos, 6):
                    row_color = "F5F5F5" if row_idx % 2 == 0 else "FFFFFF"
                    for col_idx, value in enumerate(prod, 1):
                        cell = ws.cell(row=row_idx, column=col_idx, value=value)
                        cell.font = data_font
                        cell.fill = PatternFill(start_color=row_color, end_color=row_color, fill_type="solid")
                        cell.border = thin_border
                        
                        if isinstance(value, datetime):
                            cell.value = value.strftime('%d/%m/%Y %H:%M')
                        elif isinstance(value, float):
                            cell.value = f"{value:,.2f}"
                        
                        cell.alignment = Alignment(horizontal='center', vertical='center')
            
            # Ajustar anchos
            ws.column_dimensions['A'].width = 15  # CÓDIGO
            ws.column_dimensions['B'].width = 30  # PRODUCTO
            ws.column_dimensions['C'].width = 20  # CATEGORÍA
            ws.column_dimensions['D'].width = 12  # CANTIDAD
            ws.column_dimensions['E'].width = 8   # UM
            ws.column_dimensions['F'].width = 20  # PROCEDENCIA
            ws.column_dimensions['G'].width = 35  # DESCRIPCIÓN
            ws.column_dimensions['H'].width = 18  # FECHA INGRESO
            ws.column_dimensions['I'].width = 15  # ESTADO
            
            # Altura de filas
            ws.row_dimensions[1].height = 35  # Título
            ws.row_dimensions[2].height = 25  # Subtítulo
            ws.row_dimensions[3].height = 25  # Info adicional
            ws.row_dimensions[5].height = 30  # Encabezados
            
            # Pie de página
            footer_row = len(productos) + 7
            ws.merge_cells(f'A{footer_row}:I{footer_row}')
            footer_cell = ws[f'A{footer_row}']
            footer_cell.value = "ÁREA DE LOGÍSTICA - ALMACÉN CENTRAL"
            footer_cell.font = Font(name='Arial', size=8, bold=True, color='666666')
            footer_cell.alignment = Alignment(horizontal='center')
                
    except Exception as e:
        print(f"Error en generar_excel_inventario: {str(e)}")
        raise
    finally:
        conn.close()
    
    excel_file = BytesIO()
    wb.save(excel_file)
    excel_file.seek(0)
    
    return excel_file

def generar_excel_prestamos():
    conn = conexion()
    wb = Workbook()
    ws = wb.active
    ws.title = "PRÉSTAMOS ACTIVOS"
    
    try:
        # Estilos
        title_font = Font(name='Arial', size=16, bold=True, color="FFFFFF")
        subtitle_font = Font(name='Arial', size=11, color="666666", bold=True)
        header_font = Font(name='Arial', size=11, bold=True, color="FFFFFF")
        data_font = Font(name='Arial', size=10)
        
        # Colores
        title_fill = PatternFill(start_color="2F75B5", end_color="2F75B5", fill_type="solid")
        subtitle_fill = PatternFill(start_color="F2F2F2", end_color="F2F2F2", fill_type="solid")
        header_fill = PatternFill(start_color="305496", end_color="305496", fill_type="solid")
        
        # Consulta SQL mejorada
        with conn.cursor() as cursor:
            cursor.execute("""
                SELECT 
                    p.fecha_prestamo,
                    CONCAT(p.nombre_solicitante, ' ', p.apellido_solicitante) as solicitante,
                    p.dni_solicitante as dni,
                    p.telefono_solicitante as contacto,
                    pr.codigo,
                    pr.nombre as producto,
                    dp.cantidad,
                    pr.unidad_medida,
                    p.fecha_devolucion_planeada,
                    pr.procedencia,
                    dp.observacion,
                    CASE 
                        WHEN d.fecha_devolucion IS NOT NULL THEN 'Devuelto'
                        WHEN p.fecha_devolucion_planeada < NOW() THEN 'Vencido'
                        ELSE 'Activo'
                    END as estado
                FROM Prestamo p
                JOIN DetallePrestamo dp ON p.id_prestamo = dp.id_prestamo
                JOIN Producto pr ON dp.id_producto = pr.id_producto
                LEFT JOIN Devolucion d ON p.id_prestamo = d.id_prestamo
                WHERE d.id_devolucion IS NULL
                ORDER BY p.fecha_prestamo DESC, p.id_prestamo DESC
            """)
            prestamos = cursor.fetchall()
            
            # Título y encabezados
            ws.merge_cells('A1:L1')
            title_cell = ws['A1']
            title_cell.value = "MUNICIPALIDAD DISTRITAL DE PIÁS"
            title_cell.font = title_font
            title_cell.fill = title_fill
            title_cell.alignment = Alignment(horizontal='center', vertical='center')
            
            ws.merge_cells('A2:L2')
            subtitle_cell = ws['A2']
            subtitle_cell.value = f"REPORTE DE PRÉSTAMOS ACTIVOS - Generado el {datetime.now().strftime('%d/%m/%Y %H:%M')}"
            subtitle_cell.font = subtitle_font
            subtitle_cell.fill = subtitle_fill
            subtitle_cell.alignment = Alignment(horizontal='center', vertical='center')
            
            headers = ['FECHA PRÉSTAMO', 'SOLICITANTE', 'DNI', 'CONTACTO', 'CÓDIGO', 'PRODUCTO', 
                      'CANTIDAD', 'UM', 'FECHA DEVOLUCIÓN', 'PROCEDENCIA', 'OBSERVACIÓN', 'ESTADO']
            
            for col, header in enumerate(headers, 1):
                cell = ws.cell(row=4, column=col, value=header)
                cell.font = header_font
                cell.fill = header_fill
                cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
            
            # Datos
            for row_idx, pres in enumerate(prestamos, 5):
                row_color = "F5F5F5" if row_idx % 2 == 0 else "FFFFFF"
                for col_idx, value in enumerate(pres, 1):
                    cell = ws.cell(row=row_idx, column=col_idx, value=value)
                    cell.font = data_font
                    cell.fill = PatternFill(start_color=row_color, end_color=row_color, fill_type="solid")
                    cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                    
                    if isinstance(value, datetime):
                        cell.value = value.strftime('%d/%m/%Y %H:%M')
                    elif isinstance(value, float):
                        cell.value = f"{value:,.2f}"
            
            # Ajustar anchos
            ws.column_dimensions['A'].width = 18  # FECHA PRÉSTAMO
            ws.column_dimensions['B'].width = 30  # SOLICITANTE
            ws.column_dimensions['C'].width = 12  # DNI
            ws.column_dimensions['D'].width = 15  # CONTACTO
            ws.column_dimensions['E'].width = 15  # CÓDIGO
            ws.column_dimensions['F'].width = 30  # PRODUCTO
            ws.column_dimensions['G'].width = 12  # CANTIDAD
            ws.column_dimensions['H'].width = 8   # UM
            ws.column_dimensions['I'].width = 18  # FECHA DEVOLUCIÓN
            ws.column_dimensions['J'].width = 15  # PROCEDENCIA
            ws.column_dimensions['K'].width = 35  # OBSERVACIÓN
            ws.column_dimensions['L'].width = 12  # ESTADO
            
            # Pie de página
            last_row = len(prestamos) + 6
            ws.merge_cells(f'A{last_row}:L{last_row}')
            footer_cell = ws[f'A{last_row}']
            footer_cell.value = "ÁREA DE LOGÍSTICA - ALMACÉN MUNICIPAL"
            footer_cell.font = Font(name='Arial', size=10, bold=True)
            footer_cell.alignment = Alignment(horizontal='center')
            
    except Exception as e:
        print(f"Error en generar_excel_prestamos: {str(e)}")
        raise
    finally:
        conn.close()
    
    excel_file = BytesIO()
    wb.save(excel_file)
    excel_file.seek(0)
    
    return excel_file

def generar_excel_historial():
    conn = conexion()
    wb = Workbook()
    ws = wb.active
    ws.title = "HISTORIAL DE MOVIMIENTOS"
    
    try:
        # Estilos
        title_font = Font(name='Arial', size=16, bold=True, color="FFFFFF")
        subtitle_font = Font(name='Arial', size=11, color="666666", bold=True)
        header_font = Font(name='Arial', size=11, bold=True, color="FFFFFF")
        data_font = Font(name='Arial', size=10)
        
        # Colores por tipo de movimiento
        color_ingreso = "C6EFCE"      # Verde claro
        color_salida = "FFC7CE"       # Rojo claro
        color_prestamo = "B4C6E7"     # Azul claro
        color_devolucion = "FFEB9C"   # Amarillo claro
        
        # Estilos de título y encabezados
        title_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
        subtitle_fill = PatternFill(start_color="F2F2F2", end_color="F2F2F2", fill_type="solid")
        header_fill = PatternFill(start_color="2F5597", end_color="2F5597", fill_type="solid")
        
        # Bordes
        thin_border = Border(
            left=Side(style='thin', color="CCCCCC"),
            right=Side(style='thin', color="CCCCCC"),
            top=Side(style='thin', color="CCCCCC"),
            bottom=Side(style='thin', color="CCCCCC")
        )
        
        # Consulta SQL
        with conn.cursor() as cursor:
            cursor.execute("""
                SELECT 
                    h.fecha_modificacion,
                    p.codigo,
                    p.nombre as producto,
                    c.nombre as categoria,
                    h.tipo_modificacion,
                    h.cantidad,
                    p.unidad_medida,
                    h.detalle
                FROM HistorialModificaciones h
                JOIN Producto p ON h.id_producto = p.id_producto
                LEFT JOIN Categoria c ON p.id_categoria = c.id_categoria
                ORDER BY h.fecha_modificacion DESC, h.id_historial DESC
            """)
            movimientos = cursor.fetchall()
            
            # Título principal
            ws.merge_cells('A1:H1')
            title_cell = ws['A1']
            title_cell.value = "MUNICIPALIDAD DISTRITAL DE PIÁS"
            title_cell.font = title_font
            title_cell.fill = title_fill
            title_cell.alignment = Alignment(horizontal='center', vertical='center')
            
            # Subtítulo con fecha
            ws.merge_cells('A2:H2')
            subtitle_cell = ws['A2']
            subtitle_cell.value = f"HISTORIAL DE MOVIMIENTOS DE ALMACÉN - Generado el {datetime.now().strftime('%d/%m/%Y %H:%M')}"
            subtitle_cell.font = subtitle_font
            subtitle_cell.fill = subtitle_fill
            subtitle_cell.alignment = Alignment(horizontal='center', vertical='center')
            
            # Leyenda de colores
            ws.merge_cells('A3:B3')
            ws['A3'].value = "TIPOS DE MOVIMIENTO:"
            ws['A3'].font = Font(bold=True, size=10)
            
            legend_row = 3
            legend_items = [
                ("INGRESO", color_ingreso),
                ("SALIDA", color_salida),
                ("PRÉSTAMO", color_prestamo),
                ("DEVOLUCIÓN", color_devolucion)
            ]
            
            for col, (text, color) in enumerate(legend_items, 3):
                cell = ws.cell(row=legend_row, column=col)
                cell.value = text
                cell.fill = PatternFill(start_color=color, end_color=color, fill_type="solid")
                cell.font = Font(size=9, bold=True)
                cell.alignment = Alignment(horizontal='center')
                cell.border = thin_border
            
            # Encabezados
            headers = ['FECHA', 'CÓDIGO', 'PRODUCTO', 'CATEGORÍA', 'TIPO', 
                      'CANTIDAD', 'UM', 'DETALLE']
            
            for col, header in enumerate(headers, 1):
                cell = ws.cell(row=5, column=col, value=header)
                cell.font = header_font
                cell.fill = header_fill
                cell.border = thin_border
                cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
            
            # Datos con colores por tipo
            if not movimientos:
                ws.merge_cells('A6:H6')
                no_data_cell = ws['A6']
                no_data_cell.value = "No hay movimientos registrados"
                no_data_cell.alignment = Alignment(horizontal='center')
                no_data_cell.font = data_font
            else:
                for row_idx, mov in enumerate(movimientos, 6):
                    # Seleccionar color según tipo de movimiento
                    tipo = mov[4].lower()
                    if 'ingreso' in tipo:
                        row_color = color_ingreso
                    elif 'salida' in tipo:
                        row_color = color_salida
                    elif 'prestamo' in tipo:
                        row_color = color_prestamo
                    else:  # devolución
                        row_color = color_devolucion
                    
                    for col_idx, value in enumerate(mov, 1):
                        cell = ws.cell(row=row_idx, column=col_idx, value=value)
                        cell.font = data_font
                        cell.fill = PatternFill(start_color=row_color, end_color=row_color, fill_type="solid")
                        cell.border = thin_border
                        
                        if isinstance(value, datetime):
                            cell.value = value.strftime('%d/%m/%Y %H:%M')
                        elif isinstance(value, float):
                            cell.value = f"{value:,.2f}"
                        
                        cell.alignment = Alignment(horizontal='center', vertical='center')
            
            # Ajustar anchos
            ws.column_dimensions['A'].width = 18  # FECHA
            ws.column_dimensions['B'].width = 15  # CÓDIGO
            ws.column_dimensions['C'].width = 30  # PRODUCTO
            ws.column_dimensions['D'].width = 20  # CATEGORÍA
            ws.column_dimensions['E'].width = 12  # TIPO
            ws.column_dimensions['F'].width = 12  # CANTIDAD
            ws.column_dimensions['G'].width = 8   # UM
            ws.column_dimensions['H'].width = 40  # DETALLE
            
            # Altura de filas
            ws.row_dimensions[1].height = 35  # Título
            ws.row_dimensions[2].height = 25  # Subtítulo
            ws.row_dimensions[3].height = 20  # Leyenda
            ws.row_dimensions[5].height = 30  # Encabezados
            
            # Pie de página
            footer_row = len(movimientos) + 7
            ws.merge_cells(f'A{footer_row}:H{footer_row}')
            footer_cell = ws[f'A{footer_row}']
            footer_cell.value = "ÁREA DE LOGÍSTICA - CONTROL DE INVENTARIO"
            footer_cell.font = Font(size=8, color="666666", bold=True)
            footer_cell.alignment = Alignment(horizontal='center')
                
    except Exception as e:
        print(f"Error en generar_excel_historial: {str(e)}")
        raise
    finally:
        conn.close()
    
    excel_file = BytesIO()
    wb.save(excel_file)
    excel_file.seek(0)
    
    return excel_file

@app.context_processor
def utility_processor():
    return {
        'fecha_en_español': fecha_en_español
    }



@app.route("/ver_salidas")
@requiere_login('admin')
def ver_salidas():
    conn = conexion()
    try:
        with conn.cursor() as cursor:
            cursor.execute("""
                SELECT 
                    s.fecha_salida,
                    CONCAT(p.nombre, ' ', p.apellido) as solicitante,
                    p.dni,
                    pr.codigo,
                    pr.nombre as producto,
                    s.cantidad,
                    pr.unidad_medida,
                    s.motivo,
                    pr.procedencia
                FROM SalidaProducto s
                JOIN Prestatario p ON s.id_prestatario = p.id_prestatario
                JOIN Producto pr ON s.id_producto = pr.id_producto
                ORDER BY s.fecha_salida DESC
            """)
            salidas = cursor.fetchall()
            return render_template("ver_salidas.html", salidas=salidas)
    except Exception as e:
        print("Error al cargar salidas:", str(e))
        flash('Error al cargar las salidas', 'danger')
        return redirect(url_for('reportes'))
    finally:
        conn.close()

@app.route("/descargar_salidas")    
@requiere_login('admin')
def descargar_salidas():
    try:
        excel_file = generar_excel_salidas()
        return send_file(
            excel_file,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=f'Salidas_{datetime.now().strftime("%Y%m%d_%H%M")}.xlsx'
        )
    except Exception as e:
        print("Error al descargar salidas:", str(e))
        flash('Error al generar el reporte de salidas', 'danger')
        return redirect(url_for('reportes'))

def generar_excel_salidas():
    conn = conexion()
    wb = Workbook()
    ws = wb.active
    ws.title = "SALIDAS DE MATERIALES"
    
    try:
        # Estilos
        title_font = Font(name='Arial', size=16, bold=True, color="FFFFFF")
        subtitle_font = Font(name='Arial', size=11, color="666666", bold=True)
        header_font = Font(name='Arial', size=11, bold=True, color="FFFFFF")
        data_font = Font(name='Arial', size=10)
        
        # Colores
        title_fill = PatternFill(start_color="C00000", end_color="C00000", fill_type="solid")
        subtitle_fill = PatternFill(start_color="F2F2F2", end_color="F2F2F2", fill_type="solid")
        header_fill = PatternFill(start_color="A50000", end_color="A50000", fill_type="solid")
        
        # Bordes
        thin_border = Border(
            left=Side(style='thin', color="000000"),
            right=Side(style='thin', color="000000"),
            top=Side(style='thin', color="000000"),
            bottom=Side(style='thin', color="000000")
        )
        
        # Consulta SQL modificada
        with conn.cursor() as cursor:
            cursor.execute("""
                SELECT DISTINCT
                    ds.fecha_salida,
                    CONCAT(p.nombre, ' ', p.apellido) as solicitante,
                    p.dni,
                    p.telefono,
                    ds.area_usuario,
                    pr.codigo,
                    pr.nombre as producto,
                    ds.cantidad_salida as cantidad,
                    pr.unidad_medida,
                    s.motivo,
                    ds.observacion,
                    pr.procedencia
                FROM SalidaProducto s
                JOIN DetalleSalida ds ON s.id_salida = ds.id_salida
                JOIN Prestatario p ON s.id_prestatario = p.id_prestatario
                JOIN Producto pr ON ds.id_producto = pr.id_producto
                JOIN Categoria c ON pr.id_categoria = c.id_categoria
                WHERE c.nombre NOT IN ('Gasolina', 'Petróleo')
                ORDER BY ds.fecha_salida DESC
            """)
            salidas = cursor.fetchall()
            
            # Título y encabezados
            ws.merge_cells('A1:L1')
            title_cell = ws['A1']
            title_cell.value = "MUNICIPALIDAD DISTRITAL DE PIÁS"
            title_cell.font = title_font
            title_cell.fill = title_fill
            title_cell.alignment = Alignment(horizontal='center', vertical='center')
            title_cell.border = thin_border
            
            ws.merge_cells('A2:L2')
            subtitle_cell = ws['A2']
            subtitle_cell.value = f"REPORTE DE SALIDAS DE MATERIALES - Generado el {datetime.now().strftime('%d/%m/%Y %H:%M')}"
            subtitle_cell.font = subtitle_font
            subtitle_cell.fill = subtitle_fill
            subtitle_cell.alignment = Alignment(horizontal='center', vertical='center')
            subtitle_cell.border = thin_border
            
            headers = [
                'FECHA', 'SOLICITANTE', 'DNI', 'TELÉFONO', 'ÁREA USUARIO', 
                'CÓDIGO', 'PRODUCTO', 'CANTIDAD', 'UM', 'MOTIVO', 
                'OBSERVACIÓN', 'PROCEDENCIA'
            ]
            
            for col, header in enumerate(headers, 1):
                cell = ws.cell(row=4, column=col, value=header)
                cell.font = header_font
                cell.fill = header_fill
                cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                cell.border = thin_border
            
            # Datos
            for row_idx, salida in enumerate(salidas, 5):
                row_color = "FFF2F2" if row_idx % 2 == 0 else "FFFFFF"
                for col_idx, value in enumerate(salida, 1):
                    cell = ws.cell(row=row_idx, column=col_idx, value=value)
                    cell.font = data_font
                    cell.fill = PatternFill(start_color=row_color, end_color=row_color, fill_type="solid")
                    cell.border = thin_border
                    
                    if isinstance(value, datetime):
                        cell.value = value.strftime('%d/%m/%Y %H:%M')
                    elif isinstance(value, float):
                        cell.value = f"{value:,.2f}"
                    
                    cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
            
            # Ajustar anchos
            column_widths = {
                'A': 18,  # FECHA
                'B': 30,  # SOLICITANTE
                'C': 12,  # DNI
                'D': 12,  # TELÉFONO
                'E': 20,  # ÁREA USUARIO
                'F': 15,  # CÓDIGO
                'G': 30,  # PRODUCTO
                'H': 12,  # CANTIDAD
                'I': 8,   # UM
                'J': 35,  # MOTIVO
                'K': 30,  # OBSERVACIÓN
                'L': 20   # PROCEDENCIA
            }
            
            for col, width in column_widths.items():
                ws.column_dimensions[col].width = width
            
            # Agregar pie de página
            last_row = len(salidas) + 6
            ws.merge_cells(f'A{last_row}:L{last_row}')
            footer_cell = ws[f'A{last_row}']
            footer_cell.value = "ÁREA DE LOGÍSTICA - ALMACÉN MUNICIPAL"
            footer_cell.font = Font(name='Arial', size=10, bold=True)
            footer_cell.alignment = Alignment(horizontal='center')
            footer_cell.border = thin_border
            
    except Exception as e:
        print(f"Error en generar_excel_salidas: {str(e)}")
        raise
    finally:
        conn.close()
    
    excel_file = BytesIO()
    wb.save(excel_file)
    excel_file.seek(0)
    
    return excel_file

@app.route("/solicitudes_salida")
@requiere_login('admin')
def solicitudes_salida():
    conn = conexion()
    try:
        with conn.cursor() as cursor:
            cursor.execute("""
                SELECT 
                    s.fecha_salida,
                    s.id_salida,
                    ns.numero_secuencial,
                    CONCAT(p.nombre, ' ', p.apellido) as solicitante,
                    p.dni,
                    p.telefono,
                    MAX(ds.area_usuario) as area_usuario,
                    GROUP_CONCAT(DISTINCT CONCAT(pr.nombre, ' (', ds.cantidad_salida, ' ', 
                        pr.unidad_medida, ')') SEPARATOR ', ') as productos,
                    CASE 
                        WHEN s.motivo LIKE '%- Observación:%' 
                        THEN SUBSTRING_INDEX(s.motivo, '- Observación:', 1)
                        ELSE s.motivo 
                    END as motivo,
                    CASE 
                        WHEN s.motivo LIKE '%- Observación:%' 
                        THEN TRIM(SUBSTRING_INDEX(s.motivo, '- Observación:', -1))
                        ELSE NULL 
                    END as observacion_autorizacion
                FROM SalidaProducto s
                LEFT JOIN NumeroSalida ns ON s.id_salida = ns.id_salida
                JOIN Prestatario p ON s.id_prestatario = p.id_prestatario
                JOIN DetalleSalida ds ON s.id_salida = ds.id_salida
                JOIN Producto pr ON ds.id_producto = pr.id_producto
                JOIN Categoria c ON pr.id_categoria = c.id_categoria
                WHERE c.nombre NOT IN ('Gasolina', 'Petróleo')
                GROUP BY 
                    s.fecha_salida,
                    s.id_salida,
                    ns.numero_secuencial,
                    solicitante,
                    p.dni,
                    p.telefono,
                    s.motivo
                ORDER BY s.fecha_salida DESC, ns.numero_secuencial DESC
            """)
            solicitudes = cursor.fetchall()
            return render_template("solicitudes_salida.html", solicitudes=solicitudes)
    except Exception as e:
        print("Error al cargar solicitudes:", str(e))
        return render_template("solicitudes_salida.html", error="Error al cargar las solicitudes", solicitudes=[])
    finally:
        conn.close()
@app.route("/descargar_solicitud_salida/<int:id_salida>")
@requiere_login('admin')
def descargar_solicitud_salida(id_salida):
    conn = conexion()
    try:
        with conn.cursor() as cursor:
            # Consulta principal para obtener datos de la salida
            cursor.execute("""
                SELECT 
                    sp.id_salida,
                    sp.fecha_salida,
                    p.nombre,
                    p.apellido,
                    p.dni,
                    p.telefono,
                    sp.autorizado_por,
                    CASE 
                        WHEN sp.motivo LIKE '%%- Observación:%%' 
                        THEN SUBSTRING_INDEX(sp.motivo, ' - Observación:', 1)
                        ELSE sp.motivo 
                    END as motivo,
                    COALESCE(ns.numero_secuencial, 0) as numero_salida,
                    CASE 
                        WHEN sp.motivo LIKE '%%- Observación:%%' 
                        THEN SUBSTRING_INDEX(sp.motivo, ' - Observación:', -1)
                        ELSE '' 
                    END as observacion
                FROM SalidaProducto sp
                JOIN Prestatario p ON sp.id_prestatario = p.id_prestatario
                LEFT JOIN NumeroSalida ns ON sp.id_salida = ns.id_salida
                WHERE sp.id_salida = %s
            """, (id_salida,))
            
            salida = cursor.fetchone()
            if not salida:
                flash('Salida no encontrada', 'danger')
                return redirect(url_for('solicitudes_salida'))
            
            # Asignar número secuencial si no existe
            numero_salida = salida[8]
            if numero_salida == 0:
                numero_salida = asignar_numero_secuencial(
                    conn,
                    'NumeroSalida',
                    'id_salida',
                    id_salida,
                    'fecha_salida',
                    salida[1]
                )
            
            # Obtener los productos y sus detalles
            cursor.execute("""
                SELECT 
                    pr.codigo,
                    pr.nombre,
                    ds.cantidad_salida,
                    pr.unidad_medida,
                    ds.area_usuario,
                    ds.observacion
                FROM DetalleSalida ds
                JOIN Producto pr ON ds.id_producto = pr.id_producto
                WHERE ds.id_salida = %s
            """, (id_salida,))
            
            productos_raw = cursor.fetchall()
            
            if not productos_raw:
                flash('No se encontraron productos para esta salida', 'warning')
                return redirect(url_for('solicitudes_salida'))
            
            # Formatear productos para el documento
            productos = []
            for p in productos_raw:
                productos.append({
                    'codigo': p[0],
                    'nombre': p[1],
                    'cantidad': p[2],
                    'unidad': p[3],
                    'area_usuario': p[4] if p[4] else 'No especificada',
                    'observacion': p[5] if p[5] else 'Ninguna',
                    'motivo': str(salida[7]).replace('%', '%%')  # Escapar % para evitar errores de formato
                })

            # Generar el documento
            doc = generar_documento_salida(
                nombre=salida[2],
                apellido=salida[3],
                dni=salida[4],
                productos=productos,
                fecha_salida=salida[1],
                id_salida=salida[0],
                autorizado_por=salida[6] if salida[6] else 'No especificado',
                observacion_autorizacion=str(salida[9]).strip() or 'Ninguna',
                telefono=salida[5] if salida[5] else 'No especificado',
                numero_salida=numero_salida
            )

            conn.commit()

            # Preparar la respuesta con el documento
            response = send_file(
                doc,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                as_attachment=True,
                download_name=f'salida_{str(numero_salida).zfill(4)}.docx'
            )
            
            return response

    except Exception as e:
        conn.rollback()
        print(f"Error al descargar solicitud de salida: {str(e)}")
        flash('Error al generar el documento de salida', 'danger')
        return redirect(url_for('solicitudes_salida'))
    finally:
        conn.close()

@app.route("/salida_combustible")
@requiere_login('admin')
def salida_combustible():
    return render_template("salida_combustible.html")

@app.route("/salida_gasolina", methods=["GET", "POST"])
@requiere_login('admin')
def salida_gasolina():
    if request.method == "GET":
        conn = conexion()
        try:
            with conn.cursor(pymysql.cursors.DictCursor) as cursor:
                # Obtener productos de gasolina
                cursor.execute("""
                    SELECT p.* 
                    FROM Producto p
                    JOIN Categoria c ON p.id_categoria = c.id_categoria
                    WHERE c.nombre = 'Gasolina' 
                    AND p.estado = 'disponible'
                """)
                productos = cursor.fetchall()
                
                # Obtener áreas asegurando que no estén vacías
                cursor.execute("""
                    SELECT DISTINCT area 
                    FROM AreaUsuario 
                    WHERE area IS NOT NULL AND TRIM(area) != ''
                    ORDER BY area
                """)
                areas = cursor.fetchall()
                
                # Convertir el resultado a un formato más simple si es necesario
                areas_formateadas = [{'area': area['area']} for area in areas]
                
                return render_template(
                    "salida_gasolina.html",
                    productos=productos,
                    areas=areas_formateadas,
                    fecha_actual=datetime.now().strftime('%Y-%m-%d'),
                    hora_actual=datetime.now().strftime('%H:%M')
                )
        except Exception as e:
            print(f"Error al cargar datos: {str(e)}")
            flash("Error al cargar los datos", "danger")
            return redirect(url_for('index'))
        finally:
            conn.close()
    
    else:  # POST
        conn = conexion()
        try:
            # Obtener datos del formulario con los nombres correctos
            nombre_solicitante = request.form['nombre_solicitante']
            apellido_solicitante = request.form['apellido_solicitante']
            dni = request.form['dni']
            autorizado_por = request.form['autorizado_por']
            id_producto = request.form['id_producto']
            cantidad = float(request.form['cantidad'])
            area = request.form['area']
            placa = request.form['placa']
            unidad = request.form['unidad']
            motivo = request.form['motivo_salida']  # Cambiado de 'motivo' a 'motivo_salida'
            destino = request.form['destino']
            observacion = request.form.get('observacion', '')  # Campo opcional
            fecha = request.form['fecha_salida']  # Cambio de 'fecha' a 'fecha_salida'
            hora = request.form['hora_salida']    # Cambio de 'hora' a 'hora_salida'
            
            with conn.cursor() as cursor:
                # Verificar stock disponible
                cursor.execute("""
                    SELECT cantidad, nombre FROM Producto WHERE id_producto = %s
                """, (id_producto,))
                producto = cursor.fetchone()
                if not producto or float(producto[0]) < cantidad:
                    raise Exception("Stock insuficiente")
                
                nombre_producto = producto[1]
                
                # Registrar prestatario
                cursor.execute("""
                    INSERT INTO Prestatario (nombre, apellido, dni, telefono)
                    VALUES (%s, %s, %s, %s)
                """, (nombre_solicitante, apellido_solicitante, dni, ""))
                
                id_prestatario = cursor.lastrowid
                
                # Crear motivo completo
                motivo_completo = f"Unidad: {unidad} - Placa: {placa} - Área: {area} - Destino: {destino} - Motivo: {motivo}"
                if observacion:
                    motivo_completo += f" - Observación: {observacion}"
                
                # Registrar salida con el campo autorizado_por
                fecha_hora = datetime.strptime(f"{fecha} {hora}", "%Y-%m-%d %H:%M")
                cursor.execute("""
                    INSERT INTO SalidaProducto 
                    (id_prestatario, id_producto, cantidad, fecha_salida, motivo, autorizado_por)
                VALUES (%s, %s, %s, %s, %s, %s)
                """, (id_prestatario, id_producto, cantidad, fecha_hora, motivo_completo, autorizado_por))
                
                id_salida = cursor.lastrowid
                
                # Asignar número secuencial para gasolina
                cursor.execute("""
                    INSERT INTO NumeroSalidaGasolina 
                    (id_salida_gasolina, numero_secuencial, fecha_salida)
                    SELECT %s, 
                           COALESCE(MAX(numero_secuencial), 0) + 1,
                           %s
                    FROM NumeroSalidaGasolina
                    WHERE YEAR(fecha_salida) = YEAR(%s)
                """, (id_salida, fecha_hora, fecha_hora))
                
                # Obtener el número secuencial asignado
                cursor.execute("""
                    SELECT numero_secuencial 
                    FROM NumeroSalidaGasolina 
                    WHERE id_salida_gasolina = %s
                """, (id_salida,))
                
                numero_vale = cursor.fetchone()[0]
                
                # Actualizar stock
                cursor.execute("""
                    UPDATE Producto 
                    SET cantidad = cantidad - %s,
                        estado = CASE
                            WHEN cantidad - %s <= 0 THEN 'agotado'
                            ELSE estado
                        END
                    WHERE id_producto = %s
                """, (cantidad, cantidad, id_producto))
                
                # Registrar en historial
                cursor.execute("""
                    INSERT INTO HistorialModificaciones 
                    (id_producto, tipo_modificacion, cantidad, fecha_modificacion, detalle)
                    VALUES (%s, 'salida', %s, %s, %s)
                """, (
                    id_producto,
                    cantidad,
                    datetime.now(),
                    f"Salida de gasolina - Responsable: {nombre_solicitante} {apellido_solicitante} - DNI: {dni}"
                ))
                
                # Preparar datos para el documento
                datos_salida = {
                    "id_salida": numero_vale,  # Usar número secuencial en lugar de id
                    "responsable": nombre_solicitante + " " + apellido_solicitante,
                    "dni": dni,
                    "autorizado_por": autorizado_por,
                    "fecha": fecha_hora,
                    "hora": hora,
                    "placa": placa,
                    "unidad": unidad,
                    "area": area,
                    "tipo_combustible": nombre_producto,
                    "cantidad": cantidad,
                    "destino": destino,
                    "motivo": motivo,
                    "observacion": observacion
                }

                # Generar documento Word
                doc = generar_documento_salida_combustible(datos_salida, [])
                
                conn.commit()

                response = send_file(
                    BytesIO(doc.getvalue()),
                    mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                    as_attachment=True,
                    download_name=f'vale_gasolina_{str(numero_vale).zfill(4)}.docx'
                )
                
                response.headers["X-Download-Complete-Redirect"] = url_for('salida_gasolina')
                return response

        except Exception as e:
            conn.rollback()
            print(f"Error en salida_gasolina: {str(e)}")
            flash(f"Error al registrar la salida: {str(e)}", "danger")
            return redirect(url_for('salida_gasolina'))
        finally:
            conn.close()
@app.route("/salida_petroleo", methods=["GET", "POST"])
@requiere_login('admin')
def salida_petroleo():
    if request.method == "GET":
        conn = conexion()
        try:
            with conn.cursor(pymysql.cursors.DictCursor) as cursor:
                # Modificamos la consulta para mostrar todos los productos de petróleo
                cursor.execute("""
                    SELECT p.* 
                    FROM Producto p
                    JOIN Categoria c ON p.id_categoria = c.id_categoria
                    WHERE c.nombre = 'Petróleo' 
                    AND p.cantidad > 0  -- Solo productos con stock disponible
                    ORDER BY p.fecha_ingreso DESC  -- Ordenamos por fecha de ingreso
                """)
                productos = cursor.fetchall()
                
                cursor.execute("SELECT * FROM AreaUsuario ORDER BY area")
                areas = cursor.fetchall()
                
                return render_template(
                    "salida_petroleo.html",
                    productos=productos,
                    areas=areas,
                    fecha_actual=datetime.now().strftime('%Y-%m-%d'),
                    hora_actual=datetime.now().strftime('%H:%M')
                )   
        except Exception as e:
            print(f"Error al cargar datos: {str(e)}")
            flash("Error al cargar los datos", "danger")
            return redirect(url_for('index'))
        finally:
            conn.close()
    
    else:  # POST
        conn = conexion()
        try:
            with conn.cursor(pymysql.cursors.DictCursor) as cursor:
                # Obtener datos del formulario
                nombre_solicitante = request.form['nombre_solicitante']
                apellido_solicitante = request.form['apellido_solicitante']
                dni = request.form['dni']
                autorizado_por = request.form['autorizado_por']
                id_producto = request.form['id_producto']
                cantidad = float(request.form['cantidad'])
                area = request.form['area']
                placa = request.form['placa']
                motivo = request.form['motivo_salida']
                destino = request.form['destino']
                observacion = request.form.get('observacion', '')
                fecha = request.form['fecha_salida']
                hora = request.form['hora_salida']
                
                # Verificar stock disponible
                cursor.execute("""
                    SELECT cantidad, nombre FROM Producto WHERE id_producto = %s
                """, (id_producto,))
                producto = cursor.fetchone()
                
                if not producto or float(producto['cantidad']) < cantidad:
                    raise Exception("Stock insuficiente")
                
                nombre_producto = producto['nombre']
                
                # Registrar prestatario
                cursor.execute("""
                    INSERT INTO Prestatario (nombre, apellido, dni, telefono)
                    VALUES (%s, %s, %s, %s)
                """, (nombre_solicitante, apellido_solicitante, dni, ""))
                
                id_prestatario = cursor.lastrowid
                
                # Crear motivo completo
                motivo_completo = f"Placa: {placa} - Área: {area} - Destino: {destino} - Motivo: {motivo}"
                if observacion:
                    motivo_completo += f" - Observación: {observacion}"
                
                # Registrar salida
                fecha_hora = datetime.strptime(f"{fecha} {hora}", "%Y-%m-%d %H:%M")
                cursor.execute("""
                    INSERT INTO SalidaProducto 
                    (id_prestatario, id_producto, cantidad, fecha_salida, motivo, autorizado_por)
                    VALUES (%s, %s, %s, %s, %s, %s)
                """, (id_prestatario, id_producto, cantidad, fecha_hora, motivo_completo, autorizado_por))
                
                id_salida = cursor.lastrowid
                
                # Asignar número secuencial para petróleo
                cursor.execute("""
                    INSERT INTO NumeroSalidaPetroleo 
                    (id_salida_petroleo, numero_secuencial, fecha_salida)
                    SELECT %s, 
                           COALESCE(MAX(numero_secuencial), 0) + 1,
                           %s
                    FROM NumeroSalidaPetroleo
                    WHERE YEAR(fecha_salida) = YEAR(%s)
                """, (id_salida, fecha_hora, fecha_hora))
                
                # Obtener el número secuencial asignado
                cursor.execute("""
                    SELECT numero_secuencial 
                    FROM NumeroSalidaPetroleo 
                    WHERE id_salida_petroleo = %s
                """, (id_salida,))
                
                numero_vale = cursor.fetchone()['numero_secuencial']
                
                # Actualizar stock
                cursor.execute("""
                    UPDATE Producto 
                    SET cantidad = cantidad - %s,
                        estado = CASE
                            WHEN cantidad - %s <= 0 THEN 'agotado'
                            ELSE estado
                        END
                    WHERE id_producto = %s
                """, (cantidad, cantidad, id_producto))
                
                # Registrar en historial
                cursor.execute("""
                    INSERT INTO HistorialModificaciones 
                    (id_producto, tipo_modificacion, cantidad, fecha_modificacion, detalle)
                    VALUES (%s, 'salida', %s, %s, %s)
                """, (
                    id_producto,
                    cantidad,
                    datetime.now(),
                    f"Salida de petróleo - Responsable: {nombre_solicitante} {apellido_solicitante} - DNI: {dni}"
                ))
                
                # Preparar datos para el documento
                datos_salida = {
                    "id_salida": numero_vale,
                    "responsable": f"{nombre_solicitante} {apellido_solicitante}",
                    "dni": dni,
                    "autorizado_por": autorizado_por,
                    "fecha": fecha_hora,
                    "hora": hora,
                    "placa": placa,
                    "tipo_combustible": nombre_producto,
                    "cantidad": cantidad,
                    "area": area,
                    "motivo": motivo,
                    "destino": destino,
                    "observacion": observacion
                }

                # Generar documento
                doc = generar_documento_salida_combustible(datos_salida, [])
                
                # Confirmar transacción
                conn.commit()

                # Preparar respuesta con el documento
                response = send_file(
                    BytesIO(doc.getvalue()),
                    mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                    as_attachment=True,
                    download_name=f'vale_petroleo_{str(numero_vale).zfill(4)}.docx'
                )
                
                # Configurar headers para descarga y redirección
                response.headers["Cache-Control"] = "no-cache, no-store, must-revalidate"
                response.headers["Pragma"] = "no-cache"
                response.headers["Expires"] = "0"
                response.headers["Content-Disposition"] = f'attachment; filename=vale_petroleo_{str(numero_vale).zfill(4)}.docx'
                response.headers["X-Suggested-Filename"] = f'vale_petroleo_{str(numero_vale).zfill(4)}.docx'
                response.headers["X-Download-Complete-Redirect"] = url_for('index')
                
                return response

        except Exception as e:
            conn.rollback()
            print(f"Error en salida_petroleo: {str(e)}")
            flash(f"Error al registrar la salida: {str(e)}", "danger")
            return redirect(url_for('salida_petroleo'))
        finally:
            conn.close()    
@app.route("/agregar_area", methods=["POST"])
@requiere_login('admin')
def agregar_area():
    data = request.json
    nueva_area = data.get('area', '').strip().upper()  # Convertir a mayúsculas
    
    if not nueva_area:
        return jsonify({"success": False, "error": "El área no puede estar vacía"})
        
    conn = conexion()
    try:
        with conn.cursor() as cursor:
            # Verificar si el área ya existe (ignorando mayúsculas/minúsculas)
            cursor.execute("""
                SELECT id_area 
                FROM areausuario 
                WHERE UPPER(area) = %s
            """, (nueva_area,))
            
            if cursor.fetchone():
                return jsonify({
                    "success": False, 
                    "error": "El área ya existe en el sistema"
                })
            
            # Insertar nueva área
            cursor.execute("""
                INSERT INTO areausuario (area) 
                VALUES (%s)
            """, (nueva_area,))
            
            conn.commit()
            
            # Obtener el ID del área recién insertada
            area_id = cursor.lastrowid
            
            # Devolver la respuesta con los datos necesarios para el frontend
            return jsonify({
                "success": True,
                "message": "Área agregada correctamente",
                "area": {
                    "id_area": area_id,
                    "area": nueva_area
                }
            })
            
    except Exception as e:
        conn.rollback()
        print(f"Error al agregar área: {str(e)}")
        return jsonify({
            "success": False, 
            "error": f"Error al agregar el área: {str(e)}"
        })
        
    finally:
        conn.close()



def generar_documento_salida_combustible(datos_salida, productos_salida):
    doc = Document()
    
    # Configuración inicial del documento
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10)
    
    # Título
    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titulo_run = titulo.add_run('MUNICIPALIDAD DISTRITAL DE PIÁS')
    titulo_run.bold = True
    titulo_run.font.size = Pt(14)
    
    # Subtítulo
    subtitulo = doc.add_paragraph()
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Determinar tipo de combustible y formato de número
    es_petroleo = 'PETROLEO' in datos_salida['tipo_combustible'].upper()
    tipo_vale = 'PETRÓLEO' if es_petroleo else 'GASOLINA'
    
    subtitulo_run = subtitulo.add_run(f'VALE DE SALIDA DE {tipo_vale}')
    subtitulo_run.bold = True
    subtitulo_run.font.size = Pt(12)
    
    # Número de vale con formato específico según tipo
    num_vale = doc.add_paragraph()
    num_vale.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Verificar si el id_salida ya incluye el prefijo y año
    if isinstance(datos_salida["id_salida"], str) and ('-' in datos_salida["id_salida"]):
        numero_vale = datos_salida["id_salida"]  # Ya tiene el formato correcto
    else:
        # Si no tiene el formato, construirlo
        prefijo = 'P' if es_petroleo else 'G'
        año = datos_salida["fecha"].year
        numero = str(datos_salida["id_salida"]).zfill(4)
        numero_vale = f"{prefijo}-{numero}-{año}"
    
    num_vale_run = num_vale.add_run(f'N° {numero_vale}')
    num_vale_run.bold = True
    
    # Fecha
    fecha = doc.add_paragraph()
    fecha.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fecha.add_run(f'Piás, {fecha_en_español(datos_salida["fecha"])}')
    
    # Texto introductorio
    doc.add_paragraph(f'Por medio del presente documento se deja constancia de la entrega de {tipo_vale.lower()}, autorizado por la Unidad de Logística, según el siguiente detalle:')
    
    # Tabla de información
    tabla_info = doc.add_table(rows=1, cols=2)
    tabla_info.style = 'Table Grid'
    tabla_info.autofit = False
    tabla_info.allow_autofit = False
    tabla_info.columns[0].width = Inches(3.0)
    tabla_info.columns[1].width = Inches(3.0)
    
    # Headers de la tabla
    header_cells = tabla_info.rows[0].cells
    header_cells[0].text = "DATOS DEL SOLICITANTE"
    header_cells[1].text = "DATOS DE LA UNIDAD VEHICULAR"
    
    for cell in header_cells:
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Datos principales
    row = tabla_info.add_row().cells
    
    # Columna del solicitante (igual para ambos)
    solicitante_info = f"""Nombres y Apellidos: {datos_salida['responsable']}
DNI: {datos_salida['dni']}
Autorizado por: {datos_salida['autorizado_por']}
Área: {datos_salida['area']}"""
    row[0].text = solicitante_info
    
    # Columna de la unidad (diferente para petróleo y gasolina)
    if es_petroleo:
        # Para petróleo - solo placa y destino
        unidad_info = f"""N° de Placa: {datos_salida['placa']}
Destino: {datos_salida['destino']}"""
    else:
        # Para gasolina - incluye tipo de unidad
        unidad_info = f"""Tipo de Unidad: {datos_salida['unidad']}
N° de Placa: {datos_salida['placa']}
Destino: {datos_salida['destino']}"""
    row[1].text = unidad_info
    
    # Tabla de combustible
    doc.add_paragraph()
    tabla_combustible = doc.add_table(rows=1, cols=5)
    tabla_combustible.style = 'Table Grid'
    
    # Headers de la tabla de combustible
    header_cells = tabla_combustible.rows[0].cells
    header_cells[0].text = "Tipo de Combustible"
    header_cells[1].text = "Cantidad"
    header_cells[2].text = "Hora"
    header_cells[3].text = "Motivo"
    header_cells[4].text = "Observación"
    
    for cell in header_cells:
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Datos de combustible
    row = tabla_combustible.add_row().cells
    row[0].text = datos_salida['tipo_combustible']
    row[1].text = f"{datos_salida['cantidad']} GAL"
    row[2].text = datos_salida['hora']
    row[3].text = datos_salida['motivo']
    row[4].text = datos_salida.get('observacion', 'Ninguna')
    
    for cell in row:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Espacio para firmas
    doc.add_paragraph('\n\n')
    
    # Tabla de firmas (3 columnas)
    firma_table = doc.add_table(rows=1, cols=3)
    firma_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Firma del beneficiario
    firma_beneficiario = firma_table.rows[0].cells[0]
    p_beneficiario = firma_beneficiario.add_paragraph('_____________________')
    p_beneficiario.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_nombre_beneficiario = firma_beneficiario.add_paragraph(
        f'{datos_salida["responsable"]}\n'
        f'DNI: {datos_salida["dni"]}\n'
        'BENEFICIARIO'
    )
    p_nombre_beneficiario.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Firma del autorizante
    firma_autorizante = firma_table.rows[0].cells[1]
    p_autorizante = firma_autorizante.add_paragraph('_____________________')
    p_autorizante.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_nombre_autorizante = firma_autorizante.add_paragraph(
        f'{datos_salida["autorizado_por"]}\n'
        'AUTORIZANTE'
    )
    p_nombre_autorizante.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Firma del encargado de almacén
    firma_encargado = firma_table.rows[0].cells[2]
    p_encargado = firma_encargado.add_paragraph('_____________________')
    p_encargado.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_nombre_encargado = firma_encargado.add_paragraph(
        'FABIAN LOZANO CRUZ\n'
        'ASISTENTE DE ALMACÉN'
    )
    p_nombre_encargado.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Pie de página
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run(f'ÁREA DE LOGÍSTICA - ALMACÉN MUNICIPAL\nVALE DE {tipo_vale} {datos_salida["fecha"].year}')
    footer_run.font.size = Pt(8)
    footer_run.bold = True
    
    # Guardar documento
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output


@app.route("/reportes_combustible")
@requiere_login('admin')
def reportes_combustible():
    return render_template("reportes_combustible.html")

@app.route("/reportes_gasolina")
@requiere_login('admin')
def reportes_gasolina():
    conn = conexion()
    try:
        with conn.cursor() as cursor:
            # Primero asegurarse de que todas las salidas tengan número secuencial
            cursor.execute("""
                INSERT IGNORE INTO NumeroSalidaGasolina (id_salida_gasolina, numero_secuencial, fecha_salida)
                SELECT 
                    sp.id_salida,
                    COALESCE(
                        (
                            SELECT MAX(numero_secuencial) + 1
                            FROM NumeroSalidaGasolina nsg2
                            WHERE YEAR(nsg2.fecha_salida) = YEAR(sp.fecha_salida)
                        ),
                        1
                    ) as nuevo_numero,
                    sp.fecha_salida
                FROM SalidaProducto sp
                JOIN Producto pr ON sp.id_producto = pr.id_producto
                JOIN Categoria c ON pr.id_categoria = c.id_categoria
                LEFT JOIN NumeroSalidaGasolina nsg ON sp.id_salida = nsg.id_salida_gasolina
                WHERE c.nombre = 'Gasolina'
                AND nsg.id_salida_gasolina IS NULL
                ORDER BY sp.fecha_salida ASC
            """)
            
            # Luego obtener los registros con sus números asignados
            cursor.execute("""
                SELECT 
                    sp.id_salida,
                    COALESCE(nsg.numero_secuencial, 0) as numero_vale,
                    sp.fecha_salida,
                    CONCAT(p.nombre, ' ', p.apellido) as responsable,
                    p.dni,
                    SUBSTRING_INDEX(SUBSTRING_INDEX(sp.motivo, 'Placa: ', -1), ' - ', 1) as placa,
                    SUBSTRING_INDEX(SUBSTRING_INDEX(sp.motivo, 'Unidad: ', -1), ' - ', 1) as unidad,
                    pr.nombre as tipo_combustible,
                    sp.cantidad,
                    sp.autorizado_por,
                    DATE_FORMAT(sp.fecha_salida, '%Y-%m-%d') as fecha_formato,
                    YEAR(sp.fecha_salida) as año_vale
                FROM SalidaProducto sp
                JOIN Prestatario p ON sp.id_prestatario = p.id_prestatario
                JOIN Producto pr ON sp.id_producto = pr.id_producto
                JOIN Categoria c ON pr.id_categoria = c.id_categoria
                LEFT JOIN NumeroSalidaGasolina nsg ON sp.id_salida = nsg.id_salida_gasolina
                WHERE c.nombre = 'Gasolina'
                ORDER BY sp.fecha_salida DESC
            """)
            
            columns = [col[0] for col in cursor.description]
            registros = [dict(zip(columns, row)) for row in cursor.fetchall()]
            
            # Agregar placa_unidad para mantener compatibilidad
            for registro in registros:
                registro['placa_unidad'] = f"{registro['placa']} - {registro['unidad']}"
            
            conn.commit()  # Importante: confirmar la inserción de números secuenciales
            
            return render_template(
                "reportes_gasolina.html",
                registros=registros
            )
            
    except Exception as e:
        conn.rollback()
        print(f"Error en reportes_gasolina: {str(e)}")
        flash(f'Error al cargar los registros: {str(e)}', 'danger')
        return redirect(url_for('reportes'))
    finally:
        conn.close()
@app.route("/reportes_petroleo")
@requiere_login('admin')
def reportes_petroleo():
    conn = conexion()
    try:
        with conn.cursor() as cursor:
            # Primero asegurarse de que todas las salidas tengan número secuencial
            cursor.execute("""
                INSERT IGNORE INTO NumeroSalidaPetroleo (id_salida_petroleo, numero_secuencial, fecha_salida)
                SELECT 
                    sp.id_salida,
                    COALESCE(
                        (
                            SELECT MAX(numero_secuencial) + 1
                            FROM NumeroSalidaPetroleo nsp2
                            WHERE YEAR(nsp2.fecha_salida) = YEAR(sp.fecha_salida)
                        ),
                        1
                    ) as nuevo_numero,
                    sp.fecha_salida
                FROM SalidaProducto sp
                JOIN Producto pr ON sp.id_producto = pr.id_producto
                JOIN Categoria c ON pr.id_categoria = c.id_categoria
                LEFT JOIN NumeroSalidaPetroleo nsp ON sp.id_salida = nsp.id_salida_petroleo
                WHERE c.nombre = 'Petróleo'
                AND nsp.id_salida_petroleo IS NULL
                ORDER BY sp.fecha_salida ASC
            """)
            
            # Luego obtener los registros con sus números asignados
            cursor.execute("""
                SELECT 
                    sp.id_salida,
                    COALESCE(nsp.numero_secuencial, 0) as numero_vale,
                    sp.fecha_salida,
                    CONCAT(p.nombre, ' ', p.apellido) as responsable,
                    p.dni,
                    SUBSTRING_INDEX(SUBSTRING_INDEX(sp.motivo, 'Placa: ', -1), ' - ', 1) as placa,
                    pr.nombre as tipo_combustible,
                    sp.cantidad,
                    sp.autorizado_por,
                    DATE_FORMAT(sp.fecha_salida, '%Y-%m-%d') as fecha_formato,
                    YEAR(sp.fecha_salida) as año_vale
                FROM SalidaProducto sp
                JOIN Prestatario p ON sp.id_prestatario = p.id_prestatario
                JOIN Producto pr ON sp.id_producto = pr.id_producto
                JOIN Categoria c ON pr.id_categoria = c.id_categoria
                LEFT JOIN NumeroSalidaPetroleo nsp ON sp.id_salida = nsp.id_salida_petroleo
                WHERE c.nombre = 'Petróleo'
                ORDER BY sp.fecha_salida DESC
            """)
            
            columns = [col[0] for col in cursor.description]
            registros = [dict(zip(columns, row)) for row in cursor.fetchall()]
            
            conn.commit()  # Importante: confirmar la inserción de números secuenciales
            
            return render_template(
                "reportes_petroleo.html",
                registros=registros
            )
            
    except Exception as e:
        conn.rollback()
        print(f"Error en reportes_petroleo: {str(e)}")
        flash(f'Error al cargar los registros: {str(e)}', 'danger')
        return redirect(url_for('reportes'))
    finally:
        conn.close()
@app.route("/generar_reporte_combustible/<tipo>/<formato>")
@requiere_login('admin')
def generar_reporte_combustible(tipo, formato):
    conn = conexion()
    try:
        if formato == 'excel':
            # Crear libro y hoja de trabajo
            wb = Workbook()
            ws = wb.active
            ws.title = f"REGISTRO DE {tipo.upper()}"
            
            # Definir estilos
            header_fill = PatternFill(start_color="305496", end_color="305496", fill_type="solid")
            border = Border(
                left=Side(style='thin'),
                right=Side(style='thin'),
                top=Side(style='thin'),
                bottom=Side(style='thin')
            )
            
            # Título
            ws.merge_cells('A1:L1')
            title_cell = ws['A1']
            title_cell.value = "MUNICIPALIDAD DISTRITAL DE PIÁS"
            title_cell.font = Font(name='Arial', size=16, bold=True)
            title_cell.alignment = Alignment(horizontal='center')
            
            # Subtítulo
            ws.merge_cells('A2:L2')
            subtitle_cell = ws['A2']
            subtitle_cell.value = f"CONTROL DE {tipo.upper()} - {datetime.now().strftime('%Y')}"
            subtitle_cell.font = Font(name='Arial', size=14, bold=True)
            subtitle_cell.alignment = Alignment(horizontal='center')

            with conn.cursor() as cursor:
                # Actualizar consultas para incluir número secuencial
                if tipo.lower() == 'gasolina':
                    cursor.execute("""
                        SELECT 
                            DATE_FORMAT(sp.fecha_salida, '%d/%m/%Y') as fecha,
                            TIME(sp.fecha_salida) as hora,
                            CONCAT('G-', LPAD(COALESCE(nsg.numero_secuencial, 0), 4, '0'), '-', 
                                  YEAR(sp.fecha_salida)) as numero_vale,
                            CONCAT(p.nombre, ' ', p.apellido) as responsable,
                            p.dni,
                            pr.codigo as codigo_combustible,
                            SUBSTRING_INDEX(SUBSTRING_INDEX(sp.motivo, 'Placa: ', -1), ' - ', 1) as placa,
                            SUBSTRING_INDEX(SUBSTRING_INDEX(sp.motivo, 'Unidad: ', -1), ' - ', 1) as unidad,
                            SUBSTRING_INDEX(SUBSTRING_INDEX(sp.motivo, 'Área: ', -1), ' - ', 1) as area,
                            pr.nombre as tipo_combustible,
                            sp.cantidad,
                            'GAL' as unidad_medida,
                            sp.autorizado_por,
                            SUBSTRING_INDEX(SUBSTRING_INDEX(sp.motivo, 'Motivo: ', -1), ' - ', 1) as motivo
                        FROM SalidaProducto sp
                        JOIN Prestatario p ON sp.id_prestatario = p.id_prestatario
                        JOIN Producto pr ON sp.id_producto = pr.id_producto
                        JOIN Categoria c ON pr.id_categoria = c.id_categoria
                        LEFT JOIN NumeroSalidaGasolina nsg ON sp.id_salida = nsg.id_salida_gasolina
                        WHERE c.nombre = 'Gasolina'
                        ORDER BY sp.fecha_salida DESC
                    """)
                    headers = [
                        'Fecha', 'Hora', 'N° Vale', 'Responsable', 'DNI', 'Código', 'Placa', 
                        'Unidad', 'Área', 'Tipo Gasolina', 'Cantidad', 'U.M.', 'Autorizado por', 'Motivo'
                    ]
                    column_widths = [12, 8, 15, 25, 12, 15, 12, 15, 20, 15, 10, 8, 25, 35]
                else:
                    cursor.execute("""
                        SELECT 
                            DATE_FORMAT(sp.fecha_salida, '%d/%m/%Y') as fecha,
                            TIME(sp.fecha_salida) as hora,
                            CONCAT('P-', LPAD(COALESCE(nsp.numero_secuencial, 0), 4, '0'), '-', 
                                  YEAR(sp.fecha_salida)) as numero_vale,
                            CONCAT(p.nombre, ' ', p.apellido) as responsable,
                            p.dni,
                            pr.codigo as codigo_combustible,
                            SUBSTRING_INDEX(SUBSTRING_INDEX(sp.motivo, 'Placa: ', -1), ' - ', 1) as placa,
                            SUBSTRING_INDEX(SUBSTRING_INDEX(sp.motivo, 'Área: ', -1), ' - ', 1) as area,
                            pr.nombre as tipo_combustible,
                            sp.cantidad,
                            'GAL' as unidad_medida,
                            sp.autorizado_por,
                            SUBSTRING_INDEX(SUBSTRING_INDEX(sp.motivo, 'Motivo: ', -1), ' - ', 1) as motivo
                        FROM SalidaProducto sp
                        JOIN Prestatario p ON sp.id_prestatario = p.id_prestatario
                        JOIN Producto pr ON sp.id_producto = pr.id_producto
                        JOIN Categoria c ON pr.id_categoria = c.id_categoria
                        LEFT JOIN NumeroSalidaPetroleo nsp ON sp.id_salida = nsp.id_salida_petroleo
                        WHERE c.nombre = 'Petróleo'
                        ORDER BY sp.fecha_salida DESC
                    """)
                    headers = [
                        'Fecha', 'Hora', 'N° Vale', 'Responsable', 'DNI', 'Código', 'Placa', 
                        'Área', 'Tipo Petróleo', 'Cantidad', 'U.M.', 'Autorizado por', 'Motivo'
                    ]
                    column_widths = [12, 8, 15, 25, 12, 15, 12, 20, 15, 10, 8, 25, 35]

                registros = cursor.fetchall()

                # Encabezados
                for col, header in enumerate(headers, 1):
                    cell = ws.cell(row=5, column=col, value=header)
                    cell.font = Font(name='Arial', size=11, bold=True, color='FFFFFF')
                    cell.fill = header_fill
                    cell.alignment = Alignment(horizontal='center', vertical='center')
                    cell.border = border
                ws.row_dimensions[5].height = 30

                # Datos
                for row, registro in enumerate(registros, 6):
                    row_fill = PatternFill(start_color='F0F7FF' if row % 2 == 0 else 'FFFFFF', 
                                         end_color='F0F7FF' if row % 2 == 0 else 'FFFFFF', 
                                         fill_type='solid')
                    for col, valor in enumerate(registro, 1):
                        cell = ws.cell(row=row, column=col, value=valor)
                        cell.font = Font(name='Arial', size=10)
                        cell.alignment = Alignment(horizontal='center', vertical='center')
                        cell.border = border
                        cell.fill = row_fill

                # Ajustar anchos de columna
                for i, width in enumerate(column_widths, 1):
                    ws.column_dimensions[get_column_letter(i)].width = width

                # Pie de página
                footer_row = len(registros) + 7
                ws.merge_cells(f'A{footer_row}:L{footer_row}')
                footer_cell = ws[f'A{footer_row}']
                footer_cell.value = "ÁREA DE LOGÍSTICA - ALMACÉN MUNICIPAL"
                footer_cell.font = Font(name='Arial', size=8, bold=True, color='666666')
                footer_cell.alignment = Alignment(horizontal='center')

                # Generar archivo
                excel_file = BytesIO()
                wb.save(excel_file)
                excel_file.seek(0)

                return send_file(
                    excel_file,
                    mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                    as_attachment=True,
                    download_name=f'Control_{tipo}_{datetime.now().strftime("%Y%m%d")}.xlsx'
                )

        elif formato == 'word_salida':
            id_salida = request.args.get('id_salida')
            if not id_salida:
                flash('ID de salida no proporcionado', 'error')
                return redirect(url_for(f'reportes_{tipo}'))

            with conn.cursor() as cursor:
                # Obtener número secuencial de la tabla correspondiente
                if tipo == 'gasolina':
                    cursor.execute("""
                        SELECT numero_secuencial, fecha_salida
                        FROM NumeroSalidaGasolina
                        WHERE id_salida_gasolina = %s
                    """, (id_salida,))
                else:
                    cursor.execute("""
                        SELECT numero_secuencial, fecha_salida
                        FROM NumeroSalidaPetroleo
                        WHERE id_salida_petroleo = %s
                    """, (id_salida,))
                
                numero_info = cursor.fetchone()
                if not numero_info:
                    flash('No se encontró el número de vale', 'error')
                    return redirect(url_for(f'reportes_{tipo}'))
                
                numero_vale = numero_info[0]
                año_vale = numero_info[1].year

                # Luego obtener el resto de datos
                if tipo == 'gasolina':
                    cursor.execute("""
                        SELECT 
                            sp.id_salida,
                            sp.fecha_salida,
                            CONCAT(p.nombre, ' ', p.apellido) as responsable,
                            p.dni,
                            pr.nombre as tipo_combustible,
                            sp.cantidad,
                            SUBSTRING_INDEX(SUBSTRING_INDEX(sp.motivo, 'Unidad: ', -1), ' - ', 1) as unidad,
                            SUBSTRING_INDEX(SUBSTRING_INDEX(sp.motivo, 'Placa: ', -1), ' - ', 1) as placa,
                            SUBSTRING_INDEX(SUBSTRING_INDEX(sp.motivo, 'Área: ', -1), ' - ', 1) as area,
                            SUBSTRING_INDEX(SUBSTRING_INDEX(sp.motivo, 'Destino: ', -1), ' - ', 1) as destino,
                            SUBSTRING_INDEX(SUBSTRING_INDEX(sp.motivo, 'Motivo: ', -1), ' - ', 1) as motivo,
                            SUBSTRING_INDEX(sp.motivo, 'Observación: ', -1) as observacion,
                            sp.autorizado_por
                        FROM SalidaProducto sp
                        JOIN Prestatario p ON sp.id_prestatario = p.id_prestatario
                        JOIN Producto pr ON sp.id_producto = pr.id_producto
                        WHERE sp.id_salida = %s
                    """, (id_salida,))
                else:  # para petróleo
                    cursor.execute("""
                        SELECT 
                            sp.id_salida,
                            sp.fecha_salida,
                            CONCAT(p.nombre, ' ', p.apellido) as responsable,
                            p.dni,
                            pr.nombre as tipo_combustible,
                            sp.cantidad,
                            SUBSTRING_INDEX(SUBSTRING_INDEX(sp.motivo, 'Placa: ', -1), ' - ', 1) as placa,
                            SUBSTRING_INDEX(SUBSTRING_INDEX(sp.motivo, 'Área: ', -1), ' - ', 1) as area,
                            SUBSTRING_INDEX(SUBSTRING_INDEX(sp.motivo, 'Destino: ', -1), ' - ', 1) as destino,
                            SUBSTRING_INDEX(SUBSTRING_INDEX(sp.motivo, 'Motivo: ', -1), ' - ', 1) as motivo,
                            SUBSTRING_INDEX(sp.motivo, 'Observación: ', -1) as observacion,
                            sp.autorizado_por
                        FROM SalidaProducto sp
                        JOIN Prestatario p ON sp.id_prestatario = p.id_prestatario
                        JOIN Producto pr ON sp.id_producto = pr.id_producto
                        WHERE sp.id_salida = %s
                    """, (id_salida,))
                
                datos = cursor.fetchone()
                if datos:
                    # Crear diccionario según el tipo de combustible
                    if tipo == 'gasolina':
                        datos_salida = {
                            "id_salida": f"G-{str(numero_vale).zfill(4)}-{año_vale}",
                            "fecha": datos[1],
                            "responsable": datos[2],
                            "dni": datos[3],
                            "tipo_combustible": datos[4],
                            "cantidad": datos[5],
                            "unidad": datos[6],  # Incluido para gasolina
                            "placa": datos[7],
                            "area": datos[8],
                            "destino": datos[9],
                            "motivo": datos[10],
                            "observacion": datos[11],
                            "autorizado_por": datos[12],
                            "hora": datos[1].strftime('%H:%M')
                        }
                    else:  # para petróleo
                        datos_salida = {
                            "id_salida": f"P-{str(numero_vale).zfill(4)}-{año_vale}",
                            "fecha": datos[1],
                            "responsable": datos[2],
                            "dni": datos[3],
                            "tipo_combustible": datos[4],
                            "cantidad": datos[5],
                            "placa": datos[6],
                            "area": datos[7],
                            "destino": datos[8],
                            "motivo": datos[9],
                            "observacion": datos[10],
                            "autorizado_por": datos[11],
                            "hora": datos[1].strftime('%H:%M')
                        }
                    
                    # Generar documento Word
                    doc = generar_documento_salida_combustible(datos_salida, [])
                    
                    prefijo = 'G-' if tipo == 'gasolina' else 'P-'
                    return send_file(
                        BytesIO(doc.getvalue()),
                        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                        as_attachment=True,
                        download_name=f'vale_{tipo}_{prefijo}{str(numero_vale).zfill(4)}-{año_vale}.docx'
                    )
                else:
                    flash('No se encontró el registro de salida', 'error')
                    return redirect(url_for(f'reportes_{tipo}'))

    except Exception as e:
        print(f"Error detallado: {str(e)}")
        flash(f'Error al generar el reporte: {str(e)}', 'danger')
        return redirect(url_for(f'reportes_{tipo}'))
    finally:
        conn.close()

@app.route("/editar_descripcion", methods=["POST"])
@requiere_login('admin')
def editar_descripcion():
    conn = conexion()
    try:
        data = request.get_json()
        producto_id = data.get('producto_id')
        descripcion = data.get('descripcion', '')  # Valor por defecto vacío si no viene
        fecha_ingreso = data.get('fecha_ingreso')
        
        with conn.cursor() as cursor:
            # Actualizar producto permitiendo descripción NULL o vacía
            cursor.execute("""
                UPDATE Producto 
                SET descripcion = NULLIF(%s, ''), fecha_ingreso = %s 
                WHERE id_producto = %s
            """, (descripcion, fecha_ingreso, producto_id))
            
            # Registrar en historial
            cursor.execute("""
                INSERT INTO HistorialModificaciones 
                (id_producto, tipo_modificacion, cantidad, fecha_modificacion, detalle)
                VALUES (%s, 'ingreso', 0, NOW(), %s)
            """, (
                producto_id,
                f"Actualización de descripción y fecha de ingreso a: {fecha_ingreso}"
            ))
            
            conn.commit()
            return jsonify({"success": True})
            
    except Exception as e:
        conn.rollback()
        print(f"Error al editar descripción: {str(e)}")
        return jsonify({"success": False, "error": str(e)})
    finally:
        conn.close()
            
@app.route("/eliminar_producto/<int:id_producto>", methods=["POST"])
@requiere_login('admin')
def eliminar_producto(id_producto):
    conn = conexion()
    try:
        with conn.cursor() as cursor:
            # Verificar si el producto existe
            cursor.execute("""
                SELECT nombre, codigo 
                FROM Producto 
                WHERE id_producto = %s
            """, (id_producto,))
            
            producto = cursor.fetchone()
            if not producto:
                return jsonify({"success": False, "error": "Producto no encontrado"})
            
            # Verificar préstamos activos
            cursor.execute("""
                SELECT 
                    p.id_prestamo,
                    pr.nombre as nombre_prestatario,
                    pr.apellido as apellido_prestatario,
                    dp.cantidad,
                    p.fecha_prestamo
                FROM DetallePrestamo dp
                JOIN Prestamo p ON dp.id_prestamo = p.id_prestamo
                JOIN Prestatario pr ON p.id_prestatario = pr.id_prestatario
                LEFT JOIN Devolucion d ON dp.id_detalle_prestamo = d.id_detalle_prestamo
                WHERE dp.id_producto = %s AND d.id_devolucion IS NULL
            """, (id_producto,))
            
            prestamos = cursor.fetchall()
            if prestamos:
                # ... código de verificación de préstamos igual ...
                return jsonify({
                    "success": False,
                    "error": mensaje_error
                })
            
            # Eliminar registros relacionados en orden correcto
            
            # 1. Eliminar registros de devolución
            cursor.execute("""
                DELETE d FROM Devolucion d
                INNER JOIN DetallePrestamo dp ON d.id_detalle_prestamo = dp.id_detalle_prestamo
                WHERE dp.id_producto = %s
            """, (id_producto,))
            
            # 2. Eliminar registros de detalle de préstamo
            cursor.execute("""
                DELETE FROM DetallePrestamo 
                WHERE id_producto = %s
            """, (id_producto,))
            
            # 3. Eliminar registros de detalle de salida
            cursor.execute("""
                DELETE FROM DetalleSalida 
                WHERE id_producto = %s
            """, (id_producto,))
            
            # 4. Eliminar registros de salida de producto
            cursor.execute("""
                DELETE FROM SalidaProducto 
                WHERE id_producto = %s
            """, (id_producto,))
            
            # 5. Eliminar registros de detalle de donación
            cursor.execute("""
                DELETE FROM DetalleDonacion 
                WHERE id_producto = %s
            """, (id_producto,))
            
            # 6. Eliminar registros de historial
            cursor.execute("""
                DELETE FROM HistorialModificaciones 
                WHERE id_producto = %s
            """, (id_producto,))
            
            # 7. Eliminar registros de inventario
            cursor.execute("""
                DELETE FROM Inventario 
                WHERE id_producto = %s
            """, (id_producto,))
            
            # 8. Eliminar registros de alertas
            cursor.execute("""
                DELETE FROM Alerta 
                WHERE id_producto = %s
            """, (id_producto,))
            
            # 9. Finalmente eliminar el producto
            cursor.execute("""
                DELETE FROM Producto 
                WHERE id_producto = %s
            """, (id_producto,))
            
            conn.commit()
            
            return jsonify({
                "success": True, 
                "message": f"Producto {producto[1]} - {producto[0]} eliminado correctamente"
            })
            
    except Exception as e:
        conn.rollback()
        return jsonify({"success": False, "error": str(e)})
    finally:
        conn.close()

@app.route("/retornar_producto/<int:id_detalle_salida>", methods=["POST"])
@requiere_login('admin')
def retornar_producto(id_detalle_salida):
    conn = conexion()
    try:
        cantidad_retorno = float(request.form['cantidad_retorno'])
        
        with conn.cursor() as cursor:
            # Verificar cantidad disponible para retorno
            cursor.execute("""
                SELECT 
                    ds.cantidad_disponible_retorno,
                    ds.id_producto,
                    p.cantidad as stock_actual,
                    p.cantidad_inicial
                FROM DetalleSalida ds
                JOIN Producto p ON ds.id_producto = p.id_producto
                WHERE ds.id_detalle_salida = %s
            """, (id_detalle_salida,))
            
            resultado = cursor.fetchone()
            if not resultado:
                return jsonify({
                    "success": False,
                    "error": "Detalle de salida no encontrado"
                })
                
            disponible_retorno = float(resultado[0])
            id_producto = resultado[1]
            stock_actual = float(resultado[2])
            cantidad_inicial = float(resultado[3])
            
            if cantidad_retorno > disponible_retorno:
                return jsonify({
                    "success": False,
                    "error": "La cantidad a retornar excede la cantidad disponible"
                })
                
            if (stock_actual + cantidad_retorno) > cantidad_inicial:
                return jsonify({
                    "success": False,
                    "error": "El retorno excedería la cantidad inicial del producto"
                })
            
            # Actualizar cantidad disponible para retorno
            cursor.execute("""
                UPDATE DetalleSalida
                SET cantidad_disponible_retorno = cantidad_disponible_retorno - %s
                WHERE id_detalle_salida = %s
            """, (cantidad_retorno, id_detalle_salida))
            
            # Actualizar stock del producto
            cursor.execute("""
                UPDATE Producto
                SET cantidad = cantidad + %s
                WHERE id_producto = %s
            """, (cantidad_retorno, id_producto))
            
            # Registrar en historial
            cursor.execute("""
                INSERT INTO HistorialModificaciones 
                (id_producto, tipo_modificacion, cantidad, 
                 fecha_modificacion, detalle)
                VALUES (%s, 'retorno', %s, %s, %s)
            """, (
                id_producto,
                cantidad_retorno,
                datetime.now(),
                f"Retorno de producto - Detalle salida #{id_detalle_salida}"
            ))
            
            conn.commit()
            return jsonify({"success": True})
            
    except Exception as e:
        conn.rollback()
        return jsonify({"success": False, "error": str(e)})
    finally:
        conn.close()

@app.route("/devolver_producto", methods=["POST"])
@requiere_login('admin')
def devolver_producto():
    conn = conexion()
    try:
        data = request.get_json()
        if not data:
            return jsonify({
                "success": False,
                "message": "No se recibieron datos en la solicitud"
            })

        # Validación de datos recibidos
        producto_id = data.get('producto_id')
        cantidad_str = data.get('cantidad')
        observacion = data.get('observacion', '')

        # Validaciones específicas
        if not producto_id:
            return jsonify({
                "success": False,
                "message": "ID de producto no proporcionado"
            })

        if not cantidad_str:
            return jsonify({
                "success": False,
                "message": "Cantidad no proporcionada"
            })

        try:
            cantidad_devolver = float(cantidad_str)
            if cantidad_devolver <= 0:
                return jsonify({
                    "success": False,
                    "message": "La cantidad debe ser mayor a 0"
                })
        except (ValueError, TypeError):
            return jsonify({
                "success": False,
                "message": "Cantidad inválida"
            })

        with conn.cursor() as cursor:
            # Verificar que el producto existe
            cursor.execute("""
                SELECT id_producto 
                FROM Producto 
                WHERE id_producto = %s
            """, (producto_id,))
            
            if not cursor.fetchone():
                return jsonify({
                    "success": False,
                    "message": "Producto no encontrado"
                })

            # Verificar cantidad disponible para retorno
            cursor.execute("""
                SELECT 
                    COALESCE(SUM(cantidad_disponible_retorno), 0) as disponible
                FROM DetalleSalida 
                WHERE id_producto = %s 
                AND cantidad_disponible_retorno > 0
            """, (producto_id,))
            
            cantidad_disponible = float(cursor.fetchone()[0])
            
            if cantidad_disponible < cantidad_devolver:
                return jsonify({
                    "success": False,
                    "message": f"Solo hay {cantidad_disponible} unidades disponibles para devolución"
                })

            # Procesar la devolución
            cursor.execute("""
                UPDATE DetalleSalida 
                SET cantidad_disponible_retorno = cantidad_disponible_retorno - %s
                WHERE id_producto = %s 
                AND cantidad_disponible_retorno > 0
                ORDER BY fecha_salida ASC
                LIMIT 1
            """, (cantidad_devolver, producto_id))

            # Actualizar stock
            cursor.execute("""
                UPDATE Producto 
                SET cantidad = cantidad + %s
                WHERE id_producto = %s
            """, (cantidad_devolver, producto_id))

            # Registrar en historial
            cursor.execute("""
                INSERT INTO HistorialModificaciones 
                (id_producto, tipo_modificacion, cantidad, fecha_modificacion, detalle)
                VALUES (%s, 'devolucion', %s, NOW(), %s)
            """, (
                producto_id,
                cantidad_devolver,
                f"Devolución de salida - Observación: {observacion}"
            ))

            conn.commit()
            return jsonify({
                "success": True,
                "message": "Devolución procesada correctamente"
            })

    except Exception as e:
        conn.rollback()
        print(f"Error en devolver_producto: {str(e)}")
        return jsonify({
            "success": False,
            "message": str(e)
        })
    finally:
        conn.close()    

@app.route("/obtener_cantidad_retorno/<int:producto_id>")
@requiere_login('admin')
def obtener_cantidad_retorno(producto_id):
    conn = conexion()
    try:
        with conn.cursor() as cursor:
            cursor.execute("""
                SELECT COALESCE(SUM(cantidad_disponible_retorno), 0) as disponible_retorno
                FROM DetalleSalida
                WHERE id_producto = %s
                AND cantidad_disponible_retorno > 0
            """, (producto_id,))
            
            cantidad_disponible = float(cursor.fetchone()[0])
            
            return jsonify({
                "success": True,
                "cantidad_disponible": cantidad_disponible
            })
            
    except Exception as e:
        return jsonify({
            "success": False,
            "error": str(e)
        })
    finally:
        conn.close()

@app.route("/devoluciones_pendientes")
@requiere_login('admin')
def devoluciones_pendientes():
    conn = conexion()
    try:
        with conn.cursor() as cursor:
            cursor.execute("""
                SELECT 
                    p.id_prestamo,
                    COALESCE(np.numero_secuencial, 0) as numero_prestamo,
                    pr.id_producto,
                    pr.codigo,
                    pr.nombre as nombre_producto,
                    dp.cantidad,
                    pr.unidad_medida,
                    pre.nombre,
                    pre.apellido,
                    p.fecha_prestamo,
                    p.fecha_devolucion_planeada,
                    dp.id_detalle_prestamo,
                    dp.observacion
                FROM Prestamo p
                LEFT JOIN NumeroPrestamo np ON p.id_prestamo = np.id_prestamo
                JOIN DetallePrestamo dp ON p.id_prestamo = dp.id_prestamo
                JOIN Producto pr ON dp.id_producto = pr.id_producto
                JOIN Prestatario pre ON p.id_prestatario = pre.id_prestatario
                LEFT JOIN Devolucion d ON dp.id_detalle_prestamo = d.id_detalle_prestamo
                WHERE d.id_devolucion IS NULL
                ORDER BY p.fecha_prestamo DESC
            """)
            prestamos = cursor.fetchall()
            
            # Invertir la lista para mostrar los más recientes primero
            prestamos = list(reversed(prestamos))
            
            return render_template(
                "devoluciones_pendientes.html",
                prestamos=prestamos,
                now=datetime.now()
            )
    except Exception as e:
        print(f"Error en devoluciones_pendientes: {str(e)}")
        flash("Error al cargar las devoluciones pendientes", "danger")
        return redirect(url_for('index'))
    finally:
        conn.close()
@app.route("/donacion", methods=['GET', 'POST'])
@requiere_login('admin')
def donacion():
    conn = conexion()
    try:
        with conn.cursor() as cursor:
            cursor.execute("""
                SELECT 
                    p.id_producto,
                    p.codigo,
                    p.nombre,
                    p.descripcion,
                    p.cantidad,
                    p.unidad_medida,
                    p.procedencia,
                    p.estado,
                    CASE
                        WHEN p.cantidad <= 0 AND (
                            SELECT COALESCE(SUM(dp.cantidad), 0)
                            FROM DetallePrestamo dp
                            JOIN Prestamo pr ON dp.id_prestamo = pr.id_prestamo
                            LEFT JOIN Devolucion d ON pr.id_prestamo = d.id_prestamo
                            WHERE dp.id_producto = p.id_producto
                            AND d.id_devolucion IS NULL
                        ) > 0 THEN 'prestado'
                        WHEN p.cantidad <= 0 THEN 'agotado'
                        ELSE 'disponible'
                    END as estado_actual,
                    c.nombre as categoria_nombre
                FROM Producto p
                LEFT JOIN Categoria c ON p.id_categoria = c.id_categoria
                GROUP BY 
                    p.id_producto, p.codigo, p.nombre, p.descripcion, p.cantidad,
                    p.unidad_medida, p.procedencia, p.estado, c.nombre
                ORDER BY p.nombre
            """)
            productos = cursor.fetchall()
            
            if not productos:
                flash('No hay productos disponibles para donación', 'warning')
                return render_template('donacion.html', productos=[])
                
            return render_template('donacion.html', productos=productos)
            
    except Exception as e:
        flash(f'Error al cargar los productos: {str(e)}', 'danger')
        return render_template('donacion.html', productos=[])
    finally:
        conn.close()

@app.route("/confirmar_donacion", methods=['POST'])
@requiere_login('admin')
def confirmar_donacion():
    if request.method == 'POST':
        productos_seleccionados = request.form.getlist('productos[]')
        if not productos_seleccionados:
            flash('No se han seleccionado productos', 'warning')
            return redirect(url_for('donacion'))
            
        conn = None
        try:
            conn = conexion()
            cursor = conn.cursor()
            productos_info = []
            
            for id_producto in productos_seleccionados:
                cursor.execute("""
                    SELECT 
                        id_producto,
                        codigo,
                        nombre,
                        descripcion,
                        cantidad,
                        unidad_medida
                    FROM Producto 
                    WHERE id_producto = %s
                """, (id_producto,))
                
                producto = cursor.fetchone()
                if producto:
                    productos_info.append({
                        'id_producto': producto[0],
                        'codigo': producto[1],
                        'nombre': producto[2],
                        'descripcion': producto[3],
                        'cantidad': producto[4],
                        'unidad_medida': producto[5]
                    })
            
            if not productos_info:
                flash('No se encontraron los productos seleccionados', 'warning')
                return redirect(url_for('donacion'))
                
            return render_template('confirmar_donacion.html', productos=productos_info)
            
        except Exception as e:
            print(f"Error en confirmar_donacion: {str(e)}")  # Debug print
            flash('Error al procesar los productos seleccionados', 'danger')
            return redirect(url_for('donacion'))
        finally:
            if conn:
                conn.close()
    
    return redirect(url_for('donacion'))

@app.route("/procesar_donacion", methods=['POST'])
@requiere_login('admin')
def procesar_donacion():
    if request.method == 'POST':
        conn = conexion()
        try:
            # Obtener datos del formulario
            nombre = request.form['nombre_solicitante']
            apellidos = request.form['apellidos_solicitante']
            dni = request.form['dni']
            telefono = request.form.get('telefono_solicitante', '')
            fecha_salida = request.form['fecha_salida']
            hora_salida = request.form['hora_salida']
            autorizado_por = request.form['autorizado_por']
            observacion_autorizacion = request.form.get('observacion_autorizacion', '')
            motivo = request.form['motivo_donacion']
            productos = request.form.getlist('productos[]')
            
            productos_info = []
            
            with conn.cursor() as cursor:
                # Registrar prestatario
                cursor.execute("""
                    INSERT INTO Prestatario (nombre, apellido, dni, telefono)
                    VALUES (%s, %s, %s, %s)
                    ON DUPLICATE KEY UPDATE id_prestatario=LAST_INSERT_ID(id_prestatario)
                """, (nombre, apellidos, dni, telefono))
                id_prestatario = cursor.lastrowid
                
                # Registrar donación
                cursor.execute("""
                    INSERT INTO Donacion (id_prestatario, fecha_donacion, motivo, 
                                        autorizado_por, observacion_autorizacion)
                    VALUES (%s, NOW(), %s, %s, %s)
                """, (id_prestatario, motivo, autorizado_por, observacion_autorizacion))
                id_donacion = cursor.lastrowid

                # Asignar número secuencial
                numero_donacion = asignar_numero_secuencial(
                    conn,
                    'NumeroDonacion',
                    'id_donacion',
                    id_donacion,
                    'fecha_donacion',
                    datetime.now()
                )
                
                # Procesar productos
                for id_producto in productos:
                    cantidad_donacion = float(request.form[f'cantidad_{id_producto}'])
                    observacion = request.form.get(f'observacion_producto_{id_producto}', '')
                    
                    # Obtener información del producto
                    cursor.execute("""
                        SELECT codigo, nombre, unidad_medida 
                        FROM Producto 
                        WHERE id_producto = %s
                    """, (id_producto,))
                    prod_info = cursor.fetchone()
                    
                    productos_info.append({
                        'codigo': prod_info[0],
                        'nombre': prod_info[1],
                        'cantidad': cantidad_donacion,
                        'unidad': prod_info[2],
                        'observacion': observacion
                    })
                    
                    # Registrar detalle de donación
                    cursor.execute("""
                        INSERT INTO DetalleDonacion (id_donacion, id_producto, cantidad, observacion)
                        VALUES (%s, %s, %s, %s)
                    """, (id_donacion, id_producto, cantidad_donacion, observacion))
                    
                    # Actualizar stock
                    cursor.execute("""
                        UPDATE Producto 
                        SET cantidad = cantidad - %s 
                        WHERE id_producto = %s
                    """, (cantidad_donacion, id_producto))
                    
                    # Registrar en historial
                    cursor.execute("""
                        INSERT INTO HistorialModificaciones 
                        (id_producto, tipo_modificacion, cantidad, fecha_modificacion, detalle)
                        VALUES (%s, 'donacion', %s, NOW(), %s)
                    """, (
                        id_producto,
                        cantidad_donacion,
                        f"Donación a {nombre} {apellidos} - DNI: {dni}"
                    ))
                
                conn.commit()
                
                # Generar documento
                fecha_limpia = fecha_salida.split('T')[0] if 'T' in fecha_salida else fecha_salida
                doc = generar_documento_donacion({
                    'id_donacion': id_donacion,
                    'numero_donacion': numero_donacion,
                    'nombre': nombre,
                    'apellidos': apellidos,
                    'dni': dni,
                    'telefono': telefono,
                    'fecha_salida': fecha_limpia,
                    'hora_salida': hora_salida,
                    'autorizado_por': autorizado_por,
                    'observacion_autorizacion': observacion_autorizacion,
                    'motivo': motivo
                }, productos_info)
                
                flash('Donación registrada exitosamente', 'success')
                
                return send_file(
                    doc,
                    mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                    as_attachment=True,
                    download_name=f'donacion_{str(numero_donacion).zfill(4)}.docx'
                )
                
        except Exception as e:
            conn.rollback()
            print(f"Error al procesar donación: {str(e)}")
            flash(f'Error al procesar la donación: {str(e)}', 'danger')
            return redirect(url_for('donacion'))
        finally:
            conn.close()
    
    return redirect(url_for('donacion'))

def generar_documento_donacion(datos, productos):
    doc = Document()
    
    # Configuración de estilos
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    
    # Número de donación en esquina superior derecha
    num_donacion = doc.add_paragraph()
    num_donacion.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    num_run = num_donacion.add_run(f'DONACIÓN N° {str(datos["numero_donacion"]).zfill(4)}')
    num_run.font.size = Pt(11)
    num_run.bold = True
    
    # Título
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run('\nCONSTANCIA DE DONACIÓN DE MATERIALES\n')
    title_run.bold = True
    title_run.font.size = Pt(14)
    
    # Fecha y hora
    fecha_hora = doc.add_paragraph()
    fecha_hora.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    try:
        # Convertir la fecha correctamente
        fecha_obj = datetime.strptime(datos["fecha_salida"].split('T')[0], "%Y-%m-%d")
        fecha_hora.add_run(f'Piás, {fecha_en_español(fecha_obj)}')
        fecha_hora.add_run(f'\nHora: {datos["hora_salida"]}')
    except Exception as e:
        print(f"Error al procesar fecha: {str(e)}")
        # Usar fecha actual como respaldo
        fecha_hora.add_run(f'Piás, {fecha_en_español(datetime.now())}')
        fecha_hora.add_run(f'\nHora: {datetime.now().strftime("%H:%M")}')
    
    # Datos del beneficiario
    doc.add_paragraph()
    datos_beneficiario = doc.add_paragraph()
    datos_beneficiario.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    datos_beneficiario.add_run('Por medio del presente documento se deja constancia que el Sr(a). ')
    datos_beneficiario.add_run(f'{datos["nombre"]} {datos["apellidos"]}').bold = True
    datos_beneficiario.add_run(', identificado con DNI N° ')
    datos_beneficiario.add_run(datos["dni"]).bold = True
    datos_beneficiario.add_run(', con número de teléfono ')
    datos_beneficiario.add_run(datos["telefono"]).bold = True
    datos_beneficiario.add_run(', ha recibido en calidad de DONACIÓN los siguientes materiales del Almacén Municipal:')
    
    doc.add_paragraph()
    
    # Tabla de materiales
    tabla = doc.add_table(rows=1, cols=5)
    tabla.style = 'Table Grid'
    tabla.autofit = True
    
    # Encabezados de tabla
    encabezados = tabla.rows[0].cells
    for i, texto in enumerate(['CÓDIGO', 'DESCRIPCIÓN', 'CANTIDAD', 'U.M.', 'OBSERVACIÓN']):
        encabezados[i].text = texto
        encabezados[i].paragraphs[0].runs[0].bold = True
        encabezados[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Productos
    for producto in productos:
        row = tabla.add_row().cells
        row[0].text = producto['codigo']
        row[1].text = producto['nombre']
        row[2].text = str(producto['cantidad'])
        row[3].text = producto['unidad']
        row[4].text = producto['observacion']
        for cell in row:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Motivo de la donación
    doc.add_paragraph()
    motivo_parrafo = doc.add_paragraph()
    motivo_parrafo.add_run('MOTIVO DE LA DONACIÓN: ').bold = True
    motivo_parrafo.add_run(datos['motivo'])
    
    # Autorización
    doc.add_paragraph()
    auth_parrafo = doc.add_paragraph()
    auth_parrafo.add_run('AUTORIZADO POR: ').bold = True
    auth_parrafo.add_run(datos['autorizado_por'])
    
    # Observación de la autorización
    if datos['observacion_autorizacion']:
        doc.add_paragraph()
        obs_auth_parrafo = doc.add_paragraph()
        obs_auth_parrafo.add_run('OBSERVACIÓN DE LA AUTORIZACIÓN: ').bold = True
        obs_auth_parrafo.add_run(datos['observacion_autorizacion'])
    
    # Espacio para firmas
    doc.add_paragraph('\n\n')
    
    # Tabla de firmas (3 columnas)
    firma_table = doc.add_table(rows=1, cols=3)
    firma_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Firma del beneficiario
    firma_beneficiario = firma_table.rows[0].cells[0]
    p_beneficiario = firma_beneficiario.add_paragraph('_____________________')
    p_beneficiario.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_nombre_beneficiario = firma_beneficiario.add_paragraph(
        f'{datos["nombre"]} {datos["apellidos"]}\n'
        f'DNI: {datos["dni"]}\n'
        'BENEFICIARIO'
    )
    p_nombre_beneficiario.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Firma del autorizante
    firma_autorizante = firma_table.rows[0].cells[1]
    p_autorizante = firma_autorizante.add_paragraph('_____________________')
    p_autorizante.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_nombre_autorizante = firma_autorizante.add_paragraph(
        f'{datos["autorizado_por"]}\n'
        'AUTORIZANTE'
    )
    p_nombre_autorizante.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Firma del encargado de almacén
    firma_encargado = firma_table.rows[0].cells[2]
    p_encargado = firma_encargado.add_paragraph('_____________________')
    p_encargado.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_nombre_encargado = firma_encargado.add_paragraph(
        'FABIAN LOZANO CRUZ\n'
        'ASISTENTE DE ALMACÉN'
    )
    p_nombre_encargado.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Pie de página
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run('ÁREA DE LOGÍSTICA - ALMACÉN MUNICIPAL')
    footer_run.font.size = Pt(8)
    footer_run.bold = True
    
    # Guardar documento
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output
@app.route("/solicitudes_donacion")
@requiere_login('admin')
def solicitudes_donacion():
    conn = conexion()
    try:
        with conn.cursor() as cursor:
            # Obtener autorizantes únicos
            cursor.execute("""
                SELECT DISTINCT autorizado_por 
                FROM Donacion 
                WHERE autorizado_por IS NOT NULL
                ORDER BY autorizado_por
            """)
            autorizantes = [row[0] for row in cursor.fetchall()]
            
            # Consulta principal actualizada para incluir el número secuencial
            cursor.execute("""
                SELECT 
                    d.id_donacion,
                    d.fecha_donacion,
                    p.nombre,
                    p.apellido,
                    p.dni,
                    p.telefono,
                    d.autorizado_por,
                    d.motivo,
                    GROUP_CONCAT(
                        CONCAT(pr.nombre, ' (', dd.cantidad, ' ', pr.unidad_medida, ')')
                        SEPARATOR ', '
                    ) as productos,
                    COALESCE(nd.numero_secuencial, 0) as numero_donacion
                FROM Donacion d
                LEFT JOIN NumeroDonacion nd ON d.id_donacion = nd.id_donacion
                JOIN Prestatario p ON d.id_prestatario = p.id_prestatario
                JOIN DetalleDonacion dd ON d.id_donacion = dd.id_donacion
                JOIN Producto pr ON dd.id_producto = pr.id_producto
                GROUP BY d.id_donacion, d.fecha_donacion, p.nombre, p.apellido, 
                         p.dni, p.telefono, d.autorizado_por, d.motivo, nd.numero_secuencial
                ORDER BY d.fecha_donacion DESC, d.id_donacion DESC
            """)
            donaciones = cursor.fetchall()
            
            return render_template('solicitudes_donacion.html', 
                                 donaciones=donaciones,
                                 autorizantes=autorizantes)
            
    except Exception as e:
        print(f"Error en solicitudes_donacion: {str(e)}")
        flash('Error al cargar las donaciones', 'danger')
        return redirect(url_for('inicio'))
    finally:
        conn.close()
@app.route("/descargar_donaciones")
@requiere_login('admin')
def descargar_donaciones():
    conn = conexion()
    try:
        wb = Workbook()
        ws = wb.active
        ws.title = "REGISTRO DE DONACIONES"
        
        # Estilos
        title_font = Font(name='Arial', size=16, bold=True, color="FFFFFF")
        subtitle_font = Font(name='Arial', size=11, color="666666", bold=True)
        header_font = Font(name='Arial', size=11, bold=True, color="FFFFFF")
        data_font = Font(name='Arial', size=10)
        
        # Bordes
        thin_border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )
        
        # Colores
        title_fill = PatternFill(start_color="2F75B5", end_color="2F75B5", fill_type="solid")
        subtitle_fill = PatternFill(start_color="F2F2F2", end_color="F2F2F2", fill_type="solid")
        header_fill = PatternFill(start_color="305496", end_color="305496", fill_type="solid")
        
        # Consulta SQL mejorada
        with conn.cursor() as cursor:
            cursor.execute("""
                WITH DonacionesNumeradas AS (
                    SELECT 
                        d.id_donacion,
                        COALESCE(nd.numero_secuencial, 
                            (SELECT COUNT(*) + 1 
                             FROM NumeroDonacion nd2 
                             WHERE nd2.fecha_donacion < d.fecha_donacion)
                        ) as numero_donacion,
                        DATE_FORMAT(d.fecha_donacion, '%d/%m/%Y %H:%i') as fecha,
                        CONCAT(p.nombre, ' ', p.apellido) as beneficiario,
                        p.dni,
                        p.telefono,
                        GROUP_CONCAT(
                            CONCAT(
                                pr.codigo, ' - ',
                                pr.nombre, ' (',
                                dd.cantidad, ' ',
                                pr.unidad_medida, ')'
                            )
                            SEPARATOR '\n'
                        ) as productos,
                        d.autorizado_por,
                        d.motivo,
                        d.observacion_autorizacion,
                        GROUP_CONCAT(dd.observacion SEPARATOR '\n') as observaciones_productos
                    FROM Donacion d
                    LEFT JOIN NumeroDonacion nd ON d.id_donacion = nd.id_donacion
                    JOIN Prestatario p ON d.id_prestatario = p.id_prestatario
                    JOIN DetalleDonacion dd ON d.id_donacion = dd.id_donacion
                    JOIN Producto pr ON dd.id_producto = pr.id_producto
                    GROUP BY d.id_donacion, d.fecha_donacion, p.nombre, p.apellido, 
                             p.dni, p.telefono, d.autorizado_por, d.motivo, 
                             d.observacion_autorizacion, nd.numero_secuencial
                )
                SELECT * FROM DonacionesNumeradas
                ORDER BY fecha DESC
            """)
            donaciones = cursor.fetchall()
            
            # Título y encabezados
            ws.merge_cells('A1:J1')
            title_cell = ws['A1']
            title_cell.value = "MUNICIPALIDAD DISTRITAL DE PIÁS"
            title_cell.font = title_font
            title_cell.fill = title_fill
            title_cell.alignment = Alignment(horizontal='center', vertical='center')
            
            ws.merge_cells('A2:J2')
            subtitle_cell = ws['A2']
            subtitle_cell.value = f"REPORTE GENERAL DE DONACIONES - Generado el {datetime.now().strftime('%d/%m/%Y %H:%M')}"
            subtitle_cell.font = subtitle_font
            subtitle_cell.fill = subtitle_fill
            subtitle_cell.alignment = Alignment(horizontal='center', vertical='center')
            
            # Información adicional
            ws.merge_cells('A3:J3')
            info_cell = ws['A3']
            info_cell.value = f"Total de donaciones registradas: {len(donaciones)}"
            info_cell.font = Font(name='Arial', size=10, italic=True)
            info_cell.alignment = Alignment(horizontal='right')
            
            headers = [
                'N° DONACIÓN',
                'FECHA Y HORA',
                'BENEFICIARIO',
                'DNI',
                'TELÉFONO',
                'PRODUCTOS',
                'AUTORIZADO POR',
                'MOTIVO',
                'OBS. AUTORIZACIÓN',
                'OBS. PRODUCTOS'
            ]
            
            # Aplicar estilos a los encabezados
            for col, header in enumerate(headers, 1):
                cell = ws.cell(row=4, column=col, value=header)
                cell.font = header_font
                cell.fill = header_fill
                cell.border = thin_border
                cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
            
            # Datos
            for row_idx, don in enumerate(donaciones, 5):
                row_color = "F5F5F5" if row_idx % 2 == 0 else "FFFFFF"
                
                # Formatear número de donación usando el número secuencial
                ws.cell(row=row_idx, column=1, value=f"DON-{don[1]:04d}")
                
                # Procesar el resto de columnas, excluyendo id_donacion y numero_donacion
                for col_idx, value in enumerate(don[2:], 2):
                    cell = ws.cell(row=row_idx, column=col_idx, value=value)
                    cell.font = data_font
                    cell.fill = PatternFill(start_color=row_color, end_color=row_color, fill_type="solid")
                    cell.border = thin_border
                    cell.alignment = Alignment(horizontal='center' if col_idx < 6 else 'left', 
                                            vertical='center', 
                                            wrap_text=True)
            
            # Ajustar anchos de columna
            column_widths = {
                'A': 15,  # N° DONACIÓN
                'B': 18,  # FECHA Y HORA
                'C': 30,  # BENEFICIARIO
                'D': 12,  # DNI
                'E': 12,  # TELÉFONO
                'F': 40,  # PRODUCTOS
                'G': 25,  # AUTORIZADO POR
                'H': 35,  # MOTIVO
                'I': 35,  # OBS. AUTORIZACIÓN
                'J': 35   # OBS. PRODUCTOS
            }
            
            for col, width in column_widths.items():
                ws.column_dimensions[col].width = width
            
            # Ajustar altura de filas para productos
            for row in range(5, len(donaciones) + 5):
                ws.row_dimensions[row].height = 45
            
            # Pie de página
            last_row = len(donaciones) + 6
            ws.merge_cells(f'A{last_row}:J{last_row}')
            footer_cell = ws[f'A{last_row}']
            footer_cell.value = "ÁREA DE LOGÍSTICA - ALMACÉN MUNICIPAL"
            footer_cell.font = Font(name='Arial', size=10, bold=True)
            footer_cell.alignment = Alignment(horizontal='center')
            
            # Agregar bordes a todas las celdas
            for row in ws.iter_rows(min_row=1, max_row=last_row, min_col=1, max_col=len(headers)):
                for cell in row:
                    cell.border = thin_border
        
        # Guardar archivo
        output = BytesIO()
        wb.save(output)
        output.seek(0)
        
        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=f'reporte_donaciones_{datetime.now().strftime("%Y%m%d_%H%M")}.xlsx'
        )
        
    except Exception as e:
        print(f"Error en descargar_donaciones: {str(e)}")
        flash('Error al generar el reporte de donaciones', 'danger')
        return redirect(url_for('reportes'))
    finally:
        conn.close()

@app.route("/descargar_constancia_donacion/<int:id_donacion>")
@requiere_login('admin')
def descargar_constancia_donacion(id_donacion):
    conn = conexion()
    try:
        with conn.cursor() as cursor:
            # Obtener datos de la donación incluyendo el número secuencial
            cursor.execute("""
                SELECT 
                    d.id_donacion,
                    d.fecha_donacion,
                    COALESCE(nd.numero_secuencial, 0) as numero_donacion,
                    p.nombre,
                    p.apellido,
                    p.dni,
                    p.telefono,
                    d.autorizado_por,
                    d.motivo,
                    d.observacion_autorizacion
                FROM Donacion d
                LEFT JOIN NumeroDonacion nd ON d.id_donacion = nd.id_donacion
                JOIN Prestatario p ON d.id_prestatario = p.id_prestatario
                WHERE d.id_donacion = %s
            """, (id_donacion,))
            donacion = cursor.fetchone()
            
            if not donacion:
                flash('Donación no encontrada', 'danger')
                return redirect(url_for('solicitudes_donacion'))
            
            # Asignar número secuencial si no existe
            numero_donacion = donacion[2]
            if numero_donacion == 0:
                numero_donacion = asignar_numero_secuencial(
                    conn,
                    'NumeroDonacion',
                    'id_donacion',
                    donacion[0],
                    'fecha_donacion',
                    donacion[1]
                )
            
            # Obtener productos de la donación
            cursor.execute("""
                SELECT 
                    pr.codigo,
                    pr.nombre,
                    dd.cantidad,
                    pr.unidad_medida,
                    dd.observacion
                FROM DetalleDonacion dd
                JOIN Producto pr ON dd.id_producto = pr.id_producto
                WHERE dd.id_donacion = %s
            """, (id_donacion,))
            productos = cursor.fetchall()
            
            # Generar documento
            datos = {
                'id_donacion': donacion[0],
                'numero_donacion': numero_donacion,
                'fecha_salida': donacion[1].strftime('%Y-%m-%d'),
                'hora_salida': donacion[1].strftime('%H:%M'),
                'nombre': donacion[3],
                'apellidos': donacion[4],
                'dni': donacion[5],
                'telefono': donacion[6] or 'NO ESPECIFICADO',
                'autorizado_por': donacion[7],
                'motivo': donacion[8],
                'observacion_autorizacion': donacion[9] or ''
            }
            
            productos_info = [{
                'codigo': p[0],
                'nombre': p[1],
                'cantidad': p[2],
                'unidad': p[3],
                'observacion': p[4] or ''
            } for p in productos]
            
            # Generar el documento
            doc = generar_documento_donacion(datos, productos_info)
            
            # Confirmar la transacción si se asignó un nuevo número
            if donacion[2] == 0:
                conn.commit()
            
            # Enviar el archivo
            response = send_file(
                doc,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                as_attachment=True,
                download_name=f'donacion_{str(numero_donacion).zfill(4)}.docx'
            )
            
            # Agregar header para redirección después de la descarga
            response.headers["X-Download-Complete-Redirect"] = url_for('solicitudes_donacion')
            return response
            
    except Exception as e:
        conn.rollback()
        print(f"Error al descargar constancia: {str(e)}")
        flash('Error al generar la constancia', 'danger')
        return redirect(url_for('solicitudes_donacion'))
    finally:
        conn.close()
def generar_excel_donaciones():
    conn = conexion()
    wb = Workbook()
    ws = wb.active
    ws.title = "REGISTRO DE DONACIONES"
    
    try:
        # Estilos
        title_font = Font(name='Arial', size=16, bold=True, color="FFFFFF")
        subtitle_font = Font(name='Arial', size=11, color="666666", bold=True)
        header_font = Font(name='Arial', size=11, bold=True, color="FFFFFF")
        data_font = Font(name='Arial', size=10)
        
        # Bordes
        thin_border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )
        
        # Colores
        title_fill = PatternFill(start_color="2F75B5", end_color="2F75B5", fill_type="solid")
        subtitle_fill = PatternFill(start_color="F2F2F2", end_color="F2F2F2", fill_type="solid")
        header_fill = PatternFill(start_color="305496", end_color="305496", fill_type="solid")
        
        # Consulta SQL mejorada
        with conn.cursor() as cursor:
            cursor.execute("""
                WITH DonacionesNumeradas AS (
                    SELECT 
                        d.id_donacion,
                        COALESCE(nd.numero_secuencial, 
                            (SELECT COUNT(*) + 1 
                             FROM NumeroDonacion nd2 
                             WHERE nd2.fecha_donacion < d.fecha_donacion)
                        ) as numero_donacion,
                        d.fecha_donacion,
                        CONCAT(p.nombre, ' ', p.apellido) as beneficiario,
                        p.dni,
                        p.telefono,
                        pr.codigo,
                        pr.nombre as producto,
                        c.nombre as categoria,
                        dd.cantidad,
                        pr.unidad_medida,
                        pr.procedencia,
                        d.autorizado_por,
                        d.motivo,
                        dd.observacion,
                        d.observacion_autorizacion,
                        CONCAT(u.nombre, ' ', u.apellido) as usuario_registro
                    FROM Donacion d
                    LEFT JOIN NumeroDonacion nd ON d.id_donacion = nd.id_donacion
                    JOIN Prestatario p ON d.id_prestatario = p.id_prestatario
                    JOIN DetalleDonacion dd ON d.id_donacion = dd.id_donacion
                    JOIN Producto pr ON dd.id_producto = pr.id_producto
                    LEFT JOIN Categoria c ON pr.id_categoria = c.id_categoria
                    LEFT JOIN Usuario u ON d.id_usuario = u.id_usuario
                )
                SELECT * FROM DonacionesNumeradas
                ORDER BY fecha_donacion DESC, id_donacion DESC
            """)
            donaciones = cursor.fetchall()
            
            # Título y encabezados
            ws.merge_cells('A1:P1')
            title_cell = ws['A1']
            title_cell.value = "MUNICIPALIDAD DISTRITAL DE PIÁS"
            title_cell.font = title_font
            title_cell.fill = title_fill
            title_cell.alignment = Alignment(horizontal='center', vertical='center')
            
            ws.merge_cells('A2:P2')
            subtitle_cell = ws['A2']
            subtitle_cell.value = f"REPORTE GENERAL DE DONACIONES - Generado el {datetime.now().strftime('%d/%m/%Y %H:%M')}"
            subtitle_cell.font = subtitle_font
            subtitle_cell.fill = subtitle_fill
            subtitle_cell.alignment = Alignment(horizontal='center', vertical='center')
            
            # Información adicional
            ws.merge_cells('A3:P3')
            info_cell = ws['A3']
            info_cell.value = f"Total de registros: {len(donaciones)}"
            info_cell.font = Font(name='Arial', size=10, italic=True)
            info_cell.alignment = Alignment(horizontal='right')
            
            headers = [
                'N° DONACIÓN',
                'FECHA Y HORA',
                'BENEFICIARIO',
                'DNI',
                'TELÉFONO',
                'CÓDIGO',
                'PRODUCTO',
                'CATEGORÍA',
                'CANTIDAD',
                'U.M.',
                'PROCEDENCIA',
                'AUTORIZADO POR',
                'MOTIVO',
                'OBS. PRODUCTO',
                'OBS. AUTORIZACIÓN',
                'REGISTRADO POR'
            ]
            
            # Aplicar estilos a los encabezados
            for col, header in enumerate(headers, 1):
                cell = ws.cell(row=4, column=col, value=header)
                cell.font = header_font
                cell.fill = header_fill
                cell.border = thin_border
                cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
            
            # Datos
            for row_idx, don in enumerate(donaciones, 5):
                row_color = "F5F5F5" if row_idx % 2 == 0 else "FFFFFF"
                
                # Formatear número de donación usando el número secuencial
                ws.cell(row=row_idx, column=1, value=f"DON-{don[1]:04d}")
                
                # Procesar el resto de columnas, excluyendo id_donacion y numero_donacion
                for col_idx, value in enumerate(don[2:], 2):
                    cell = ws.cell(row=row_idx, column=col_idx, value=value)
                    cell.font = data_font
                    cell.fill = PatternFill(start_color=row_color, end_color=row_color, fill_type="solid")
                    cell.border = thin_border
                    cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                    
                    # Formato especial para fechas y números
                    if isinstance(value, datetime):
                        cell.value = value.strftime('%d/%m/%Y %H:%M')
                    elif isinstance(value, float):
                        cell.value = f"{value:,.2f}"
            
            # Ajustar anchos de columna
            column_widths = {
                'A': 15,  # N° DONACIÓN
                'B': 18,  # FECHA Y HORA
                'C': 30,  # BENEFICIARIO
                'D': 12,  # DNI
                'E': 12,  # TELÉFONO
                'F': 12,  # CÓDIGO
                'G': 30,  # PRODUCTO
                'H': 15,  # CATEGORÍA
                'I': 12,  # CANTIDAD 
                'J': 8,   # U.M.
                'K': 15,  # PROCEDENCIA
                'L': 25,  # AUTORIZADO POR
                'M': 35,  # MOTIVO
                'N': 35,  # OBS. PRODUCTO
                'O': 35,  # OBS. AUTORIZACIÓN
                'P': 25   # REGISTRADO POR
            }
            
            for col, width in column_widths.items():
                ws.column_dimensions[col].width = width
            
            # Pie de página
            last_row = len(donaciones) + 6
            ws.merge_cells(f'A{last_row}:P{last_row}')
            footer_cell = ws[f'A{last_row}']
            footer_cell.value = "ÁREA DE LOGÍSTICA - ALMACÉN MUNICIPAL"  
            footer_cell.font = Font(name='Arial', size=10, bold=True)
            footer_cell.alignment = Alignment(horizontal='center')
            
            # Agregar bordes a todas las celdas
            for row in ws.iter_rows(min_row=1, max_row=last_row, min_col=1, max_col=len(headers)):
                for cell in row:
                    cell.border = thin_border
            
    except Exception as e:
        print(f"Error en generar_excel_donaciones: {str(e)}")
        raise
    finally:
        conn.close()
    
    excel_file = BytesIO()
    wb.save(excel_file)
    excel_file.seek(0)
    
    return excel_file


@app.route("/agregar_stock", methods=["POST"])
@requiere_login('admin')
def agregar_stock():
    conn = conexion()
    try:
        data = request.get_json()
        if not data:
            return jsonify({
                "success": False,
                "message": "No se recibieron datos en la solicitud"
            })

        # Obtener datos
        producto_id = data.get('producto_id')
        cantidad = float(data.get('cantidad', 0))
        motivo = data.get('motivo', '')
        observacion = data.get('observacion', '')

        # Validaciones básicas
        if not producto_id or cantidad <= 0:
            return jsonify({
                "success": False,
                "message": "Datos inválidos"
            })

        with conn.cursor() as cursor:
            # Actualizar stock del producto
            cursor.execute("""
                UPDATE Producto 
                SET cantidad = cantidad + %s
                WHERE id_producto = %s
            """, (cantidad, producto_id))

            # Registrar en historial
            cursor.execute("""
                INSERT INTO HistorialModificaciones 
                (id_producto, tipo_modificacion, cantidad, fecha_modificacion, detalle)
                VALUES (%s, 'ingreso', %s, NOW(), %s)
            """, (
                producto_id,
                cantidad,
                f"Ingreso adicional - Motivo: {motivo} - Observación: {observacion}"
            ))

            conn.commit()
            return jsonify({
                "success": True,
                "message": "Stock actualizado correctamente"
            })

    except Exception as e:
        conn.rollback()
        print(f"Error en agregar_stock: {str(e)}")
        return jsonify({
            "success": False,
            "message": str(e)
        })
    finally:
        conn.close()

@app.route("/obtener_registros/<tipo>/<int:producto_id>")
@requiere_login('admin')
def obtener_registros(tipo, producto_id):
    conn = conexion()
    try:
        with conn.cursor() as cursor:
            if tipo == 'donaciones':
                cursor.execute("""
                    SELECT 
                        dd.id_detalle_donacion as id,
                        d.fecha,
                        d.beneficiario as persona,
                        dd.cantidad,
                        d.motivo as detalle
                    FROM DetalleDonacion dd
                    JOIN Donacion d ON dd.id_donacion = d.id_donacion
                    WHERE dd.id_producto = %s
                    ORDER BY d.fecha DESC
                """, (producto_id,))
            
            elif tipo == 'solicitudes':
                cursor.execute("""
                    SELECT 
                        ds.id_detalle_solicitud as id,
                        s.fecha,
                        s.solicitante as persona,
                        ds.cantidad,
                        s.estado as detalle
                    FROM DetalleSolicitud ds
                    JOIN Solicitud s ON ds.id_solicitud = s.id_solicitud
                    WHERE ds.id_producto = %s
                    ORDER BY s.fecha DESC
                """, (producto_id,))
            
            elif tipo == 'prestamos':
                cursor.execute("""
                    SELECT 
                        dp.id_detalle_prestamo as id,
                        p.fecha,
                        p.responsable as persona,
                        dp.cantidad,
                        CASE 
                            WHEN d.id_devolucion IS NULL THEN 'Pendiente'
                            ELSE 'Devuelto'
                        END as detalle
                    FROM DetallePrestamo dp
                    JOIN Prestamo p ON dp.id_prestamo = p.id_prestamo
                    LEFT JOIN Devolucion d ON dp.id_detalle_prestamo = d.id_detalle_prestamo
                    WHERE dp.id_producto = %s
                    ORDER BY p.fecha DESC
                """, (producto_id,))
            
            elif tipo == 'salidas':
                cursor.execute("""
                    SELECT 
                        ds.id_detalle_salida as id,
                        s.fecha,
                        s.destino as persona,
                        ds.cantidad,
                        s.motivo as detalle
                    FROM DetalleSalida ds
                    JOIN Salida s ON ds.id_salida = s.id_salida
                    WHERE ds.id_producto = %s
                    ORDER BY s.fecha DESC
                """, (producto_id,))
            
            elif tipo == 'combustible':
                cursor.execute("""
                    SELECT 
                        dc.id_detalle_combustible as id,
                        sc.fecha,
                        v.placa as persona,
                        dc.cantidad,
                        sc.conductor as detalle
                    FROM DetalleCombustible dc
                    JOIN SalidaCombustible sc ON dc.id_salida = sc.id_salida
                    JOIN Vehiculo v ON sc.id_vehiculo = v.id_vehiculo
                    WHERE dc.id_producto = %s
                    ORDER BY sc.fecha DESC
                """, (producto_id,))
            
            registros = cursor.fetchall()
            registros_formateados = []
            for r in registros:
                registros_formateados.append({
                    'id': r[0],
                    'fecha': r[1].strftime('%d/%m/%Y %H:%M') if r[1] else '',
                    'persona': r[2],
                    'cantidad': float(r[3]),
                    'detalle': r[4]
                })
            
            return jsonify({
                "success": True,
                "registros": registros_formateados
            })
            
    except Exception as e:
        return jsonify({
            "success": False,
            "error": str(e)
        })
    finally:
        conn.close()

@app.route("/eliminar_registros", methods=['POST'])
@requiere_login('admin')
def eliminar_registros():
    conn = conexion()
    try:
        data = request.get_json()
        registros = data.get('registros', {})
        
        with conn.cursor() as cursor:
            for tipo, ids in registros.items():
                if tipo == 'donaciones':
                    cursor.execute("DELETE FROM DetalleDonacion WHERE id_detalle_donacion IN %s", (tuple(ids),))
                elif tipo == 'solicitudes':
                    cursor.execute("DELETE FROM DetalleSolicitud WHERE id_detalle_solicitud IN %s", (tuple(ids),))
                elif tipo == 'prestamos':
                    cursor.execute("DELETE FROM DetallePrestamo WHERE id_detalle_prestamo IN %s", (tuple(ids),))
                elif tipo == 'salidas':
                    cursor.execute("DELETE FROM DetalleSalida WHERE id_detalle_salida IN %s", (tuple(ids),))
                elif tipo == 'combustible':
                    cursor.execute("DELETE FROM DetalleCombustible WHERE id_detalle_combustible IN %s", (tuple(ids),))
        
        conn.commit()
        return jsonify({"success": True, "message": "Registros eliminados correctamente"})
        
    except Exception as e:
        conn.rollback()
        return jsonify({"success": False, "error": str(e)})
    finally:
        conn.close()
@app.route("/obtener_documentos/<tipo>/<int:producto_id>")
@requiere_login('admin')
def obtener_documentos(tipo, producto_id):
    conn = conexion()
    try:
        with conn.cursor() as cursor:
            if tipo == 'prestamos':
                cursor.execute("""
                    SELECT DISTINCT
                        p.id_prestamo as id,
                        p.id_prestamo as numero,
                        p.fecha_prestamo as fecha,
                        CONCAT(p.nombre_solicitante, ' ', p.apellido_solicitante) as responsable,
                        dp.cantidad,
                        CASE 
                            WHEN d.fecha_devolucion IS NOT NULL THEN 'Devuelto'
                            WHEN p.fecha_devolucion_planeada < NOW() THEN 'Vencido'
                            ELSE 'Activo'
                        END as estado
                    FROM Prestamo p
                    JOIN DetallePrestamo dp ON p.id_prestamo = dp.id_prestamo
                    LEFT JOIN Devolucion d ON p.id_prestamo = d.id_prestamo
                    WHERE dp.id_producto = %s
                    ORDER BY p.fecha_prestamo DESC
                """, (producto_id,))
            
            documentos = cursor.fetchall()
            documentos_formateados = []
            for doc in documentos:
                documentos_formateados.append({
                    'id': str(doc[0]),
                    'numero': str(doc[1]),
                    'fecha': doc[2].strftime('%d/%m/%Y %H:%M'),
                    'responsable': doc[3],
                    'cantidad': float(doc[4]),
                    'estado': doc[5]
                })
            
            return jsonify({
                "success": True,
                "documentos": documentos_formateados
            })
            
    except Exception as e:
        return jsonify({
            "success": False,
            "error": str(e)
        })
    finally:
        conn.close()

@app.route("/eliminar_documento", methods=['POST'])
@requiere_login('admin')
def eliminar_documento():
    conn = conexion()
    try:
        data = request.get_json()
        tipo = data.get('tipo')
        id_documento = data.get('id_documento')
        
        with conn.cursor() as cursor:
            if tipo == 'salidas':
                cursor.execute("DELETE FROM DetalleSalida WHERE id_salida = %s", (id_documento,))
                cursor.execute("DELETE FROM Salida WHERE id_salida = %s", (id_documento,))
            elif tipo == 'prestamos':
                cursor.execute("DELETE FROM DetallePrestamo WHERE id_prestamo = %s", (id_documento,))
                cursor.execute("DELETE FROM Prestamo WHERE id_prestamo = %s", (id_documento,))
            elif tipo == 'donaciones':
                cursor.execute("DELETE FROM DetalleDonacion WHERE id_donacion = %s", (id_documento,))
                cursor.execute("DELETE FROM Donacion WHERE id_donacion = %s", (id_documento,))
            elif tipo == 'solicitudes':
                cursor.execute("DELETE FROM DetalleSolicitud WHERE id_solicitud = %s", (id_documento,))
                cursor.execute("DELETE FROM Solicitud WHERE id_solicitud = %s", (id_documento,))
            elif tipo == 'combustible':
                cursor.execute("DELETE FROM DetalleCombustible WHERE id_salida = %s", (id_documento,))
                cursor.execute("DELETE FROM SalidaCombustible WHERE id_salida = %s", (id_documento,))
        
        conn.commit()
        return jsonify({"success": True, "message": "Documento eliminado correctamente"})
        
    except Exception as e:
        conn.rollback()
        return jsonify({"success": False, "error": str(e)})
    finally:
        conn.close()



if __name__ == '__main__':
    app.run(debug=True)

    